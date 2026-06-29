import streamlit as st
from datetime import datetime, time, date
import pandas as pd
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
import logging
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import time as time_module
import textwrap
import base64
import re
from googleapiclient.http import MediaIoBaseUpload
import io
st.cache_data.clear()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

st.set_page_config(page_title="منصة مذكرات الماستر", page_icon="📘", layout="wide")

DEPOSIT_DEADLINE = datetime(2026, 6, 25, 23, 59)  # الدورة الاستدراكية
REGISTRATION_DEADLINE = datetime(2027, 1, 28, 23, 59)

def get_days_remaining():
    delta = DEPOSIT_DEADLINE - datetime.now()
    return max(0, delta.days)

def render_countdown_banner():
    days = get_days_remaining()
    now = datetime.now()
    if now > DEPOSIT_DEADLINE:
        st.markdown("""<div style="background:linear-gradient(135deg,#1a472a,#2d6a4f);border-radius:16px;padding:16px 24px;margin-bottom:20px;display:flex;align-items:center;gap:14px;box-shadow:0 8px 24px rgba(45,106,79,0.3);"><span style="font-size:2rem;">✅</span><div><div style="color:#fff;font-size:1.05rem;font-weight:800;">انتهى أجل إيداع المذكرات</div><div style="color:rgba(255,255,255,0.9);font-size:0.82rem;">25 جوان 2026 (الدورة الاستدراكية)</div></div></div>""", unsafe_allow_html=True)
        return
    if days == 0:
        urgency = "⚡ اليوم الأخير للإيداع!"
        bg = "linear-gradient(135deg,#4A0000,#C0392B)"
        shadow = "rgba(192,57,43,0.6)"
    elif days <= 3:
        urgency = f"⚠️ تبقى {days} أيام فقط — أودع الآن!"
        bg = "linear-gradient(135deg,#7C0A02,#E74C3C)"
        shadow = "rgba(231,76,60,0.5)"
    elif days <= 7:
        urgency = f"📌 تبقى {days} أيام على آخر أجل"
        bg = "linear-gradient(135deg,#8B4513,#E67E22)"
        shadow = "rgba(230,126,34,0.4)"
    else:
        urgency = "📌 آخر أجل لإيداع المذكرات (دورة استدراكية): 25 جوان 2026"
        bg = "linear-gradient(135deg,#1A3A5C,#2F6F7E)"
        shadow = "rgba(47,111,126,0.4)"
    st.markdown(f"""<div style="background:{bg};border-radius:16px;padding:16px 24px;margin-bottom:20px;display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:10px;box-shadow:0 8px 28px {shadow};"><div style="display:flex;align-items:center;gap:12px;"><span style="font-size:1.9rem;">⏳</span><div><div style="color:#fff;font-size:1.05rem;font-weight:800;">{urgency}</div><div style="color:rgba(255,255,255,0.9);font-size:0.8rem;margin-top:2px;">آخر أجل لإيداع المذكرات النهائية (دورة استدراكية): 25 جوان 2026</div></div></div><div style="background:rgba(0,0,0,0.3);border:2px solid rgba(255,255,255,0.4);border-radius:12px;padding:8px 20px;text-align:center;"><div style="font-size:2.6rem;font-weight:900;color:#FFD700;line-height:1;text-shadow:0 0 20px rgba(255,215,0,0.5);">{days}</div><div style="font-size:0.72rem;color:rgba(255,255,255,0.8);letter-spacing:1px;">يوم متبقي</div></div></div>""", unsafe_allow_html=True)


st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap" rel="stylesheet">
<style>
* { box-sizing: border-box; }
html, body, [class*="css"] { font-family: 'Cairo', sans-serif !important; direction: rtl; text-align: right; }
.main { background-color: #0A1B2C; color: #ffffff; }
.block-container { padding: 2rem; background-color: #1A2A3D; border-radius: 16px; margin:auto; }
h1, h2, h3, h4 { font-weight: 700; margin-bottom: 1rem; color: #ffffff !important; }
label, p, span { color: #ffffff !important; }
.stMarkdown div { color: #ffffff; }
.stMarkdown p { color: #ffffff !important; }
h1 { text-align: center; }
.stTextInput label, .stSelectbox label { color: #ffffff !important; font-weight: 600; }
    .stTextInput input, .stTextInput textarea { background-color: #ffffff !important; color: #1A2A3D !important; border: 1px solid #2F6F7E !important; border-radius: 8px !important; font-family: 'Cairo', sans-serif !important; font-size: 15px !important; }
    .stTextInput input::placeholder { color: #94A3B8 !important; }
    input[type="password"] { background-color: #ffffff !important; color: #1A2A3D !important; border: 1px solid #2F6F7E !important; border-radius: 8px !important; }
.stButton>button { background-color: #2F6F7E !important; color: #ffffff !important; font-size: 16px; font-weight: 600; padding: 14px 32px; border: none !important; border-radius: 12px !important; cursor: pointer; box-shadow: 0 4px 12px rgba(0,0,0,0.2); transition: all 0.3s ease; width: 100%; text-align: center; }
.stButton>button:hover { background-color: #285E6B !important; transform: translateY(-2px); }
div[data-testid="stFormSubmitButton"] button { background-color: #2F6F7E !important; color: #ffffff !important; font-size: 16px; font-weight: 600; border: none !important; border-radius: 12px !important; width: 100%; }
div[data-testid="stFormSubmitButton"] button:hover { background-color: #285E6B !important; }
div[data-testid="stFormSubmitButton"] button p { color: #ffffff !important; }
.stButton>button p { color: #ffffff !important; }
.card { background: rgba(30,41,59,0.95); border:1px solid rgba(255,255,255,0.08); border-radius: 20px; padding: 30px; margin-bottom: 20px; box-shadow: 0 20px 25px -5px rgba(0,0,0,0.2); border-top: 3px solid #2F6F7E; }
.kpi-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 1.2rem; margin-bottom: 2rem; }
.kpi-card { background: linear-gradient(145deg,#1E293B,#0F172A); border: 1px solid rgba(255,255,255,0.05); border-radius: 16px; padding: 2rem 1rem; text-align: center; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.3); }
.kpi-value { font-size: 2.2rem; font-weight: 900; color: #FFD700; margin: 10px 0; }
.kpi-label { font-size: 1rem; color: #ffffff !important; font-weight: 600; }
.alert-card { background: linear-gradient(90deg,#8B4513,#A0522D); border: 1px solid #CD853F; color: white; padding: 22px; border-radius: 12px; text-align: center; font-weight: bold; }
.progress-container { background-color: #0F172A; border-radius: 99px; padding: 5px; margin: 14px 0; overflow: hidden; }
.progress-bar { height: 22px; border-radius: 99px; background: linear-gradient(90deg,#2F6F7E,#285E6B,#FFD700); box-shadow: 0 0 15px rgba(47,111,126,0.5); }
.stDataFrame { border-radius: 12px; overflow: hidden; }
.stTabs [data-baseweb="tab-list"] { gap: 1.5rem; padding-bottom: 12px; }
.stTabs [data-baseweb="tab"] { background: transparent; color: #94A3B8; font-weight: 600; padding: 10px 20px; border-radius: 10px; border: 1px solid transparent; }
.stTabs [data-baseweb="tab"]:hover { background: rgba(255,255,255,0.08); color: white; }
div[role="radiogroup"] label { color: #ffffff !important; background: #1E293B !important; border: 1px solid rgba(255,255,255,0.2) !important; border-radius: 20px !important; padding: 4px 14px !important; margin: 2px !important; cursor: pointer !important; }
div[role="radiogroup"] label:hover { background: #2F6F7E !important; border-color: #2F6F7E !important; }
div[role="radiogroup"] label[data-checked="true"], div[role="radiogroup"] input:checked + div { background: #2F6F7E !important; }
div[role="radiogroup"] p { color: #ffffff !important; }
.stRadio > div { flex-direction: row !important; flex-wrap: wrap !important; gap: 6px !important; }
.stSelectbox > div > div { background-color: #1E293B !important; color: #ffffff !important; border: 1px solid rgba(255,255,255,0.15) !important; }
.stSelectbox > div > div > div { color: #ffffff !important; }
.stSelectbox svg { fill: #ffffff !important; }
.stSelectbox [data-baseweb="select"] > div { background-color: #1E293B !important; color: #ffffff !important; }
.stSelectbox [data-baseweb="select"] span { color: #ffffff !important; }
div[data-baseweb="select"] { background-color: #1E293B !important; }
div[data-baseweb="select"] > div { background-color: #1E293B !important; color: #ffffff !important; }
div[data-baseweb="popover"] { background-color: #1E293B !important; }
div[data-baseweb="popover"] li { color: #ffffff !important; background-color: #1E293B !important; }
div[data-baseweb="popover"] li:hover { background-color: #2F6F7E !important; }
.stTabs [aria-selected="true"] { background: rgba(47,111,126,0.2); color: #FFD700; border: 1px solid #2F6F7E; font-weight: bold; }
[data-testid="stExpander"] { background-color: #1E293B !important; border: 1px solid rgba(255,255,255,0.1) !important; border-radius: 10px !important; }
[data-testid="stExpander"] details { background-color: #1E293B !important; }
[data-testid="stExpander"] details summary { background-color: #1E293B !important; color: #ffffff !important; }
[data-testid="stExpander"] details summary:hover { background-color: #263548 !important; color: #FFD700 !important; }
[data-testid="stExpander"] details summary:focus { background-color: #1E293B !important; color: #ffffff !important; outline: none !important; }
[data-testid="stExpander"] details summary p,
[data-testid="stExpander"] details summary span { color: #ffffff !important; background: transparent !important; }
[data-testid="stExpander"] details summary:hover p,
[data-testid="stExpander"] details summary:hover span { color: #FFD700 !important; }
[data-testid="stExpander"] details summary:focus p,
[data-testid="stExpander"] details summary:focus span { color: #ffffff !important; }
[data-testid="stExpander"] details[open] summary { background-color: #1E293B !important; color: #FFD700 !important; }
[data-testid="stExpander"] details[open] summary p,
[data-testid="stExpander"] details[open] summary span { color: #FFD700 !important; }
[data-testid="stExpander"] details[open] > div { background-color: #1A2A3D !important; }
.students-grid { display: flex; justify-content: center; gap: 30px; flex-wrap: wrap; margin: 20px 0; }
.student-card { flex:1; max-width:400px; min-width:260px; background:rgba(255,255,255,0.03); border:1px solid rgba(255,255,255,0.08); border-radius:16px; padding:22px; text-align:center; transition:all 0.3s; }
.student-card:hover { background:rgba(255,255,255,0.06); border-color:#2F6F7E; }
.diploma-item { background:rgba(255,255,255,0.04); padding:13px 17px; border-radius:10px; margin-bottom:9px; display:flex; justify-content:space-between; align-items:center; border-right:4px solid #2F6F7E; }
.status-badge { padding:4px 13px; border-radius:20px; font-size:0.86rem; font-weight:bold; min-width:88px; text-align:center; }
.status-available { background:rgba(16,185,129,0.15); color:#10B981; border:1px solid rgba(16,185,129,0.3); }
.status-unavailable { background:rgba(239,68,68,0.15); color:#EF4444; border:1px solid rgba(239,68,68,0.3); }
.status-pending { background:rgba(245,158,11,0.15); color:#F59E0B; border:1px solid rgba(245,158,11,0.3); }
[data-testid="stFileUploader"] { background:rgba(15,41,66,0.9)!important; border:2px dashed #2F6F7E!important; border-radius:16px!important; padding:18px!important; }
[data-testid="stFileUploader"] label { color:#ffffff!important; font-weight:700!important; font-size:1rem!important; }
[data-testid="stFileUploaderDropzoneInstructions"] span,[data-testid="stFileUploaderDropzoneInstructions"] p,[data-testid="stFileUploaderDropzoneInstructions"] small { color:#E2E8F0!important; font-size:0.9rem!important; }
[data-testid="stFileUploaderDropzone"] { background:rgba(47,111,126,0.08)!important; border-radius:10px!important; }
[data-testid="stFileUploader"] button { background-color:#2F6F7E!important; color:#ffffff!important; border-radius:8px!important; border:none!important; font-weight:600!important; }
[data-testid="stFileUploader"] button:hover { background-color:#285E6B!important; }
[data-testid="stFileUploader"] button p { color:#ffffff!important; }
[data-testid="stFileUploader"] section { background:rgba(15,41,66,0.6)!important; border-radius:10px!important; }
[data-testid="stFileUploader"] section p, [data-testid="stFileUploader"] section span, [data-testid="stFileUploader"] section small { color:#E2E8F0!important; }
.deposit-hero { background:linear-gradient(135deg,#0F2942,#1A3A5C,#0F2942); border:2px solid #2F6F7E; border-radius:22px; padding:30px 34px; margin-bottom:22px; position:relative; overflow:hidden; }
.deposit-hero::before { content:''; position:absolute; top:-60px; left:-60px; width:240px; height:240px; background:radial-gradient(circle,rgba(47,111,126,0.16) 0%,transparent 70%); pointer-events:none; }
.deposit-hero-icon { font-size:3rem; display:block; margin-bottom:10px; }
.deposit-hero-title { font-size:1.5rem; font-weight:900; color:#FFD700!important; margin-bottom:7px; }
.deposit-hero-sub { font-size:0.92rem; color:#E2E8F0!important; line-height:1.6; }
.notif-card { border-radius:16px; padding:18px 22px; margin-bottom:14px; display:flex; gap:16px; align-items:flex-start; }
.notif-card-waiting { background:linear-gradient(135deg,rgba(245,158,11,0.1),rgba(245,158,11,0.04)); border:1.5px solid rgba(245,158,11,0.42); }
.notif-card-approved { background:linear-gradient(135deg,rgba(16,185,129,0.1),rgba(16,185,129,0.04)); border:1.5px solid rgba(16,185,129,0.42); }
.notif-card-rejected { background:linear-gradient(135deg,rgba(239,68,68,0.1),rgba(239,68,68,0.04)); border:1.5px solid rgba(239,68,68,0.42); }
.notif-card-scheduled { background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(99,102,241,0.04)); border:1.5px solid rgba(99,102,241,0.42); }
.notif-icon { font-size:1.9rem; flex-shrink:0; margin-top:2px; }
.notif-title { font-size:1.02rem; font-weight:700; margin-bottom:5px; }
.notif-title-waiting { color:#F59E0B!important; } .notif-title-approved { color:#10B981!important; }
.notif-title-rejected { color:#EF4444!important; } .notif-title-scheduled { color:#818CF8!important; }
.notif-desc { color:#ffffff!important; font-size:0.86rem; line-height:1.7; }
.defense-schedule-card { background:linear-gradient(135deg,rgba(99,102,241,0.1),rgba(99,102,241,0.04)); border:2px solid rgba(99,102,241,0.38); border-radius:18px; padding:24px; margin:18px 0; }
.defense-info-grid { display:grid; grid-template-columns:repeat(auto-fit,minmax(150px,1fr)); gap:12px; margin-top:12px; }
.defense-info-item { background:rgba(255,255,255,0.05); border-radius:10px; padding:13px; text-align:center; }
.defense-info-label { font-size:0.73rem; color:#E2E8F0!important; margin-bottom:4px; }
.defense-info-value { font-size:1.05rem; font-weight:700; color:#818CF8!important; }
.jury-card { background:linear-gradient(145deg,#0F1E30,#162840); border:1px solid rgba(255,255,255,0.07); border-radius:18px; padding:24px; margin:18px 0; }
.jury-header { background:linear-gradient(135deg,#1E3A5F,#2F6F7E); border-radius:11px; padding:14px 18px; margin-bottom:20px; display:flex; align-items:center; gap:12px; }
.jury-header-icon { font-size:1.8rem; }
.jury-header-title { font-size:1rem; font-weight:700; color:#FFD700!important; margin:0; }
.jury-header-sub { font-size:0.8rem; color:rgba(255,255,255,0.85)!important; margin:0; }
.jury-members-grid { display:grid; grid-template-columns:repeat(auto-fit,minmax(180px,1fr)); gap:13px; margin-bottom:16px; }
.jury-member-card { background:rgba(255,255,255,0.04); border:1px solid rgba(255,255,255,0.07); border-radius:13px; padding:16px; text-align:center; transition:all 0.3s; }
.jury-member-card:hover { background:rgba(255,255,255,0.07); border-color:rgba(47,111,126,0.38); transform:translateY(-2px); }
.jury-member-avatar { width:48px; height:48px; border-radius:50%; margin:0 auto 9px; display:flex; align-items:center; justify-content:center; font-size:1.25rem; }
.avatar-president { background:rgba(255,215,0,0.1); border:2px solid rgba(255,215,0,0.38); }
.avatar-supervisor { background:rgba(47,111,126,0.16); border:2px solid rgba(47,111,126,0.38); }
.avatar-examiner { background:rgba(148,163,184,0.08); border:2px solid rgba(148,163,184,0.18); }
.jury-member-role { font-size:0.7rem; font-weight:700; letter-spacing:0.7px; padding:3px 9px; border-radius:20px; display:inline-block; margin-bottom:7px; }
.role-president { background:rgba(255,215,0,0.1); color:#FFD700!important; border:1px solid rgba(255,215,0,0.28); }
.role-supervisor { background:rgba(47,111,126,0.14); color:#2F9EA0!important; border:1px solid rgba(47,111,126,0.28); }
.role-examiner { background:rgba(148,163,184,0.08); color:#E2E8F0!important; border:1px solid rgba(148,163,184,0.18); }
.jury-member-name { font-size:0.9rem; font-weight:700; color:#ffffff!important; line-height:1.4; }
.prof-deposit-alert { background:linear-gradient(135deg,#0D2010,#0F2020); border:2px solid rgba(16,185,129,0.42); border-radius:18px; margin-bottom:22px; overflow:hidden; box-shadow:0 8px 28px rgba(16,185,129,0.1); }
.prof-deposit-alert-header { background:linear-gradient(135deg,rgba(16,185,129,0.17),rgba(16,185,129,0.04)); padding:17px 21px; display:flex; align-items:center; gap:13px; border-bottom:1px solid rgba(16,185,129,0.17); }
.prof-deposit-alert-icon { font-size:1.9rem; }
.prof-deposit-alert-title { font-size:1.1rem; font-weight:800; color:#10B981!important; margin:0; }
.prof-deposit-alert-sub { font-size:0.8rem; color:#ffffff!important; margin:0; }
.prof-deposit-list { padding:13px 19px; }
.prof-deposit-item { background:rgba(255,255,255,0.03); border:1px solid rgba(16,185,129,0.17); border-radius:11px; padding:13px 17px; margin-bottom:9px; display:flex; justify-content:space-between; align-items:center; flex-wrap:wrap; gap:7px; transition:background 0.2s; }
.prof-deposit-item:hover { background:rgba(16,185,129,0.07); }
.prof-deposit-memo-num { font-size:1.2rem; font-weight:900; color:#FFD700!important; }
.prof-deposit-memo-title { font-size:0.86rem; color:#E2E8F0!important; margin-top:2px; }
.prof-deposit-memo-date { font-size:0.76rem; color:#E2E8F0!important; margin-top:3px; }
.declaration-card { background:linear-gradient(145deg,#0A1628,#0F1E30); border:1px solid rgba(255,255,255,0.08); border-radius:20px; overflow:hidden; margin:17px 0; }
.declaration-card-header { background:linear-gradient(135deg,#1E3A5F,#0F2942); padding:19px 24px; border-bottom:1px solid rgba(255,255,255,0.07); }
.declaration-card-title { font-size:1.15rem; font-weight:800; color:#FFD700!important; margin:0 0 3px; }
.declaration-card-sub { font-size:0.8rem; color:rgba(255,255,255,0.8)!important; margin:0; }
.declaration-card-body { padding:22px 24px; }
.declaration-step-label { font-size:0.8rem; font-weight:600; color:#ffffff!important; margin-bottom:7px; letter-spacing:0.3px; }
.declaration-preview { background:rgba(255,255,255,0.03); border:1px solid rgba(255,255,255,0.07); border-right:4px solid #2F6F7E; border-radius:9px; padding:13px 15px; font-size:0.85rem; color:#ffffff!important; line-height:1.9; font-style:italic; }
.declaration-preview strong { color:#FFD700!important; font-style:normal; }
.quick-memo-btn {
    background: linear-gradient(135deg,rgba(245,158,11,0.12),rgba(245,158,11,0.05));
    border: 1.5px solid rgba(245,158,11,0.4);
    border-radius: 14px;
    padding: 16px 20px;
    margin-bottom: 10px;
    cursor: pointer;
    transition: all 0.25s;
}
.quick-memo-btn:hover {
    background: linear-gradient(135deg,rgba(245,158,11,0.22),rgba(245,158,11,0.1));
    border-color: rgba(245,158,11,0.7);
    transform: translateY(-2px);
}
</style>
""", unsafe_allow_html=True)


SCOPES_SHEETS = ['https://www.googleapis.com/auth/spreadsheets']
SCOPES_DRIVE  = ['https://www.googleapis.com/auth/drive']

try:
    info = st.secrets["service_account"]
    credentials = Credentials.from_service_account_info(info, scopes=SCOPES_SHEETS)
    sheets_service = build('sheets', 'v4', credentials=credentials)
except Exception as e:
    st.error("⚠️ خطأ في الاتصال بـ Google Sheets"); st.stop()

DRIVE_UPLOAD_FOLDER_ID = "1fvckcOGegVD4Ofs-UnVZbVCbYBQlToWs"
try:
    drive_info = st.secrets["drive_service_account"]
    drive_credentials = Credentials.from_service_account_info(drive_info, scopes=SCOPES_DRIVE)
    drive_service = build('drive', 'v3', credentials=drive_credentials)
except:
    drive_service = None

STUDENTS_SHEET_ID   = "1_uhwdyOERttICR6uqftU9eyJrmetpA6idL3vz5WtWX4"
MEMOS_SHEET_ID      = "1LNJMBAye4QIQy7JHz6F8mQ6-XNC1weZx1ozDZFfjD5s"
PROF_MEMOS_SHEET_ID = "1OnZi1o-oPMUI_W_Ew-op0a1uOhSj006hw_2jrMD6FSE"
REQUESTS_SHEET_ID   = "1sTJ6BZRM4Qgt0w2xUkpFZqquL-hfriMYTSN3x1_12_o"
STUDENTS_RANGE  = "Feuille 1!A1:V1000"  # K=البريد المهني included
MEMOS_RANGE     = "Feuille 1!A1:AI1000"
PROF_MEMOS_RANGE= "Feuille 1!A1:S1000"
REQUESTS_RANGE  = "Feuille 1!A1:K1000"
ADMIN_CREDENTIALS = {"admin": "admin2026", "dsp": "dsp@2026"}
PRINTER_CREDENTIALS  = {"mem": "1234"}   # فضاء المحاضر فقط
LIBRARY_CREDENTIALS  = {"bib": "1234"}   # مسؤول المكتبة — غيّر كلمة السر هنا
IDAA_FOLDER_ID       = "12zHRa5dDfGTotK0QBkvN5xUQC5bTZMQV"  # مجلد الإيداع النهائي
EMAIL_SENDER  = "domaine.dsp@univ-bba.dz"
MAHDAR_FOLDER_ID = "1T8ABqZeSOwBBbMwpLRyxyRNrf0diYzJB"  # مجلد Drive لحفظ محاضر المناقشة
EMAIL_PASSWORD= "oqwejylusjllfvhc"
SMTP_SERVER   = "smtp.gmail.com"
SMTP_PORT     = 587
ADMIN_EMAIL   = "domaine.dsp@univ-bba.dz"

def format_arabic_date(date_input):
    try:
        if isinstance(date_input, str): date_obj = datetime.strptime(date_input, '%Y-%m-%d %H:%M:%S')
        elif isinstance(date_input, datetime): date_obj = date_input
        else: return str(date_input)
        months_map = {1:"جانفي",2:"فيفري",3:"مارس",4:"أفريل",5:"ماي",6:"جوان",7:"جويلية",8:"أوت",9:"سبتمبر",10:"أكتوبر",11:"نوفمبر",12:"ديسمبر"}
        return f"{date_obj.day:02d} {months_map.get(date_obj.month,'')} {date_obj.year}"
    except: return str(date_input)

def col_letter(n):
    result = ""
    while n > 0:
        n, remainder = divmod(n-1, 26)
        result = chr(65+remainder) + result
    return result

def sanitize_input(text):
    if not text: return ""
    cleaned = str(text).strip()
    for c in ['<','>','"',"'",'&','|','`']: cleaned = cleaned.replace(c,'')
    return cleaned

def normalize_text(val):
    v = str(val).strip()
    if v.endswith('.0'): v = v[:-2]
    return v

def validate_username(u):
    u = sanitize_input(u)
    if not u: return False, "⚠️ فارغ"
    return True, u

def validate_note_number(n):
    n = sanitize_input(n)
    if not n: return False, "⚠️ فارغ"
    if len(n)>20: return False, "⚠️ غير صالح"
    return True, n

def is_phone_valid(phone_val):
    val = str(phone_val).strip()
    if val in ['','0','-','nan','None','NaN','.0','0.0']: return False, "فارغ"
    cleaned = val.replace(' ','').replace('-','')
    if not cleaned.isdigit(): return False, "ليس أرقام"
    if len(cleaned) < 9: return False, "قصير"
    return True, "صالح"

def is_nin_valid(nin_val):
    val = str(nin_val).strip().replace('.0','')
    if val in ['','0','-','nan','None','NaN']: return False, "فارغ"
    if not val.isdigit(): return False, "ليس أرقام"
    return True, "صالح"

def get_student_name_display(student_dict):
    lname = ""
    for k in ["اللقب","لقب","Nom","nom"]:
        val = str(student_dict.get(k,"")).strip()
        if val and val not in ['nan','None','-']: lname = val; break
    fname = ""
    for k in ["الإسم","إسم","اسم","الاسم","Prénom","prenom"]:
        val = str(student_dict.get(k,"")).strip()
        if val and val not in ['nan','None','-']: fname = val; break
    return lname, fname

def get_email_smart(row):
    if isinstance(row, dict):
        priority_keys = ["الإيميل", "البريد المهني", "البريد الإلكتروني", "email", "Email", "E-mail"]
        for key in priority_keys:
            val = str(row.get(key, "")).strip()
            if "@" in val and val != "nan": return val
        for val in row.values():
            v = str(val).strip()
            if "@" in v and v != "nan": return v
        return ""
    try:
        values_list = row.tolist()
        for i in range(9, 13):
            if i < len(values_list):
                val = str(values_list[i]).strip()
                if "@" in val and val != "nan": return val
        for col in row.index:
            clean_col = str(col).strip()
            if clean_col in ["الإيميل", "البريد المهني", "البريد الإلكتروني", "email", "Email", "E-mail"]:
                val = str(row[col]).strip()
                if "@" in val and val != "nan": return val
        return ""
    except:
        return ""

def load_student2_for_memo(memo_row, current_student_reg, df_students):
    memo_list = memo_row.tolist() if hasattr(memo_row,'tolist') else list(memo_row.values())
    reg_s = normalize_text(memo_row.get("رقم تسجيل الطالب 1", memo_list[18] if len(memo_list)>18 else ""))
    reg_t = normalize_text(memo_row.get("رقم تسجيل الطالب 2", memo_list[19] if len(memo_list)>19 else ""))
    current_reg = normalize_text(current_student_reg)
    other_reg = reg_t if current_reg == reg_s else (reg_s if current_reg == reg_t else reg_t)
    if not other_reg or other_reg in ["","nan"]: return None
    df_students['رقم التسجيل_norm'] = df_students['رقم التسجيل'].astype(str).apply(normalize_text)
    s2 = df_students[df_students["رقم التسجيل_norm"] == other_reg]
    return s2.iloc[0].to_dict() if not s2.empty else None

def get_student_info_from_memo(memo_row, df_students):
    s1_name = str(memo_row.get("الطالب الأول","")).strip()
    s2_name = str(memo_row.get("الطالب الثاني","")).strip()
    s1_email = s2_email = s1_reg = s2_reg = ""
    try:
        reg1 = normalize_text(str(memo_row.get("رقم تسجيل الطالب 1","")).strip())
        reg2 = normalize_text(str(memo_row.get("رقم تسجيل الطالب 2","")).strip())
    except Exception:
        reg1 = ""; reg2 = ""
    if df_students is not None and not df_students.empty and "رقم التسجيل" in df_students.columns:
        df_students = df_students.copy()
        df_students['رقم التسجيل_norm'] = df_students['رقم التسجيل'].astype(str).apply(normalize_text)
        if reg1:
            s = df_students[df_students["رقم التسجيل_norm"]==reg1]
            if not s.empty: s1_email = get_email_smart(s.iloc[0]); s1_reg = reg1
        if s2_name and reg2:
            s = df_students[df_students["رقم التسجيل_norm"]==reg2]
            if not s.empty: s2_email = get_email_smart(s.iloc[0]); s2_reg = reg2
    return {"s1_name":s1_name,"s1_email":s1_email,"s1_reg":s1_reg,"s2_name":s2_name,"s2_email":s2_email,"s2_reg":s2_reg}

@st.cache_data(ttl=60)
def load_students():
    try:
        result = sheets_service.spreadsheets().values().get(spreadsheetId=STUDENTS_SHEET_ID, range=STUDENTS_RANGE).execute()
        values = result.get('values',[])
        if not values: return pd.DataFrame()
        df = pd.DataFrame(values[1:], columns=values[0]); df.columns = df.columns.str.strip(); return df
    except Exception as e: logger.error(f"خطأ الطلاب: {e}"); return pd.DataFrame()

@st.cache_data(ttl=60)
def load_memos():
    try:
        result = sheets_service.spreadsheets().values().get(spreadsheetId=MEMOS_SHEET_ID, range="Feuille 1!A1:AV1000").execute()
        values = result.get('values',[])
        if not values: return pd.DataFrame()
        headers = values[0]; rows = values[1:]
        padded = [r+['']*(len(headers)-len(r)) for r in rows]
        return pd.DataFrame(padded, columns=headers)
    except Exception as e: logger.error(f"خطأ المذكرات: {e}"); return pd.DataFrame()

@st.cache_data(ttl=60)
def upload_to_drive(file_bytes, filename, folder_id, mimetype='application/pdf'):
    """رفع ملف على Google Drive"""
    try:
        from googleapiclient.discovery import build as _gbuild
        from googleapiclient.http import MediaIoBaseUpload as _MIU
        import io as _io_d
        drive_service = _gbuild('drive', 'v3', credentials=drive_credentials)
        file_meta = {'name': filename, 'parents': [folder_id]}
        media = _MIU(_io_d.BytesIO(file_bytes), mimetype=mimetype, resumable=True)
        f = drive_service.files().create(body=file_meta, media_body=media, fields='id,webViewLink').execute()
        drive_service.permissions().create(fileId=f['id'], body={'type':'anyone','role':'reader'}).execute()
        return True, f['webViewLink']
    except Exception as e:
        return False, str(e)

def upload_mahdar_pdf(file_bytes, filename, folder_id=MAHDAR_FOLDER_ID):
    """رفع PDF محضر المناقشة على Google Drive"""
    try:
        from googleapiclient.discovery import build as _gbuild
        from googleapiclient.http import MediaIoBaseUpload as _MIU
        import io as _io_d
        drive_service = _gbuild('drive', 'v3', credentials=drive_credentials)
        file_meta = {'name': filename, 'parents': [folder_id]}
        media = _MIU(_io_d.BytesIO(file_bytes), mimetype='application/pdf', resumable=True)
        f = drive_service.files().create(body=file_meta, media_body=media, fields='id,webViewLink').execute()
        # جعل الملف قابلاً للعرض للجميع
        drive_service.permissions().create(fileId=f['id'], body={'type':'anyone','role':'reader'}).execute()
        return True, f['webViewLink']
    except Exception as e:
        return False, str(e)

def load_prof_memos():
    try:
        result = sheets_service.spreadsheets().values().get(spreadsheetId=PROF_MEMOS_SHEET_ID, range=PROF_MEMOS_RANGE).execute()
        values = result.get('values',[])
        if not values: return pd.DataFrame()
        return pd.DataFrame(values[1:], columns=values[0])
    except Exception as e: logger.error(f"خطأ الأساتذة: {e}"); return pd.DataFrame()

@st.cache_data(ttl=60)
def load_requests():
    try:
        result = sheets_service.spreadsheets().values().get(spreadsheetId=REQUESTS_SHEET_ID, range=REQUESTS_RANGE).execute()
        values = result.get('values',[])
        if not values: return pd.DataFrame()
        return pd.DataFrame(values[1:], columns=values[0])
    except: return pd.DataFrame()

def clear_cache_and_reload(): st.cache_data.clear()


def _email_style():
    return """<style>body{font-family:Arial,sans-serif;background:#f4f4f4;padding:20px;direction:rtl;text-align:right;}.container{background:#fff;padding:28px;border-radius:12px;box-shadow:0 2px 10px rgba(0,0,0,.1);max-width:600px;margin:auto;}.header{background:linear-gradient(135deg,#0F2942,#2F6F7E);color:#fff;padding:20px;border-radius:8px;text-align:center;margin-bottom:18px;}.header h2{margin:0;font-size:1.3rem;}.info-box{background:#f0f9ff;padding:14px;border-right:4px solid #2F6F7E;margin:12px 0;border-radius:6px;}.action-box{background:#fff8e1;padding:14px;border-right:4px solid #F59E0B;margin:12px 0;border-radius:6px;}.success-box{background:#f0fdf4;padding:14px;border-right:4px solid #10B981;margin:12px 0;border-radius:6px;}.warning-box{background:#fff1f2;padding:14px;border-right:4px solid #EF4444;margin:12px 0;border-radius:6px;}.reject-box{background:#fff1f2;padding:14px;border-right:4px solid #EF4444;margin:12px 0;border-radius:6px;}.platform-btn{display:inline-block;background:#2F6F7E;color:#fff!important;padding:12px 28px;border-radius:8px;text-decoration:none;font-weight:bold;margin-top:10px;}.footer{text-align:center;color:#888;font-size:12px;margin-top:24px;border-top:1px solid #eee;padding-top:12px;}p{color:#333;line-height:1.8;}</style>"""


# ================================================================
# ✅ التعديل 2: دالة إرسال إيميل الإيداع مع بيانات الدخول
# ================================================================
def send_deposit_email_to_professor(prof_name, memo_number, memo_title, student1_name, student2_name=None):
    try:
        df_profs = load_prof_memos()
        prof_rows = df_profs[
            (df_profs["الأستاذ"].astype(str).str.strip()==prof_name.strip()) &
            (df_profs["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number))
        ]
        if prof_rows.empty:
            all_prof_rows = df_profs[df_profs["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
            prof_rows = all_prof_rows
        if prof_rows.empty: return False, f"الأستاذ غير موجود في قاعدة البيانات: {prof_name}"

        prof_email = ""
        prof_username = ""
        prof_password_val = ""

        for _, prow in prof_rows.iterrows():
            # البريد الإلكتروني
            email_candidate = get_email_smart(prow)
            if email_candidate and "@" in email_candidate:
                prof_email = email_candidate

            # اسم المستخدم — من شيت الأساتذة فقط
            for col in ["إسم المستخدم", "اسم المستخدم", "Identifiant"]:
                if col in prow.index:
                    val = str(prow.get(col, "")).strip()
                    if val and val.lower() != "nan":
                        prof_username = val; break

            # كلمة المرور — من شيت الأساتذة فقط
            for col in ["كلمة المرور", "كلمة السر", "Password"]:
                if col in prow.index:
                    val = str(prow.get(col, "")).strip()
                    if val and val.lower() != "nan":
                        prof_password_val = val; break

            if prof_email and prof_username and prof_password_val:
                break

        if not prof_email: return False, f"البريد غير متوفر للأستاذ: {prof_name}"

        students_html = f"<p>👤 <strong>الطالب الأول:</strong> {student1_name}</p>"
        if student2_name and student2_name.strip() and student2_name != student1_name:
            students_html += f"<p>👤 <strong>الطالب الثاني:</strong> {student2_name}</p>"

        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')

        # بناء صندوق بيانات الدخول
        login_html = ""
        if prof_username or prof_password_val:
            login_html = f"""
            <div style="background:#f0fdf4;padding:16px;border-right:4px solid #10B981;margin:14px 0;border-radius:6px;">
                <p style="margin:0 0 10px;font-weight:800;color:#065f46;font-size:1rem;">
                    🔐 تذكير ببيانات الدخول إلى المنصة
                </p>
                <p style="margin:6px 0;color:#333;">
                    👤 <strong>اسم المستخدم:</strong>&nbsp;
                    <span style="font-family:monospace;background:#dbeafe;padding:3px 10px;border-radius:5px;
                          font-size:1rem;color:#1e3a5f;font-weight:700;letter-spacing:0.5px;">
                        {prof_username if prof_username else '—'}
                    </span>
                </p>
                <p style="margin:6px 0;color:#333;">
                    🔑 <strong>كلمة المرور:</strong>&nbsp;
                    <span style="font-family:monospace;background:#dbeafe;padding:3px 10px;border-radius:5px;
                          font-size:1rem;color:#1e3a5f;font-weight:700;letter-spacing:0.5px;">
                        {prof_password_val if prof_password_val else '—'}
                    </span>
                </p>
            </div>"""

        body = f"""<html dir="rtl"><head><meta charset="UTF-8">{_email_style()}</head><body>
        <div class="container" dir="rtl">
            <div class="header">
                <h2>📥 إيداع مذكرة — للإطلاع والمراجعة</h2>
                <p style="color:rgba(255,255,255,0.8);margin:5px 0 0;font-size:0.88rem;">
                    جامعة محمد البشير الإبراهيمي — كلية الحقوق
                </p>
            </div>
            <p>الأستاذ(ة) الفاضل(ة) <strong>{prof_name}</strong>،</p>
            <p>نحيطكم علماً بأن الطالب(ين) أودعوا نسختهم النهائية من المذكرة رقم
                <strong>{memo_number}</strong> عبر منصة مذكرات الماستر بتاريخ <strong>{timestamp}</strong>.
            </p>
            <div class="info-box">
                <p>📄 <strong>رقم المذكرة:</strong> {memo_number}</p>
                <p>📑 <strong>العنوان:</strong> {memo_title}</p>
                {students_html}
            </div>
            <div class="action-box">
                <p>يُرجى الدخول إلى المنصة للاطلاع على المذكرة المودعة،
                ثم اتخاذ قرار الموافقة أو الإعادة مع الملاحظات.</p>
            </div>
            {login_html}
            <div style="text-align:center;margin:20px 0;">
                <a href="https://memoires2026.streamlit.app" class="platform-btn">
                    🔗 الدخول إلى منصة المذكرات
                </a>
            </div>
            <div class="footer">
                <p>مسؤول الميدان: البروفيسور لخضر رفاف</p>
                <p>جامعة محمد البشير الإبراهيمي</p>
            </div>
        </div></body></html>"""

        for recipient in [prof_email, ADMIN_EMAIL]:
            msg = MIMEMultipart('alternative')
            msg['From']=EMAIL_SENDER; msg['To']=recipient
            msg['Subject']=f"📥 إيداع مذكرة للمراجعة — رقم {memo_number}"
            msg.attach(MIMEText(body,'html','utf-8'))
            with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
                server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅ تم إرسال الإشعار"
    except Exception as e: return False, f"❌ {str(e)}"


def send_approval_email_to_students(memo_number, memo_title, prof_name, student1_data, student2_data=None):
    try:
        recipients = []; names = []
        for s in [student1_data, student2_data]:
            if s is None: continue
            e = get_email_smart(s)
            if e: recipients.append(e)
            ln, fn = get_student_name_display(s); names.append(f"{ln} {fn}".strip())
        if not recipients: return False, "لا يوجد بريد"
        names_str = " و ".join(names)
        body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header"><h2>🟢 مذكرتك معتمدة — قابلة للمناقشة</h2></div><p>تحية طيبة، {names_str}،</p><p>يسعدنا إعلامكم بأن الأستاذ المشرف <strong>{prof_name}</strong> وافق على مذكرتكم رسمياً.</p><div class="success-box"><p>📄 <strong>رقم المذكرة:</strong> {memo_number}</p><p>📑 <strong>العنوان:</strong> {memo_title}</p><p>✅ <strong>الحالة:</strong> قابلة للمناقشة</p></div><p>ستتلقون إشعاراً من الإدارة بموعد ومكان المناقشة.</p><div style="text-align:center;margin:16px 0;"><a href="https://memoires2026.streamlit.app" class="platform-btn">🔗 متابعة على المنصة</a></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=recipients[0]
        if len(recipients)>1: msg['Cc']=", ".join(recipients[1:])
        msg['Bcc']=ADMIN_EMAIL
        msg['Subject']=f"🟢 مذكرتك معتمدة — رقم {memo_number}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅ تم إرسال إشعار الموافقة"
    except Exception as e: return False, f"❌ {str(e)}"

def send_rejection_email_to_students(memo_number, memo_title, prof_name, reason, s1_data, s2_data=None):
    try:
        recipients = []; names = []
        for s in [s1_data, s2_data]:
            if s is None: continue
            e = get_email_smart(s)
            if e: recipients.append(e)
            ln, fn = get_student_name_display(s); names.append(f"{ln} {fn}".strip())
        if not recipients: return False, "لا يوجد بريد"
        names_str = " و ".join(names)
        body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header" style="background:linear-gradient(135deg,#450a0a,#EF4444);"><h2>🔴 المذكرة بحاجة لمراجعة</h2></div><p>تحية طيبة، {names_str}،</p><p>أعاد الأستاذ المشرف <strong>{prof_name}</strong> مذكرتكم للمراجعة.</p><div class="info-box"><p>📄 <strong>رقم المذكرة:</strong> {memo_number}</p><p>📑 <strong>العنوان:</strong> {memo_title}</p></div><div class="reject-box"><p><strong>📋 ملاحظات المشرف وسبب الإعادة:</strong></p><p style="font-style:italic;color:#C0392B;">{reason}</p></div><div class="action-box"><p>⚠️ يُرجى مراجعة الملاحظات وإعادة رفع النسخة المُصححة عبر المنصة.</p></div><div style="text-align:center;margin:16px 0;"><a href="https://memoires2026.streamlit.app" class="platform-btn">🔗 الدخول للمنصة</a></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=recipients[0]
        if len(recipients)>1: msg['Cc']=", ".join(recipients[1:])
        msg['Bcc']=ADMIN_EMAIL
        msg['Subject']=f"🔴 ملاحظات على مذكرتك — رقم {memo_number}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅ تم إرسال إشعار الإعادة"
    except Exception as e: return False, f"❌ {str(e)}"

def send_notes_email_to_students(memo_number, memo_title, prof_name, notes, s1_data, s2_data=None):
    try:
        recipients = []; names = []
        for s in [s1_data, s2_data]:
            if s is None: continue
            e = get_email_smart(s)
            if e: recipients.append(e)
            ln, fn = get_student_name_display(s); names.append(f"{ln} {fn}".strip())
        if not recipients: return False, "لا يوجد بريد"
        names_str = " و ".join(names)
        body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header" style="background:linear-gradient(135deg,#1e1b4b,#6366F1);"><h2>📝 ملاحظات المشرف</h2></div><p>تحية طيبة، {names_str}،</p><p>أرسل الأستاذ المشرف <strong>{prof_name}</strong> ملاحظاته على مذكرتكم رقم <strong>{memo_number}</strong>.</p><div style="background:#f0f0ff;padding:14px;border-right:4px solid #6366F1;margin:12px 0;border-radius:6px;"><p><strong>💬 الملاحظات:</strong></p><p style="font-style:italic;color:#4338CA;">{notes}</p></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=recipients[0]
        if len(recipients)>1: msg['Cc']=", ".join(recipients[1:])
        msg['Subject']=f"📝 ملاحظات المشرف — رقم {memo_number}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅ تم إرسال الملاحظات"
    except Exception as e: return False, f"❌ {str(e)}"

def send_defense_schedule_email(memo_number, memo_title, prof_name, defense_date, defense_time, defense_room, s1_data, s2_data=None):
    try:
        recipients = []; names = []
        for s in [s1_data, s2_data]:
            if s is None: continue
            e = get_email_smart(s)
            if e: recipients.append(e)
            ln, fn = get_student_name_display(s); names.append(f"{ln} {fn}".strip())
        if not recipients: return False, "لا يوجد بريد"
        names_str = " و ".join(names)
        body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header"><h2>📅 موعد مناقشة مذكرتك</h2></div><p>تحية طيبة، {names_str}،</p><p>تم تحديد موعد مناقشتكم رسمياً.</p><div class="info-box"><p>📄 <strong>رقم المذكرة:</strong> {memo_number}</p><p>📑 <strong>العنوان:</strong> {memo_title}</p><p>👨‍🏫 <strong>المشرف:</strong> {prof_name}</p></div><div class="success-box"><p>📆 <strong>التاريخ:</strong> <span style="font-size:1.05rem;color:#10B981;">{defense_date}</span></p><p>🕐 <strong>التوقيت:</strong> <span style="font-size:1.05rem;color:#10B981;">{defense_time}</span></p><p>🏛️ <strong>القاعة:</strong> <span style="font-size:1.05rem;color:#10B981;">{defense_room}</span></p></div><div class="warning-box"><p>⚠️ يرجى الحضور قبل الموعد بـ 15 دقيقة مع جميع الوثائق.</p></div><div style="text-align:center;margin:16px 0;"><a href="https://memoires2026.streamlit.app" class="platform-btn">🔗 متابعة على المنصة</a></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=recipients[0]
        if len(recipients)>1: msg['Cc']=", ".join(recipients[1:])
        msg['Bcc']=ADMIN_EMAIL
        msg['Subject']=f"📅 موعد مناقشتك — رقم {memo_number}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅ تم إرسال موعد المناقشة"
    except Exception as e: return False, f"❌ {str(e)}"

def _send_email_to_professor_row(row):
    prof_name = row.get("الأستاذ","")
    email = ""; username = ""; password = ""
    for col in ["البريد الإلكتروني","الإيميل","email","Email","E-mail"]:
        if col in row.index:
            val = str(row[col]).strip()
            if "@" in val and val != "nan": email = val; break
    for col in ["إسم المستخدم","اسم المستخدم","Identifiant"]:
        if col in row.index:
            val = str(row[col]).strip()
            if val != "nan" and val: username = val; break
    for col in ["كلمة المرور","كلمة السر","Password"]:
        if col in row.index:
            val = str(row[col]).strip()
            if val != "nan" and val: password = val; break
    if not email or not username or not password: return False, "⚠️ بيانات ناقصة"
    body = f"""<html dir="rtl"><head><style>body{{font-family:Arial,sans-serif;direction:rtl;background:#f4f4f4;padding:20px;}}.container{{background:#fff;padding:28px;border-radius:12px;max-width:600px;margin:auto;}}.header{{background:linear-gradient(135deg,#003366,#005580);color:white;padding:18px;border-radius:8px;text-align:center;margin-bottom:18px;}}.info-box{{background:#eef7fb;border-right:5px solid #005580;padding:18px;margin:14px 0;border-radius:4px;}}.footer{{text-align:center;color:#666;font-size:12px;margin-top:24px;border-top:1px solid #eee;padding-top:14px;}}p{{color:#333;line-height:1.7;}}</style></head><body><div class="container"><div class="header"><h2>جامعة محمد البشير الإبراهيمي</h2><h3>كلية الحقوق والعلوم السياسية</h3></div><p>تحية طيبة، الأستاذ(ة) <strong>{prof_name}</strong>،</p><p>تم تفعيل فضاء الأساتذة على منصة متابعة مذكرات الماستر.</p><div class="info-box"><p>🔗 <a href="https://memoires2026.streamlit.app">https://memoires2026.streamlit.app</a></p><p>👤 <strong>اسم المستخدم:</strong> {username}</p><p>🔑 <strong>كلمة المرور:</strong> {password}</p></div><div class="footer"><p>مسؤول الميدان: البروفيسور لخضر رفاف</p></div></div></body></html>"""
    try:
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=email
        msg['Subject']="تفعيل حساب فضاء الأساتذة - منصة المذكرات"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, f"✅ تم الإرسال إلى {email}"
    except Exception as e: return False, f"❌ {str(e)}"

def send_welcome_emails_to_all_profs():
    df_profs = load_prof_memos(); sent=0; failed=0; logs=[]
    pb = st.progress(0); total = len(df_profs)
    with st.spinner("⏳ جاري الإرسال..."):
        for i,(_,row) in enumerate(df_profs.iterrows()):
            ok,msg = _send_email_to_professor_row(row)
            if ok: sent+=1
            else: failed+=1
            logs.append(msg); pb.progress((i+1)/total); time_module.sleep(0.5)
    return sent, failed, logs

def send_welcome_email_to_one(prof_name):
    df_profs = load_prof_memos()
    rows = df_profs[df_profs["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
    if rows.empty: return False, f"❌ لم يتم العثور على: {prof_name}"
    return _send_email_to_professor_row(rows.iloc[0])

def send_email_to_professor(prof_name, memo_info, student1, student2=None):
    df_profs = load_prof_memos()
    prof_row = df_profs[df_profs["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
    if prof_row.empty: return False, "لم يُعثر على الأستاذ"
    prof_email = get_email_smart(prof_row.iloc[0])
    if "@" not in prof_email: return False, "البريد فارغ"
    s1_ln,s1_fn = get_student_name_display(student1)
    s2_info = ""
    if student2:
        s2_ln,s2_fn = get_student_name_display(student2)
        s2_info = f"<p>👤 <strong>الطالب الثاني:</strong> {s2_ln} {s2_fn}</p>"
    body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header"><h2>✅ تسجيل مذكرة جديدة</h2></div><p>الأستاذ(ة) <strong>{prof_name}</strong>،</p><div class="info-box"><p>📄 <strong>رقم المذكرة:</strong> {memo_info['رقم المذكرة']}</p><p>📑 <strong>العنوان:</strong> {memo_info['عنوان المذكرة']}</p><p>🎓 <strong>التخصص:</strong> {memo_info['التخصص']}</p><p>👤 <strong>الطالب الأول:</strong> {s1_ln} {s1_fn}</p>{s2_info}<p>🕒 <strong>تاريخ التسجيل:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
    try:
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=prof_email
        msg['Subject']=f"✅ تسجيل مذكرة رقم {memo_info['رقم المذكرة']}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "تم الإرسال"
    except Exception as e: return False, f"❌ {str(e)}"

def send_session_emails(students_data, session_info, prof_name):
    df_s = load_students(); df_s['رقم التسجيل_norm'] = df_s['رقم التسجيل'].astype(str).apply(normalize_text)
    emails = []
    for s in students_data:
        row = df_s[df_s["رقم التسجيل_norm"]==s['reg']]
        if not row.empty:
            e = get_email_smart(row.iloc[0])
            if e: emails.append(e)
    body = f"""<html dir="rtl"><head>{_email_style()}</head><body><div class="container"><div class="header"><h2>📅 ملاحظة من المشرف</h2></div><p>الأستاذ(ة) <strong>{prof_name}</strong> يُعلن عن ملاحظة من المشرف:</p><div class="success-box"><p><strong>📆 الموعد:</strong> {session_info}</p></div><div class="footer"><p>جامعة محمد البشير الإبراهيمي</p></div></div></body></html>"""
    try:
        msg = MIMEMultipart('alternative')
        msg['From']=EMAIL_SENDER; msg['To']=ADMIN_EMAIL
        if emails: msg['Bcc']=", ".join(emails)
        msg['Subject']=f"🔔 ملاحظة من المشرف — {prof_name}"
        msg.attach(MIMEText(body,'html','utf-8'))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "تم"
    except Exception as e: return False, str(e)


def upload_memo_to_drive(pdf_bytes, memo_number, memo_title):
    if drive_service is None: return False, "", "❌ Drive غير متاح"
    try:
        safe_title = re.sub(r'[\\/:*?"<>|]','',str(memo_title).strip())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # البحث عن النسخ القديمة لإعادة تسميتها كنسخ احتياطية
        existing = drive_service.files().list(
            q=f"'{DRIVE_UPLOAD_FOLDER_ID}' in parents and name contains '{memo_number}.' and not name contains '_v' and trashed=false",
            fields="files(id,name)"
        ).execute()
        for f in existing.get('files',[]):
            old_name = f.get('name','')
            # إعادة تسمية الملف القديم كنسخة احتياطية فقط إذا كان نفس المذكرة
            if old_name.startswith(f"{memo_number}."):
                backup_name = old_name.replace('.pdf', f'_v{timestamp}.pdf')
                drive_service.files().update(
                    fileId=f['id'],
                    body={'name': backup_name}
                ).execute()

        # رفع الملف الجديد
        file_name = f"{memo_number}.{safe_title}.pdf"
        media = MediaIoBaseUpload(io.BytesIO(pdf_bytes), mimetype='application/pdf', resumable=True)
        uploaded = drive_service.files().create(
            body={'name': file_name, 'parents': [DRIVE_UPLOAD_FOLDER_ID]},
            media_body=media, fields='id,webViewLink'
        ).execute()
        file_id = uploaded.get('id')
        drive_service.permissions().create(fileId=file_id, body={'type':'anyone','role':'reader'}).execute()
        link = uploaded.get('webViewLink', f"https://drive.google.com/file/d/{file_id}/view")
        return True, link, "✅ تم رفع الملف"
    except Exception as e: return False, "", f"❌ {str(e)}"

def save_memo_deposit(memo_number, file_link):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":[{"range":f"Feuille 1!T{row_idx}","values":[["مودعة"]]},{"range":f"Feuille 1!U{row_idx}","values":[[file_link]]},{"range":f"Feuille 1!V{row_idx}","values":[[timestamp]]}]}).execute()
        clear_cache_and_reload(); return True, "✅ تم حفظ الإيداع"
    except Exception as e: return False, f"❌ {str(e)}"

def save_approval_declaration(memo_number, prof_name, signature, declaration_text):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().update(spreadsheetId=MEMOS_SHEET_ID, range=f"Feuille 1!Z{row_idx}", valueInputOption="USER_ENTERED", body={"values":[[declaration_text]]}).execute()
        return True, "✅ تم حفظ التصريح"
    except Exception as e: return False, f"❌ {str(e)}"

def approve_memo_for_defense(memo_number):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().update(spreadsheetId=MEMOS_SHEET_ID, range=f"Feuille 1!T{row_idx}", valueInputOption="USER_ENTERED", body={"values":[["قابلة للمناقشة"]]}).execute()
        clear_cache_and_reload(); return True, "✅ تمت الموافقة"
    except Exception as e: return False, f"❌ {str(e)}"

def reject_memo_and_reopen(memo_number, prof_name, rejection_reason):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        rejection_full = f"مرفوضة بتاريخ {timestamp} | المشرف: {prof_name} | السبب: {rejection_reason}"
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":[{"range":f"Feuille 1!T{row_idx}","values":[["مرفوضة"]]},{"range":f"Feuille 1!U{row_idx}","values":[[""]]},{"range":f"Feuille 1!V{row_idx}","values":[[""]]},{"range":f"Feuille 1!Z{row_idx}","values":[[rejection_full]]}]}).execute()
        clear_cache_and_reload(); return True, "✅ تم تسجيل الإعادة وفتح الإيداع"
    except Exception as e: return False, f"❌ {str(e)}"

def save_prof_notes(memo_number, prof_name, notes_text):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        note_full = f"ملاحظات المشرف {prof_name} [{timestamp}]: {notes_text}"
        sheets_service.spreadsheets().values().update(spreadsheetId=MEMOS_SHEET_ID, range=f"Feuille 1!Z{row_idx}", valueInputOption="USER_ENTERED", body={"values":[[note_full]]}).execute()
        clear_cache_and_reload(); return True, "✅ تم حفظ الملاحظات"
    except Exception as e: return False, f"❌ {str(e)}"

def save_defense_schedule(memo_number, defense_date, defense_time, defense_room):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":[{"range":f"Feuille 1!W{row_idx}","values":[[str(defense_date)]]},{"range":f"Feuille 1!X{row_idx}","values":[[str(defense_time)]]},{"range":f"Feuille 1!Y{row_idx}","values":[[defense_room]]}]}).execute()
        clear_cache_and_reload(); return True, "✅ تم حفظ الموعد"
    except Exception as e: return False, f"❌ {str(e)}"

def save_jury(memo_number, president, exam1, exam2):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":[{"range":f"Feuille 1!AA{row_idx}","values":[[president]]},{"range":f"Feuille 1!AB{row_idx}","values":[[exam1]]},{"range":f"Feuille 1!AC{row_idx}","values":[[exam2]]}]}).execute()
        clear_cache_and_reload(); return True, "✅ تم حفظ اللجنة"
    except Exception as e: return False, f"❌ {str(e)}"

def save_notes_by_member(memo_number, member_role, notes_text):
    col_map = {"رئيس لجنة":"AE","مناقش 1":"AF","مناقش 2":"AG"}
    col = col_map.get(member_role)
    if not col: return False, "دور غير معروف"
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().update(spreadsheetId=MEMOS_SHEET_ID, range=f"Feuille 1!{col}{row_idx}", valueInputOption="USER_ENTERED", body={"values":[[notes_text]]}).execute()
        clear_cache_and_reload(); return True, "✅ تم حفظ الملاحظات"
    except Exception as e: return False, f"❌ {str(e)}"

def publish_memos(memo_numbers=None):
    try:
        df_memos = load_memos()
        if memo_numbers:
            targets = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text).isin([normalize_text(m) for m in memo_numbers])]
        else:
            col = "حالة الإيداع"
            targets = df_memos[df_memos[col].astype(str).str.strip()=="قابلة للمناقشة"] if col in df_memos.columns else pd.DataFrame()
        if targets.empty: return False, "لا توجد مذكرات"
        updates = [{"range":f"Feuille 1!AD{idx+2}","values":[["نعم"]]} for idx in targets.index]
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates}).execute()
        clear_cache_and_reload(); return True, f"✅ تم نشر {len(updates)} مذكرة"
    except Exception as e: return False, f"❌ {str(e)}"

def update_progress(memo_number, progress_value):
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0]+2
        sheets_service.spreadsheets().values().update(spreadsheetId=MEMOS_SHEET_ID, range=f"Feuille 1!Q{row_idx}", valueInputOption="USER_ENTERED", body={"values":[[str(progress_value)]]}).execute()
        clear_cache_and_reload(); return True, "✅ تم تحديث نسبة التقدم"
    except Exception as e: return False, f"❌ {str(e)}"

def save_and_send_request(req_type, prof_name, memo_id, memo_title, details_text, status="قيد المراجعة"):
    try:
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        new_row = ["", timestamp, req_type, status, prof_name, memo_id, "", "", details_text, "", ""]
        sheets_service.spreadsheets().values().append(spreadsheetId=REQUESTS_SHEET_ID, range="Feuille 1!A2", valueInputOption="USER_ENTERED", body={"values":[new_row]}, insertDataOption="INSERT_ROWS").execute()
        return True, "✅ تم تسجيل الطلب"
    except Exception as e: return False, f"❌ {str(e)}"

def update_student_profile(username, phone, nin):
    try:
        df_students = load_students()
        df_students['username_norm'] = df_students["اسم المستخدم"].astype(str).apply(normalize_text)
        student_row = df_students[df_students["username_norm"]==normalize_text(username)]
        if student_row.empty: return False, "❌ لم يتم العثور على الطالب"
        row_idx = student_row.index[0]+2
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=STUDENTS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":[{"range":f"Feuille 1!M{row_idx}","values":[[phone]]},{"range":f"Feuille 1!U{row_idx}","values":[[nin]]}]}).execute()
        clear_cache_and_reload(); return True, "✅ تم التحديث"
    except Exception as e: return False, f"❌ {str(e)}"

def update_session_date_in_sheets(prof_name, date_str):
    try:
        df_memos = load_memos()
        masks = (df_memos["الأستاذ"].astype(str).str.strip()==prof_name) & (df_memos["تم التسجيل"].astype(str).str.strip()=="نعم")
        target_indices = df_memos[masks].index
        if target_indices.empty: return True, "لا توجد مذكرات"
        col_names = df_memos.columns.tolist()
        target_col = "موعد الجلسة القادمة"
        col_idx = col_names.index(target_col)+1 if target_col in col_names else len(col_names)
        col_l = col_letter(col_idx)
        updates = [{"range":f"Feuille 1!{col_l}{idx+2}","values":[[date_str]]} for idx in target_indices]
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates}).execute()
        return True, "تم التحديث"
    except Exception as e: return False, str(e)

def sync_student_registration_numbers():
    try:
        df_s = load_students(); df_m = load_memos(); updates = []
        df_s['رقم المذكرة_norm'] = df_s['رقم المذكرة'].astype(str).apply(normalize_text)
        students_with_memo = df_s[df_s['رقم المذكرة_norm'].notna() & (df_s['رقم المذكرة_norm']!="")]
        for index, row in df_m.iterrows():
            memo_num = normalize_text(row.get("رقم المذكرة",""))
            if not memo_num: continue
            matched = students_with_memo[students_with_memo["رقم المذكرة_norm"]==memo_num]
            if matched.empty: continue
            s1_name = str(row.get("الطالب الأول","")).strip(); s2_name = str(row.get("الطالب الثاني","")).strip()
            reg_s1=""; reg_s2=""
            for _, s_row in matched.iterrows():
                ln = s_row.get('لقب', s_row.get('اللقب','')); fn = s_row.get('إسم', s_row.get('الإسم',''))
                full = f"{ln} {fn}".strip()
                if full==s1_name: reg_s1 = normalize_text(s_row.get("رقم التسجيل",""))
                elif s2_name and full==s2_name: reg_s2 = normalize_text(s_row.get("رقم التسجيل",""))
            if not reg_s1 and len(matched)>0: reg_s1 = normalize_text(matched.iloc[0].get("رقم التسجيل",""))
            row_idx = index+2
            if reg_s1: updates.append({"range":f"Feuille 1!S{row_idx}","values":[[reg_s1]]})
            if reg_s2: updates.append({"range":f"Feuille 1!T{row_idx}","values":[[reg_s2]]})
        if updates:
            sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates}).execute()
            return True, f"✅ تم تحديث {len(updates)} خلية"
        return False, "ℹ️ لا توجد تغييرات"
    except Exception as e: return False, f"❌ {str(e)}"

def update_diploma_status(username, status_dict):
    try:
        df_students = load_students()
        df_students['username_norm'] = df_students["اسم المستخدم"].astype(str).apply(normalize_text)
        row = df_students[df_students["username_norm"]==normalize_text(username)]
        if row.empty: return False, "❌ لم يتم العثور على الطالب"
        row_idx = row.index[0]+2
        updates = [{"range":f"Feuille 1!{k}{row_idx}","values":[[v]]} for k,v in status_dict.items()]
        if updates:
            sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=STUDENTS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates}).execute()
            clear_cache_and_reload(); return True, "✅ تم التحديث"
        return False, "لا شيء"
    except Exception as e: return False, f"❌ {str(e)}"

def get_students_of_professor(prof_name, df_memos):
    prof_memos = df_memos[(df_memos["الأستاذ"].astype(str).str.strip()==prof_name.strip()) & (df_memos["تم التسجيل"].astype(str).str.strip()=="نعم")]
    result = []
    for _, memo in prof_memos.iterrows():
        s1_name = str(memo.get("الطالب الأول","")).strip(); s1_reg = normalize_text(memo.get("رقم تسجيل الطالب 1",""))
        if s1_name and s1_name!="--" and s1_reg: result.append({"name":s1_name,"reg":s1_reg,"memo":memo.get("رقم المذكرة")})
        s2_name = str(memo.get("الطالب الثاني","")).strip(); s2_reg = normalize_text(memo.get("رقم تسجيل الطالب 2",""))
        if s2_name and s2_name!="--" and s2_reg: result.append({"name":s2_name,"reg":s2_reg,"memo":memo.get("رقم المذكرة")})
    return result

def format_datetime_ar(date_obj, time_str):
    days_ar = ["الاثنين","الثلاثاء","الأربعاء","الخميس","الجمعة","السبت","الأحد"]
    return f"{days_ar[date_obj.weekday()]} {date_obj.strftime('%Y-%m-%d')} الساعة {time_str}"

def verify_student(username, password, df_students):
    valid, result = validate_username(username)
    if not valid: return False, result
    username = result; password = sanitize_input(password)
    if df_students.empty or "اسم المستخدم" not in df_students.columns: return False, "❌ خطأ في البيانات"
    db = df_students["اسم المستخدم"].astype(str).apply(normalize_text)
    student = df_students[db==username]
    if student.empty: return False, "❌ اسم المستخدم غير موجود"
    if str(student.iloc[0]["كلمة السر"]).strip() != password: return False, "❌ كلمة السر غير صحيحة"
    return True, student.iloc[0].to_dict()

def verify_professor(username, password, df_prof_memos):
    username = sanitize_input(username); password = sanitize_input(password)
    if df_prof_memos.empty or "إسم المستخدم" not in df_prof_memos.columns: return False, "❌ خطأ"
    db = df_prof_memos["إسم المستخدم"].astype(str).apply(normalize_text)
    mask = (db==username) & (df_prof_memos["كلمة المرور"].astype(str).str.strip()==password)
    prof = df_prof_memos[mask]
    if prof.empty: return False, "❌ بيانات غير صحيحة"
    return True, prof.iloc[0].to_dict()

def verify_admin(username, password):
    username = sanitize_input(username); password = sanitize_input(password)
    if username in ADMIN_CREDENTIALS and ADMIN_CREDENTIALS[username]==password: return True, username
    return False, "❌ بيانات غير صحيحة"

def verify_professor_password(note_number, prof_password, df_memos, df_prof_memos):
    valid, result = validate_note_number(note_number)
    if not valid: return False, None, result
    note_number = result; prof_password = sanitize_input(prof_password)
    if df_memos.empty or df_prof_memos.empty: return False, None, "❌ خطأ في البيانات"
    row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==note_number]
    if row.empty: return False, None, "❌ رقم المذكرة غير موجود"
    row = row.iloc[0]
    if str(row.get("تم التسجيل","")).strip()=="نعم": return False, None, "❌ هذه المذكرة مسجلة مسبقاً"
    prof_row = df_prof_memos[(df_prof_memos["الأستاذ"].astype(str).str.strip()==row["الأستاذ"].strip()) & (df_prof_memos["كلمة سر التسجيل"].astype(str).str.strip()==prof_password)]
    if prof_row.empty: return False, None, "❌ كلمة سر المشرف غير صحيحة"
    return True, prof_row.iloc[0].to_dict(), None


def update_registration(note_number, student1, student2=None, s2_new_phone=None, s2_new_nin=None):
    try:
        df_memos = load_memos(); df_prof_memos = load_prof_memos(); df_students = load_students()
        memo_data = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==str(note_number).strip()]
        if memo_data.empty: return False, "❌ رقم المذكرة غير موجود"
        prof_name = memo_data["الأستاذ"].iloc[0].strip()
        used_pw = st.session_state.prof_password.strip()
        potential = df_prof_memos[(df_prof_memos["الأستاذ"].astype(str).str.strip()==prof_name) & (df_prof_memos["كلمة سر التسجيل"].astype(str).str.strip()==used_pw)]
        if potential.empty: return False, "❌ بيانات الأستاذ غير متطابقة"
        target = potential[potential["رقم المذكرة"].astype(str).apply(normalize_text)==str(note_number).strip()]
        if target.empty:
            target = potential[potential["تم التسجيل"].astype(str).str.strip()!="نعم"]
            if target.empty: return False, "❌ جميع المذكرات مسجلة"
        prof_row_idx = target.index[0]+2
        col_names = df_prof_memos.columns.tolist()
        s1_ln, s1_fn = get_student_name_display(student1)
        updates_prof = [
            {"range":f"Feuille 1!{col_letter(col_names.index('الطالب الأول')+1)}{prof_row_idx}","values":[[f"{s1_ln} {s1_fn}"]]},
            {"range":f"Feuille 1!{col_letter(col_names.index('تم التسجيل')+1)}{prof_row_idx}","values":[["نعم"]]},
            {"range":f"Feuille 1!{col_letter(col_names.index('تاريخ التسجيل')+1)}{prof_row_idx}","values":[[datetime.now().strftime('%Y-%m-%d %H:%M')]]},
            {"range":f"Feuille 1!{col_letter(col_names.index('رقم المذكرة')+1)}{prof_row_idx}","values":[[note_number]]}
        ]
        if student2:
            s2_ln, s2_fn = get_student_name_display(student2)
            updates_prof.append({"range":f"Feuille 1!{col_letter(col_names.index('الطالب الثاني')+1)}{prof_row_idx}","values":[[f"{s2_ln} {s2_fn}"]]})
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=PROF_MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates_prof}).execute()
        memo_row_idx = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==str(note_number).strip()].index[0]+2
        memo_cols = df_memos.columns.tolist()
        reg1 = normalize_text(student1.get('رقم التسجيل',''))
        updates_memo = [
            {"range":f"Feuille 1!{col_letter(memo_cols.index('الطالب الأول')+1)}{memo_row_idx}","values":[[f"{s1_ln} {s1_fn}"]]},
            {"range":f"Feuille 1!{col_letter(memo_cols.index('تم التسجيل')+1)}{memo_row_idx}","values":[["نعم"]]},
            {"range":f"Feuille 1!{col_letter(memo_cols.index('تاريخ التسجيل')+1)}{memo_row_idx}","values":[[datetime.now().strftime('%Y-%m-%d %H:%M')]]},
            {"range":f"Feuille 1!S{memo_row_idx}","values":[[reg1]]}
        ]
        if 'كلمة سر التسجيل' in memo_cols:
            updates_memo.append({"range":f"Feuille 1!{col_letter(memo_cols.index('كلمة سر التسجيل')+1)}{memo_row_idx}","values":[[used_pw]]})
        if student2:
            s2_ln2, s2_fn2 = get_student_name_display(student2); reg2 = normalize_text(student2.get('رقم التسجيل',''))
            updates_memo.append({"range":f"Feuille 1!{col_letter(memo_cols.index('الطالب الثاني')+1)}{memo_row_idx}","values":[[f"{s2_ln2} {s2_fn2}"]]})
            updates_memo.append({"range":f"Feuille 1!T{memo_row_idx}","values":[[reg2]]})
        sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=MEMOS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":updates_memo}).execute()
        students_cols = df_students.columns.tolist()
        s1_user = normalize_text(student1.get('اسم المستخدم',''))
        s1_row_idx = df_students[df_students["اسم المستخدم"].astype(str).apply(normalize_text)==s1_user].index[0]+2
        sheets_service.spreadsheets().values().update(spreadsheetId=STUDENTS_SHEET_ID, range=f"Feuille 1!{col_letter(students_cols.index('رقم المذكرة')+1)}{s1_row_idx}", valueInputOption="USER_ENTERED", body={"values":[[note_number]]}).execute()
        if student2:
            s2_user = normalize_text(student2.get('اسم المستخدم',''))
            s2_row_idx = df_students[df_students["اسم المستخدم"].astype(str).apply(normalize_text)==s2_user].index[0]+2
            sheets_service.spreadsheets().values().update(spreadsheetId=STUDENTS_SHEET_ID, range=f"Feuille 1!{col_letter(students_cols.index('رقم المذكرة')+1)}{s2_row_idx}", valueInputOption="USER_ENTERED", body={"values":[[note_number]]}).execute()
            s2_upd = []
            if s2_new_phone: s2_upd.append({"range":f"Feuille 1!M{s2_row_idx}","values":[[s2_new_phone]]})
            if s2_new_nin: s2_upd.append({"range":f"Feuille 1!U{s2_row_idx}","values":[[s2_new_nin]]})
            if s2_upd: sheets_service.spreadsheets().values().batchUpdate(spreadsheetId=STUDENTS_SHEET_ID, body={"valueInputOption":"USER_ENTERED","data":s2_upd}).execute()
        time_module.sleep(2); clear_cache_and_reload(); time_module.sleep(1)
        df_students_updated = load_students()
        st.session_state.student1 = df_students_updated[df_students_updated["اسم المستخدم"].astype(str).apply(normalize_text)==s1_user].iloc[0].to_dict()
        if student2:
            s2_user2 = normalize_text(student2.get('اسم المستخدم',''))
            st.session_state.student2 = df_students_updated[df_students_updated["اسم المستخدم"].astype(str).apply(normalize_text)==s2_user2].iloc[0].to_dict()
        memo_d = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==str(note_number).strip()].iloc[0]
        send_email_to_professor(prof_name, memo_d, st.session_state.student1, st.session_state.get('student2'))
        return True, "✅ تم تسجيل المذكرة بنجاح!"
    except Exception as e:
        logger.error(f"Registration error: {e}"); return False, f"❌ {str(e)}"

def encode_str(s): return base64.urlsafe_b64encode(s.encode()).decode()
def decode_str(s):
    try: return base64.urlsafe_b64decode(s.encode()).decode()
    except: return ""


# ================================================================
# 🎓 خوارزمية جدولة المناقشات
# ================================================================


# ================================================================
# 🧠 خوارزمية الجدولة الذكية — DSatur + Tabu Search
# ================================================================


# ================================================================
# 🧠 خوارزمية الجدولة الذكية v3 — Professor-First Algorithm
# ================================================================

def get_memo_members(row):
    """استخراج كل أعضاء لجنة المذكرة"""
    members = set()
    for col in ["الأستاذ", "الرئيس", "المناقش1", "المناقش2"]:
        v = str(row.get(col, "")).strip()
        if v and v not in ["", "nan", "—", "None"]:
            members.add(v)
    return members

def build_prof_memo_map(df_memos):
    """بناء خريطة الأساتذة ومذكراتهم"""
    prof_memos = {}  # prof -> list of memo_ids
    memo_members = {}  # memo_id -> set of profs
    
    for _, row in df_memos.iterrows():
        mid = str(row["رقم المذكرة"])
        members = get_memo_members(row)
        memo_members[mid] = members
        for prof in members:
            if prof not in prof_memos:
                prof_memos[prof] = []
            if mid not in prof_memos[prof]:
                prof_memos[prof].append(mid)
    
    return prof_memos, memo_members

def build_conflict_matrix(df_memos):
    """بناء مصفوفة التعارض"""
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()
    _, memo_members = build_prof_memo_map(df_memos)
    
    conflicts = {m: set() for m in memo_ids}
    for i in range(len(memo_ids)):
        for j in range(i+1, len(memo_ids)):
            mi, mj = memo_ids[i], memo_ids[j]
            if memo_members.get(mi, set()) & memo_members.get(mj, set()):
                conflicts[mi].add(mj)
                conflicts[mj].add(mi)
    return memo_ids, conflicts, memo_members

def plan_prof_days(prof_memos_count, days, max_per_day=3, max_consecutive=3):
    """
    تخطيط توزيع مذكرات الأستاذ على الأيام
    القواعد:
    - 3 مذكرات حد أقصى في اليوم
    - لا أكثر من 3 أيام متتالية ثم راحة
    - أقل عدد أيام ممكن
    """
    plan = []  # list of (day_index, slots_count)
    remaining = prof_memos_count
    consecutive = 0
    day_idx = 0
    
    while remaining > 0 and day_idx < len(days):
        if consecutive >= max_consecutive:
            # راحة إلزامية
            day_idx += 1
            consecutive = 0
            continue
        
        today = min(remaining, max_per_day)
        plan.append((day_idx, today))
        remaining -= today
        consecutive += 1
        day_idx += 1
    
    return plan

def professor_first_schedule(df_memos, days, slots_per_day, rooms, max_per_day=3, max_consecutive=3,
                             fixed_slots=None, memo_date_limits=None, prof_banned_days=None,
                             prof_not_before=None, prof_not_after=None, prof_one_day=None,
                             prof_allowed_days=None, prof_consecutive=None,
                             frozen_profs=None, prof_phase_split=None,
                             memo_alt_days=None, day_time_limits=None,
                             profs_accept_18=None,
                             profs_cluster_days=None):
    fixed_slots = fixed_slots or {}
    memo_date_limits = memo_date_limits or {}
    prof_banned_days = prof_banned_days or {}
    prof_not_before = prof_not_before or {}
    prof_not_after = prof_not_after or {}
    prof_one_day = prof_one_day or set()
    prof_allowed_days = prof_allowed_days or {}
    prof_consecutive = prof_consecutive or set()
    prof_phase_split = prof_phase_split or {}
    memo_alt_days = memo_alt_days or {}
    day_time_limits = day_time_limits or {}
    profs_accept_18 = profs_accept_18 or set()
    profs_cluster_days = profs_cluster_days or set()
    profs_cluster_days = profs_cluster_days or set()
    LATE_SLOT = "18:00"
    profs_accept_18 = profs_accept_18 or set()
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    """
    الخوارزمية الرئيسية — تبدأ بالأستاذ وليس بالمذكرة
    
    المبدأ:
    1. رتّب الأساتذة من الأكثر مذكرات إلى الأقل
    2. لكل أستاذ: خطط له كتلة أيام وحصص متتالية
    3. ضع مذكراته في الحصص المخصصة له
    4. تحقق من التعارضات مع الأساتذة الآخرين
    5. أي مذكرة لم تُجدول → ابحث عن أي خانة فارغة
    """
    import random
    
    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()
    
    # الجدول النهائي: memo_id -> (day, slot, room)
    schedule = {}
    
    # خريطة الخانات المشغولة: (day, slot, room) -> memo_id
    occupied = {}
    
    # خريطة الأستاذ في التوقيت: (day, slot, prof) -> memo_id
    prof_busy = {}
    # عدد مناقشات الأستاذ في اليوم: (prof, day) -> count
    prof_day_count = {}
    # تتبع أسباب رفض المذكرات
    rejection_log = {}
    # تتبع أسباب رفض المذكرات: memo_id -> [reasons]
    rejection_log = {}
    
    def can_place(memo_id, day, slot, room, log=False):
        """هل يمكن وضع المذكرة في هذه الخانة؟"""
        def _reject(reason):
            if log: rejection_log.setdefault(memo_id, set()).add(reason)
            return False

        if (day, slot, room) in occupied:
            return _reject(f"القاعة {room} مشغولة {day} {slot}")

        # لا بعد 16:00 للجميع — إلا من يقبل 18:00
        if slot > "16:00":
            _members_late = memo_members.get(memo_id, set())
            _non_acc = _members_late - profs_accept_18
            if _non_acc: return _reject(f"توقيت {slot} بعد 16:00 — غير مقبول من: {', '.join(_non_acc)}")
        if slot == LATE_SLOT:
            members_check = memo_members.get(memo_id, set())
            non_acc = members_check - profs_accept_18
            if non_acc: return _reject(f"18:00 غير مقبول من: {', '.join(non_acc)}")

        if day in day_time_limits:
            from_t, to_t = day_time_limits[day]
            if from_t and slot_to_idx.get(slot,0) < slot_to_idx.get(from_t,0):
                return _reject(f"{day}: {slot} قبل الحد {from_t}")
            if to_t and slot_to_idx.get(slot,99) > slot_to_idx.get(to_t,99):
                return _reject(f"{day}: {slot} بعد الحد {to_t}")

        if memo_id in memo_date_limits:
            earliest, latest = memo_date_limits[memo_id]
            if earliest and day < earliest: return _reject(f"لا قبل {earliest}")
            if latest and day > latest: return _reject(f"لا بعد {latest}")

        if memo_id in memo_alt_days and memo_alt_days[memo_id]:
            if day not in memo_alt_days[memo_id]:
                return _reject(f"{day} ليس من الأيام البديلة")

        members = memo_members.get(memo_id, set())
        for prof in members:
            if (day, slot, prof) in prof_busy:
                return _reject(f"{prof} مشغول {day} {slot}")
            if prof_day_count.get((prof, day), 0) >= 3:
                return _reject(f"{prof} بلغ الحد الأقصى 3 في {day}")
            if day in prof_banned_days.get(prof, set()):
                return _reject(f"{day} ممنوع على {prof}")
            if prof in prof_allowed_days and prof_allowed_days[prof]:
                if day not in prof_allowed_days[prof]:
                    return _reject(f"{day} غير مسموح لـ{prof} (مسموح: {','.join(sorted(prof_allowed_days[prof]))})")
            if slot == "18:00" and prof not in profs_accept_18:
                return _reject(f"{prof} لا يقبل 18:00")
            if prof in prof_not_before:
                if slot_to_idx.get(slot,0) < slot_to_idx.get(prof_not_before[prof],0):
                    return _reject(f"{prof} لا يحضر قبل {prof_not_before[prof]}")
            if prof in prof_not_after:
                if slot_to_idx.get(slot,99) > slot_to_idx.get(prof_not_after[prof],99):
                    return _reject(f"{prof} لا يحضر بعد {prof_not_after[prof]}")
            if prof in prof_phase_split:
                n_first, start_second = prof_phase_split[prof]
                cur = prof_first_phase_count.get(prof, 0)
                if cur < n_first:
                    if day >= start_second: return _reject(f"{prof} في الفترة الأولى — {day} في الثانية")
                else:
                    if day < start_second: return _reject(f"{prof} في الفترة الثانية — {day} في الأولى")
        return True
    
    def place_memo(memo_id, day, slot, room):
        """وضع المذكرة في الخانة"""
        schedule[memo_id] = (day, slot, room)
        occupied[(day, slot, room)] = memo_id
        members = memo_members.get(memo_id, set())
        for prof in members:
            prof_busy[(day, slot, prof)] = memo_id
            prof_day_count[(prof, day)] = prof_day_count.get((prof, day), 0) + 1
            if prof in prof_phase_split:
                prof_first_phase_count[prof] = prof_first_phase_count.get(prof, 0) + 1
    
    # رتّب الأساتذة من الأكثر مذكرات إلى الأقل
    sorted_profs = sorted(prof_memos_map.items(), key=lambda x: len(x[1]), reverse=True)
    
    # تتبع المذكرات المجدولة
    scheduled_memos = set()

    # بناء خريطة التقسيم الزمني للأساتذة
    prof_first_phase_count = {prof: 0 for prof in prof_phase_split}

    # المرحلة 0: تطبيق المواعيد المثبتة أولاً (قيود صارمة)
    for _fmid, (_fday, _fslot, _froom) in fixed_slots.items():
        if _fday in days and _fslot in slots_per_day:
            _fr = _froom if _froom and _froom in rooms else (rooms[0] if rooms else "")
            if can_place(_fmid, _fday, _fslot, _fr):
                place_memo(_fmid, _fday, _fslot, _fr)
                scheduled_memos.add(_fmid)
    
    # المرحلة 1: جدولة مذكرات كل أستاذ
    for prof, prof_memo_list in sorted_profs:
        # المذكرات التي لم تُجدول بعد لهذا الأستاذ
        unscheduled = [m for m in prof_memo_list if m not in scheduled_memos]
        if not unscheduled:
            continue
        
        # خطط توزيع أيامه
        day_plan = plan_prof_days(len(unscheduled), days, max_per_day, max_consecutive)
        
        memo_idx = 0
        for (day_idx, count) in day_plan:
            if day_idx >= len(days):
                break
            day = days[day_idx]
            
            # اختر حصص متتالية قدر الإمكان في هذا اليوم
            placed_today = 0
            slot_idx = 0
            
            while placed_today < count and memo_idx < len(unscheduled):
                if slot_idx >= len(slots_per_day):
                    break
                
                slot = slots_per_day[slot_idx]
                memo = unscheduled[memo_idx]
                
                # ابحث عن قاعة متاحة
                placed = False
                for room in rooms:
                    if can_place(memo, day, slot, room):
                        place_memo(memo, day, slot, room)
                        scheduled_memos.add(memo)
                        placed_today += 1
                        memo_idx += 1
                        placed = True
                        break
                
                if not placed:
                    # جرب التوقيت التالي في نفس اليوم
                    # لكن تحقق من الفراغ — لا فراغ أكثر من توقيت واحد
                    if placed_today > 0:
                        # هناك فراغ — نقبله إذا كان توقيتاً واحداً فقط
                        gap = slot_idx - slots_per_day.index(slots_per_day[slot_idx-1]) if placed_today > 0 else 0
                        if gap > 1:
                            # فراغ كبير — انتقل ليوم جديد
                            break
                
                slot_idx += 1
    
    # المرحلة 2: المذكرات التي لم تُجدول بعد → ابحث عن أي خانة فارغة
    unscheduled_final = [m for m in memo_ids if m not in scheduled_memos]
    
    for memo in unscheduled_final:
        placed = False
        for day in days:
            for slot in slots_per_day:
                for room in rooms:
                    if can_place(memo, day, slot, room):
                        place_memo(memo, day, slot, room)
                        scheduled_memos.add(memo)
                        placed = True
                        break
                if placed: break
            if placed: break
        
        if not placed:
            schedule[memo] = None
    
    # تأكد أن كل المذكرات في الجدول
    for memo in memo_ids:
        if memo not in schedule:
            schedule[memo] = None
    
    return schedule, memo_members, rejection_log

def calc_schedule_quality(schedule, memo_members, days, slots_per_day):
    """حساب جودة الجدول من منظور الأستاذ"""
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    day_to_idx = {d: i for i, d in enumerate(days)}
    
    prof_program = {}
    for mid, slot in schedule.items():
        if not slot: continue
        day, time_slot, room = slot
        for prof in memo_members.get(mid, set()):
            prof_program.setdefault(prof, {}).setdefault(day, []).append(slot_to_idx.get(time_slot, 0))
    
    total_idle = 0
    total_days = 0
    max_gap = 0
    
    for prof, days_dict in prof_program.items():
        total_days += len(days_dict)
        for day, slot_indices in days_dict.items():
            if len(slot_indices) > 1:
                sorted_s = sorted(slot_indices)
                for i in range(len(sorted_s)-1):
                    gap = sorted_s[i+1] - sorted_s[i] - 1
                    total_idle += gap
                    max_gap = max(max_gap, gap)
    
    placed = len([s for s in schedule.values() if s])
    unplaced = len([s for s in schedule.values() if not s])
    total = len(schedule)
    
    placement_rate = placed / total * 100 if total else 0
    quality = round(placement_rate * 0.7 + max(0, 30 - total_idle) * 1.0, 1)
    
    return min(100, quality), placed, unplaced, total_idle, total_days, max_gap

def improve_schedule(schedule, memo_members, days, slots_per_day, rooms, iterations=2000, prof_banned_days=None, prof_allowed_days=None, profs_accept_18=None, fixed_slots=None):
    """
    تحسين الجدول:
    - ابحث عن أستاذ لديه مذكرة معزولة في يوم وحدها
    - حاول نقلها ليوم فيه مذكرات أخرى له
    """
    import random
    
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    
    def get_prof_program(sched):
        prog = {}
        for mid, slot in sched.items():
            if not slot: continue
            day, time_slot, room = slot
            for prof in memo_members.get(mid, set()):
                prog.setdefault(prof, {}).setdefault(day, []).append(mid)
        return prog
    
    def can_place(sched, memo_id, day, slot, room):
        members_new = memo_members.get(memo_id, set())
        # ✅ تحقق من الأيام الممنوعة أولاً
        for p in members_new:
            _banned = prof_banned_days.get(p, set())
            if day in _banned:
                return False
            if prof_allowed_days.get(p) and day not in prof_allowed_days[p]:
                return False
            if slot > "16:00" and p not in profs_accept_18:
                return False
        day_counts = {}
        for other_mid, other_slot in sched.items():
            if other_mid == memo_id or not other_slot: continue
            if other_slot == (day, slot, room): return False
            if other_slot[0] == day and other_slot[1] == slot:
                if members_new & memo_members.get(other_mid, set()):
                    return False
            if other_slot[0] == day:
                for p in memo_members.get(other_mid, set()):
                    day_counts[p] = day_counts.get(p, 0) + 1
        for p in members_new:
            if day_counts.get(p, 0) >= 3:
                return False
        return True
    
    prof_banned_days = prof_banned_days or {}
    prof_allowed_days = prof_allowed_days or {}
    profs_accept_18 = profs_accept_18 or set()
    fixed_slots = fixed_slots or {}
    _fixed_mids = set(str(k) for k in fixed_slots.keys())

    # مذكرات محمية — لا تُلمس أبداً:
    # 1. مواعيد مثبتة
    # 2. مذكرات فيها أستاذ له قيود (أيام ممنوعة أو مسموحة فقط)
    # 3. مذكرات لها قيود زمنية (أقرب/أبعد تاريخ أو أيام بديلة)
    _protected_mids = set(str(k) for k in fixed_slots.keys())
    for mid, members in memo_members.items():
        for prof in members:
            if prof in prof_banned_days or prof in prof_allowed_days:
                _protected_mids.add(str(mid))
                break

    current = dict(schedule)
    _, _, _, cur_idle, cur_days, _ = calc_schedule_quality(current, memo_members, days, slots_per_day)
    cur_score = cur_idle + cur_days * 10
    
    for _ in range(iterations):
        prog = get_prof_program(current)
        
        # ابحث عن أستاذ له يوم فيه مذكرة واحدة فقط
        lonely_memos = []
        for prof, days_dict in prog.items():
            if len(days_dict) <= 1: continue
            for day, memos in days_dict.items():
                if len(memos) == 1:
                    lonely_memos.append((prof, day, memos[0]))
        
        if not lonely_memos:
            # تحقق من أيام منعزلة متعددة — فقط إذا المجموع >= 3
            multi_lonely = []
            for prof, days_dict in prog.items():
                total_p = sum(len(ms) for ms in days_dict.values())
                if total_p < 3: continue
                solo_days = sorted([d for d, ms in days_dict.items() if len(ms)==1])
                if len(solo_days) > 2:  # أكثر من يومين منعزلين
                    for d in solo_days[2:]:  # ابقِ الأول والثاني
                        multi_lonely.append((prof, d, days_dict[d][0]))

            # تحقق من أيام متتالية > 3
            if not multi_lonely:
                from datetime import datetime as _dt2, timedelta as _td2
                for prof, days_dict in prog.items():
                    sorted_d = sorted(days_dict.keys())
                    consec = 1
                    for i in range(1, len(sorted_d)):
                        try:
                            d1 = _dt2.strptime(sorted_d[i-1], "%Y-%m-%d")
                            d2 = _dt2.strptime(sorted_d[i], "%Y-%m-%d")
                            if (d2-d1).days <= 3:
                                consec += 1
                                if consec > 3:
                                    # انقل آخر يوم في السلسلة
                                    target_day = sorted_d[i]
                                    target_memo = days_dict[target_day][0]
                                    multi_lonely.append((prof, target_day, target_memo))
                            else:
                                consec = 1
                        except: pass

            if not multi_lonely:
                break
            lonely_memos = multi_lonely
        
        random.shuffle(lonely_memos)
        # تخطى المذكرات المحمية
        _movable = [(p, d, m) for p, d, m in lonely_memos if str(m) not in _protected_mids]
        if not _movable: break
        prof, lonely_day, lonely_memo = _movable[0]
        old_slot = current.get(lonely_memo)
        if not old_slot: continue
        
        # حاول نقلها ليوم فيه مذكرات أخرى لنفس الأستاذ
        best_day = None
        best_slot_val = None
        
        for target_day, target_memos in prog.get(prof, {}).items():
            if target_day == lonely_day: continue
            if len(target_memos) >= 3: continue  # اليوم ممتلئ
            
            # جرب وضعها في توقيت متتالي مع مذكرات الأستاذ في هذا اليوم
            for slot in slots_per_day:
                for room in rooms:
                    if can_place(current, lonely_memo, target_day, slot, room):
                        current[lonely_memo] = (target_day, slot, room)
                        _, _, _, new_idle, new_days, _ = calc_schedule_quality(current, memo_members, days, slots_per_day)
                        new_score = new_idle + new_days * 10
                        if new_score < cur_score:
                            cur_score = new_score
                            best_day = target_day
                            best_slot_val = (target_day, slot, room)
                        current[lonely_memo] = old_slot
                        break
                if best_slot_val: break
            if best_day: break
        
        if best_slot_val:
            current[lonely_memo] = best_slot_val
            _, _, _, cur_idle, cur_days, _ = calc_schedule_quality(current, memo_members, days, slots_per_day)
            cur_score = cur_idle + cur_days * 10
    
    return current


# ================================================================
# 🧠 خوارزمية 1: كتل الأساتذة (Professor Blocks)
# ================================================================

# ================================================================
# 🔒 دالة مشتركة للتحقق من القيود — تُستخدم في كل الخوارزميات
# ================================================================
def make_can_place(occupied, prof_busy, prof_day_count, memo_members,
                   fixed_slots, memo_date_limits, prof_banned_days,
                   prof_not_before, prof_not_after, prof_allowed_days,
                   profs_accept_18, memo_alt_days, prof_phase_split,
                   prof_first_phase_count, slots_per_day, rejection_log=None):
    """
    تُنشئ دالة can_place مخصصة بكل القيود
    تُرجع (can_place_fn, constraint_report)
    """
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    LATE_SLOT = "18:00"
    constraint_report = []  # سجل تطبيق القيود


    def _reject(memo_id, reason):
        if rejection_log is not None:
            rejection_log.setdefault(str(memo_id), set()).add(reason)
        return False

    def can_place(memo_id, day, slot, room, log=True):
        mid = str(memo_id)

        # 1. القاعة مشغولة
        if (day, slot, room) in occupied:
            return _reject(mid, f"القاعة {room} مشغولة {day} {slot}") if log else False

        # 2. لا بعد 16:00 للجميع إلا من يقبل 18:00
        if slot > "16:00":
            _members = memo_members.get(mid, set())
            _non_acc = _members - profs_accept_18
            if _non_acc:
                return _reject(mid, f"توقيت {slot} بعد 16:00 — لا يقبله: {', '.join(sorted(_non_acc)[:2])}") if log else False

        # 3. قيود التوقيت اليومي (من/إلى في شيت الأيام) — يُمرر من خارج
        # (يُعالج في day_time_limits خارج هذه الدالة)

        # 4. حدود تاريخ المذكرة
        if mid in memo_date_limits:
            earliest, latest = memo_date_limits[mid]
            if earliest and day < earliest:
                return _reject(mid, f"لا تُبرمج قبل {earliest}") if log else False
            if latest and day > latest:
                return _reject(mid, f"لا تُبرمج بعد {latest}") if log else False

        # 5. الأيام البديلة
        if mid in memo_alt_days and memo_alt_days[mid]:
            if day not in memo_alt_days[mid]:
                return _reject(mid, f"يوم {day} ليس من الأيام البديلة") if log else False

        members = memo_members.get(mid, set())
        for prof in members:
            # 6. تعارض في نفس التوقيت
            if (day, slot, prof) in prof_busy:
                return _reject(mid, f"{prof} مشغول {day} {slot}") if log else False

            # 7. الحد الأقصى 3 مناقشات/يوم
            if prof_day_count.get((prof, day), 0) >= 3:
                return _reject(mid, f"{prof} بلغ الحد 3 في {day}") if log else False

            # 8. أيام ممنوعة
            if day in prof_banned_days.get(prof, set()):
                if rejection_log is not None:
                    rejection_log.setdefault(str(mid), set()).add(f"يوم {day} ممنوع على {prof}")
                return False  # رفض مطلق بغض النظر عن log

            # 9. أيام مسموحة فقط
            if prof in prof_allowed_days and prof_allowed_days[prof]:
                if day not in prof_allowed_days[prof]:
                    if rejection_log is not None:
                        rejection_log.setdefault(str(mid), set()).add(f"يوم {day} غير مسموح لـ{prof}")
                    return False  # رفض مطلق

            # 10. لا قبل توقيت
            if prof in prof_not_before:
                if slot_to_idx.get(slot, 0) < slot_to_idx.get(prof_not_before[prof], 0):
                    return _reject(mid, f"{prof} لا يحضر قبل {prof_not_before[prof]}") if log else False

            # 11. لا بعد توقيت
            if prof in prof_not_after:
                if slot_to_idx.get(slot, 99) > slot_to_idx.get(prof_not_after[prof], 99):
                    return _reject(mid, f"{prof} لا يحضر بعد {prof_not_after[prof]}") if log else False

            # 12. تقسيم الفترتين
            if prof in prof_phase_split:
                n_first, start_second = prof_phase_split[prof]
                cur = prof_first_phase_count.get(prof, 0)
                if cur < n_first:
                    if day >= start_second:
                        return _reject(mid, f"{prof} لا يزال في الفترة الأولى") if log else False
                else:
                    if day < start_second:
                        return _reject(mid, f"{prof} انتقل للفترة الثانية") if log else False

        return True

    return can_place, constraint_report


def make_place_fn(schedule, occupied, prof_busy, prof_day_count, memo_members, prof_first_phase_count, prof_phase_split):
    """دالة وضع المذكرة في الجدول"""
    def place(memo_id, day, slot, room):
        mid = str(memo_id)
        schedule[mid] = (day, slot, room)
        occupied[(day, slot, room)] = mid
        for prof in memo_members.get(mid, set()):
            prof_busy[(day, slot, prof)] = mid
            prof_day_count[(prof, day)] = prof_day_count.get((prof, day), 0) + 1
            if prof in prof_phase_split:
                prof_first_phase_count[prof] = prof_first_phase_count.get(prof, 0) + 1
    return place


def apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled, rejection_log=None):
    """تطبيق المواعيد المثبتة — مشترك بين كل الخوارزميات"""
    applied = []
    failed = []
    for mid, slot_val in fixed_slots.items():
        fd, fs, fr = slot_val
        fd, fs = str(fd).strip(), str(fs).strip()
        if fd not in days:

            failed.append(f"المذكرة {mid}: يوم التثبيت {fd} غير موجود في الأيام")
            continue
        if fs not in slots_per_day:

            failed.append(f"المذكرة {mid}: توقيت التثبيت {fs} غير موجود في التوقيتات")
            continue
        placed = False
        for r in ([fr] if fr and fr in rooms else rooms):
            if can_place(mid, fd, fs, r, log=False):
                place(mid, fd, fs, r)
                scheduled.add(mid)
                applied.append(f"✅ المذكرة {mid} → {fd} {fs} {r}")
                placed = True; break
        if not placed:
            # debug: جرب مع log=True لمعرفة السبب

            failed.append(f"⚠️ المذكرة {mid}: لا يمكن تطبيق التثبيت {fd} {fs} — تعارض")
    return applied, failed


def algo_blocks(df_memos, days, slots_per_day, rooms, constraints):
    """
    كتل الأساتذة — توزيع عادل مع تجميع مناقشات الأستاذ
    المبدأ: نبني الجدول عمودياً (جولة لكل يوم) لضمان التوازن
    """
    import random
    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after, \
    prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split, \
    memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}

    schedule, occupied, prof_busy, prof_day_count, prof_first_phase_count = {}, {}, {}, {}, {}
    rejection_log = {}
    scheduled = set()

    # الطاقة اليومية المستهدفة — توزيع عادل
    n_memos = len(memo_ids)
    n_days = len(days)
    n_slots = len(slots_per_day)
    n_rooms = len(rooms)
    daily_capacity = n_slots * n_rooms  # الطاقة القصوى
    daily_target = max(1, round(n_memos / n_days))  # الهدف اليومي
    daily_max = min(daily_capacity, int(daily_target * 1.4))  # سقف مرن

    # عداد المذكرات في كل يوم
    day_count = {d: 0 for d in days}

    can_place, _ = make_can_place(
        occupied, prof_busy, prof_day_count, memo_members,
        fixed_slots, memo_date_limits, prof_banned_days,
        prof_not_before, prof_not_after, prof_allowed_days,
        profs_accept_18, memo_alt_days, prof_phase_split,
        prof_first_phase_count, slots_per_day, rejection_log)
    place = make_place_fn(schedule, occupied, prof_busy, prof_day_count,
                          memo_members, prof_first_phase_count, prof_phase_split)

    # Phase 0: Fixed slots
    applied, failed = apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled)
    for mid in scheduled:
        sv = schedule.get(mid)
        if sv: day_count[sv[0]] = day_count.get(sv[0], 0) + 1

    # Phase 1: رتّب الأساتذة من الأكثر مناقشات
    sorted_profs = sorted(prof_memos_map.items(), key=lambda x: len(x[1]), reverse=True)

    # لكل أستاذ — خصص له أيام متجاورة بناءً على عدد مناقشاته
    prof_preferred_days = {}
    for prof, pmemos in sorted_profs:
        n = len([m for m in pmemos if m not in scheduled])
        if n == 0: continue
        banned = prof_banned_days.get(prof, set())
        allowed = prof_allowed_days.get(prof, set())
        # أيام مرتبة من الأقل اكتظاظاً
        cand = [d for d in days if d not in banned and (not allowed or d in allowed)]
        cand_sorted = sorted(cand, key=lambda d: day_count.get(d, 0))
        # اختر أيام متجاورة قدر الإمكان
        n_days_needed = max(1, (n + 2) // 3)
        # ابحث عن كتلة متجاورة
        best_block = []
        for i in range(len(cand_sorted)):
            # جرب بناء كتلة من هذا اليوم
            block = [cand_sorted[i]]
            for j in range(i+1, len(cand_sorted)):
                if len(block) >= n_days_needed: break
                # هل اليوم مجاور لأحد أيام الكتلة؟
                try:
                    from datetime import datetime as _dt
                    d_new = _dt.strptime(cand_sorted[j], "%Y-%m-%d")
                    is_adj = any(
                        abs((_dt.strptime(d, "%Y-%m-%d") - d_new).days) <= 2
                        for d in block
                    )
                    if is_adj or len(block) == 0:
                        block.append(cand_sorted[j])
                except: block.append(cand_sorted[j])
            if len(block) >= min(n_days_needed, len(cand_sorted)):
                best_block = block[:n_days_needed]; break
        prof_preferred_days[prof] = best_block or cand_sorted[:n_days_needed]

    # Phase 2: توزيع عمودي — جولة لكل يوم
    remaining = [m for m in memo_ids if m not in scheduled]

    # رتب المذكرات: الأساتذة الأثقل أولاً
    def memo_priority(mid):
        members = memo_members.get(mid, set())
        max_load = max((len(prof_memos_map.get(p, [])) for p in members), default=0)
        return -max_load

    remaining.sort(key=memo_priority)

    # جولات توزيع — كل جولة تضع مذكرة في كل يوم
    max_rounds = daily_max + 5
    for round_num in range(max_rounds):
        if not [m for m in remaining if m not in scheduled]:
            break
        # رتب الأيام من الأقل اكتظاظاً
        days_ordered = sorted(days, key=lambda d: day_count.get(d, 0))

        for day in days_ordered:
            if day_count.get(day, 0) >= daily_max:
                continue
            # ابحث عن مذكرة مناسبة لهذا اليوم
            best_memo = None
            for mid in remaining:
                if mid in scheduled: continue
                # أولوية للمذكرات التي أعضاؤها لديهم مناقشات في هذا اليوم
                members = memo_members.get(mid, set())
                has_prof_here = any(
                    day in prof_preferred_days.get(p, []) for p in members
                )
                day_not_banned = all(
                    day not in prof_banned_days.get(p, set()) and
                    (not prof_allowed_days.get(p) or day in prof_allowed_days.get(p, set()))
                    for p in members
                )
                if day_not_banned:
                    if has_prof_here:
                        best_memo = mid; break
            if best_memo is None:
                for mid in remaining:
                    if mid in scheduled: continue
                    members = memo_members.get(mid, set())
                    ok = all(
                        day not in prof_banned_days.get(p, set()) and
                        (not prof_allowed_days.get(p) or day in prof_allowed_days.get(p, set()))
                        for p in members
                    )
                    if ok: best_memo = mid; break

            if best_memo is None: continue

            # ضع المذكرة في أفضل توقيت متاح
            for slot in slots_per_day:
                placed = False
                for room in rooms:
                    if can_place(best_memo, day, slot, room, log=True):
                        place(best_memo, day, slot, room)
                        scheduled.add(best_memo)
                        day_count[day] = day_count.get(day, 0) + 1
                        placed = True; break
                if placed: break

    # Phase 3: المتبقية
    for memo in memo_ids:
        if memo in scheduled: continue
        placed = False
        # رتب الأيام من الأقل
        for day in sorted(days, key=lambda d: day_count.get(d, 0)):
            for slot in slots_per_day:
                for room in rooms:
                    if can_place(memo, day, slot, room, log=True):
                        place(memo, day, slot, room)
                        scheduled.add(memo)
                        day_count[day] = day_count.get(day, 0) + 1
                        placed = True; break
                if placed: break
            if placed: break
        if not placed: schedule[memo] = None

    for memo in memo_ids:
        if memo not in schedule: schedule[memo] = None

    return schedule, memo_members, rejection_log


def algo_day_first(df_memos, days, slots_per_day, rooms, constraints):
    """الجدول أولاً — يوم بيوم مع أولوية للمذكرات ذات الأعضاء المتاحين"""
    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,     prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split,     memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()

    schedule, occupied, prof_busy, prof_day_count, prof_first_phase_count = {}, {}, {}, {}, {}
    rejection_log = {}
    scheduled = set()

    can_place, _ = make_can_place(occupied, prof_busy, prof_day_count, memo_members,
        fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,
        prof_allowed_days, profs_accept_18, memo_alt_days, prof_phase_split,
        prof_first_phase_count, slots_per_day, rejection_log)
    place = make_place_fn(schedule, occupied, prof_busy, prof_day_count, memo_members,
        prof_first_phase_count, prof_phase_split)

    applied, failed = apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled)
    if failed:
        import streamlit as _st
        for f in failed: _st.warning(f)

    def score_memo_for_day(memo_id, day):
        members = memo_members.get(str(memo_id), set())
        if not members: return 0
        free = sum(1 for p in members
            if day not in prof_banned_days.get(p, set())
            and (not prof_allowed_days.get(p) or day in prof_allowed_days.get(p, set()))
            and prof_day_count.get((p, day), 0) < 3)
        return free / len(members)

    remaining = [m for m in memo_ids if m not in scheduled]
    for day in days:
        day_cands = sorted(remaining, key=lambda m: -score_memo_for_day(m, day))
        for memo in day_cands:
            if memo in scheduled: continue
            for slot in slots_per_day:
                placed = False
                for room in rooms:
                    if can_place(memo, day, slot, room):
                        place(memo, day, slot, room)
                        scheduled.add(memo); placed = True; break
                if placed: break
        remaining = [m for m in memo_ids if m not in scheduled]

    for memo in remaining:
        placed = False
        for day in days:
            for slot in slots_per_day:
                for room in rooms:
                    if can_place(memo, day, slot, room, log=True):
                        place(memo, day, slot, room)
                        scheduled.add(memo); placed = True; break
                if placed: break
            if placed: break
        if not placed: schedule[memo] = None

    for memo in memo_ids:
        if memo not in schedule: schedule[memo] = None

    return schedule, memo_members, rejection_log


def algo_greedy(df_memos, days, slots_per_day, rooms, constraints):
    """الأثقل أولاً (Greedy) — يختار أفضل يوم لكل مذكرة"""
    import random
    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,     prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split,     memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()

    schedule, occupied, prof_busy, prof_day_count, prof_first_phase_count = {}, {}, {}, {}, {}
    prof_days_used = {}
    rejection_log = {}
    scheduled = set()

    can_place, _ = make_can_place(occupied, prof_busy, prof_day_count, memo_members,
        fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,
        prof_allowed_days, profs_accept_18, memo_alt_days, prof_phase_split,
        prof_first_phase_count, slots_per_day, rejection_log)

    def place(memo_id, day, slot, room):
        mid = str(memo_id)
        schedule[mid] = (day, slot, room)
        occupied[(day, slot, room)] = mid
        for prof in memo_members.get(mid, set()):
            prof_busy[(day, slot, prof)] = mid
            prof_day_count[(prof, day)] = prof_day_count.get((prof, day), 0) + 1
            prof_days_used.setdefault(prof, set()).add(day)
            if prof in prof_phase_split:
                prof_first_phase_count[prof] = prof_first_phase_count.get(prof, 0) + 1

    applied, failed = apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled)
    if failed:
        import streamlit as _st
        for f in failed: _st.warning(f)

    def best_day_for_memo(memo_id):
        mid = str(memo_id)
        if mid in fixed_slots and mid in scheduled:
            return fixed_slots[mid][0]
        members = memo_members.get(mid, set())
        best_day, best_score = None, -1
        for day in days:
            if mid in memo_date_limits:
                e, l = memo_date_limits[mid]
                if e and day < e: continue
                if l and day > l: continue
            if mid in memo_alt_days and memo_alt_days[mid]:
                if day not in memo_alt_days[mid]: continue
            available, already_here = 0, 0
            for p in members:
                if day in prof_banned_days.get(p, set()): continue
                if prof_allowed_days.get(p) and day not in prof_allowed_days[p]: continue
                if prof_day_count.get((p, day), 0) >= 3: continue
                available += 1
                if day in prof_days_used.get(p, set()): already_here += 1
            score = available * 2 + already_here * 3
            if score > best_score:
                best_score = score; best_day = day
        return best_day

    sorted_profs = sorted(prof_memos_map.items(), key=lambda x: len(x[1]), reverse=True)
    for prof, pmemos in sorted_profs:
        unsch = [m for m in pmemos if m not in scheduled]
        random.shuffle(unsch)
        for memo in unsch:
            best_day = best_day_for_memo(memo)
            if not best_day: continue
            for slot in slots_per_day:
                placed = False
                for room in rooms:
                    if can_place(memo, best_day, slot, room):
                        place(memo, best_day, slot, room)
                        scheduled.add(memo); placed = True; break
                if placed: break

    for memo in memo_ids:
        if memo in scheduled: continue
        placed = False
        for day in days:
            for slot in slots_per_day:
                for room in rooms:
                    if can_place(memo, day, slot, room, log=True):
                        place(memo, day, slot, room)
                        scheduled.add(memo); placed = True; break
                if placed: break
            if placed: break
        if not placed: schedule[memo] = None

    for memo in memo_ids:
        if memo not in schedule: schedule[memo] = None

    return schedule, memo_members, rejection_log



# ================================================================
# 🏆 Smart Fair Scheduling Optimizer
# ================================================================
def algo_smart_fair(df_memos, days, slots_per_day, rooms, constraints):
    """
    Smart Fair Scheduling Optimizer
    المرحلة 1: توليد schedule صالح يحترم كل القيود
    المرحلة 2: optimization بـ scoring function
    """
    import random
    from itertools import combinations

    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after, \
    prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split, \
    memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    MAX_PER_DAY = 3

    # ── Score Function ──
    def compute_score(schedule, memo_members_map):
        prof_program = {}  # prof -> {day -> [slot_idx]}
        room_usage = {}    # (day, slot) -> count

        for mid, sv in schedule.items():
            if not sv: continue
            day, slot, room = sv
            slot_idx = slot_to_idx.get(slot, 0)
            for prof in memo_members_map.get(mid, set()):
                prof_program.setdefault(prof, {}).setdefault(day, []).append(slot_idx)
            room_usage[(day, slot)] = room_usage.get((day, slot), 0) + 1

        total_score = 0.0
        prof_totals = {p: sum(len(s) for s in d.items()) for p, d in prof_program.items()
                       for _ in [None]}

        # 1. Fairness score — انحراف معياري منخفض = أفضل
        if prof_program:
            counts = [sum(len(slots) for slots in days_dict.values())
                      for days_dict in prof_program.values()]
            mean_c = sum(counts) / len(counts) if counts else 0
            std_c = (sum((c - mean_c)**2 for c in counts) / len(counts))**0.5 if counts else 0
            fairness_score = 100 / (1 + std_c)
            total_score += fairness_score * 3  # وزن عالٍ

        for prof, days_dict in prof_program.items():
            sorted_days = sorted(days_dict.keys())

            # 2. Overload penalty — تجاوز 3/يوم
            for day, slots in days_dict.items():
                if len(slots) > MAX_PER_DAY:
                    total_score -= (len(slots) - MAX_PER_DAY) * 50

            # 3. Consecutive days penalty
            consec = 1; max_consec = 1
            for i in range(1, len(sorted_days)):
                try:
                    from datetime import datetime
                    d1 = datetime.strptime(sorted_days[i-1], "%Y-%m-%d")
                    d2 = datetime.strptime(sorted_days[i], "%Y-%m-%d")
                    diff = (d2 - d1).days
                    if diff <= 3: consec += 1; max_consec = max(max_consec, consec)
                    else: consec = 1
                except: pass
            if max_consec > 3:
                total_score -= (max_consec - 3) * 20

            # 4. Gaps penalty — فجوات كبيرة في اليوم
            for day, slots in days_dict.items():
                if len(slots) > 1:
                    sorted_slots = sorted(slots)
                    max_gap = max(sorted_slots[i+1] - sorted_slots[i]
                                  for i in range(len(sorted_slots)-1))
                    if max_gap > 1:
                        total_score -= max_gap * 10

            # 5. Professor comfort — مناقشات مجمعة في أقل أيام
            n_days = len(days_dict)
            total_sessions = sum(len(s) for s in days_dict.values())
            if n_days > 0:
                sessions_per_day = total_sessions / n_days
                comfort_score = sessions_per_day * 15
                total_score += comfort_score

            # 6. Lonely days penalty — أيام منعزلة
            if total_sessions >= 3:
                lonely = sum(1 for d, s in days_dict.items() if len(s) == 1)
                if lonely > 1:
                    total_score -= (lonely - 1) * 25

        # 7. Room efficiency
        placed = sum(1 for sv in schedule.values() if sv)
        capacity = len(days) * len(slots_per_day) * len(rooms)
        if capacity > 0:
            efficiency = placed / capacity
            room_score = efficiency * 50
            total_score += room_score

        return total_score

    # ── المرحلة 1: توليد schedule صالح ──
    schedule, occupied, prof_busy, prof_day_count, prof_first_phase_count = {}, {}, {}, {}, {}
    rejection_log = {}
    scheduled = set()

    can_place, _ = make_can_place(occupied, prof_busy, prof_day_count, memo_members,
        fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,
        prof_allowed_days, profs_accept_18, memo_alt_days, prof_phase_split,
        prof_first_phase_count, slots_per_day, rejection_log)
    place = make_place_fn(schedule, occupied, prof_busy, prof_day_count, memo_members,
        prof_first_phase_count, prof_phase_split)

    # Fixed slots
    applied, failed = apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled)
    if failed:
        import streamlit as _st
        for f in failed: _st.warning(f)

    # ── نرتب المذكرات بذكاء: الأكثر قيوداً أولاً ──
    def memo_priority(mid):
        members = memo_members.get(mid, set())
        constraints_count = 0
        for p in members:
            constraints_count += len(prof_banned_days.get(p, set()))
            constraints_count += (1 if p in prof_allowed_days else 0) * 5
            constraints_count += (1 if p in prof_not_before else 0) * 2
            constraints_count += (1 if p in prof_not_after else 0) * 2
        if mid in memo_alt_days and memo_alt_days[mid]: constraints_count += 10
        if mid in memo_date_limits: constraints_count += 5
        return -constraints_count  # الأكثر قيوداً أولاً

    remaining = sorted([m for m in memo_ids if m not in scheduled], key=memo_priority)

    # ── اختيار أفضل خانة لكل مذكرة (greedy مع scoring) ──
    for memo in remaining:
        best_slot = None
        best_slot_score = -float('inf')

        members = memo_members.get(memo, set())

        for day in days:
            for slot in slots_per_day:
                for room in rooms:
                    if not can_place(memo, day, slot, room, log=False):
                        continue
                    # احسب تأثير وضع هذه المذكرة هنا
                    slot_score = 0

                    # مكافأة إذا الأستاذ موجود بالفعل في هذا اليوم
                    for p in members:
                        if prof_day_count.get((p, day), 0) > 0:
                            slot_score += 30  # يفضل التجميع
                        # عقوبة إذا الأستاذ سيبلغ 3 في هذا اليوم
                        if prof_day_count.get((p, day), 0) == 2:
                            slot_score -= 5  # اليوم سيمتلئ

                    # مكافأة إذا التوقيتات متتالية
                    for p in members:
                        existing_slots = [slot_to_idx.get(s, 0)
                                          for (d, s, _), mid2 in [(sv, m) for m, sv in schedule.items() if sv]
                                          if d == day]
                        if existing_slots:
                            min_gap = min(abs(slot_to_idx.get(slot, 0) - es) for es in existing_slots)
                            if min_gap == 1: slot_score += 20  # متتالية
                            elif min_gap > 2: slot_score -= 15  # فجوة

                    # عقوبة على أيام منعزلة
                    for p in members:
                        days_with_sessions = [d for (d, s, r), mid2 in [(sv, m) for m, sv in schedule.items() if sv]
                                              for pp in [memo_members.get(mid2, set())] if p in pp]
                        if days_with_sessions and day not in days_with_sessions:
                            slot_score -= 10  # يوم جديد = قد يكون منعزلاً

                    if slot_score > best_slot_score:
                        best_slot_score = slot_score
                        best_slot = (day, slot, room)

        if best_slot:
            d, s, r = best_slot
            place(memo, d, s, r)
            scheduled.add(memo)
        else:
            schedule[memo] = None
            rejection_log.setdefault(memo, set()).add("لا توجد خانة مناسبة مع جميع القيود")

    for memo in memo_ids:
        if memo not in schedule: schedule[memo] = None

    # ── المرحلة 2: Tabu Search optimization ──
    current_score = compute_score(schedule, memo_members)
    best_schedule = dict(schedule)
    best_score = current_score

    placed_memos = [m for m, sv in schedule.items() if sv]

    for iteration in range(200):
        if len(placed_memos) < 2: break
        m1, m2 = random.sample(placed_memos, 2)
        sv1, sv2 = schedule[m1], schedule[m2]
        if not sv1 or not sv2: continue

        # جرب المبادلة
        schedule[m1] = sv2; schedule[m2] = sv1

        # تحقق صحة المبادلة (لا تعارض)
        # بناء occupied مؤقت
        test_occupied = {}
        valid_swap = True
        for m, sv in schedule.items():
            if not sv: continue
            if sv in test_occupied:
                valid_swap = False; break
            test_occupied[sv] = m

        if valid_swap:
            new_score = compute_score(schedule, memo_members)
            if new_score > best_score:
                best_score = new_score
                best_schedule = dict(schedule)
            elif new_score < current_score - 10:  # تراجع سيء
                schedule[m1] = sv1; schedule[m2] = sv2  # تراجع
            else:
                current_score = new_score
        else:
            schedule[m1] = sv1; schedule[m2] = sv2  # تراجع

    return best_schedule, memo_members, rejection_log



# ================================================================
# 🏆 Smart Fair Scheduling Optimizer (SFSO)
# ================================================================



def _validate_hard_constraints(sched, memo_members, max_per_day=3):
    """التحقق الصارم من القيود الحرجة"""
    violations = []
    prof_slot = {}
    room_slot = {}
    memo_count = {}
    prof_day = {}
    for mid, sv in sched.items():
        if not sv: continue
        day, slot, room = sv
        memo_count[mid] = memo_count.get(mid, 0) + 1
        if memo_count[mid] > 1:
            violations.append(f"🔴 المذكرة {mid} مبرمجة أكثر من مرة")
        key_r = (day, slot, room)
        if key_r in room_slot:
            violations.append(f"🔴 تعارض قاعة: {room} في {day} {slot}")
        room_slot[key_r] = mid
        for prof in memo_members.get(mid, set()):
            key_p = (day, slot, prof)
            if key_p in prof_slot:
                violations.append(f"🔴 تعارض: {prof} في {day} {slot}")
            prof_slot[key_p] = mid
            prof_day[(prof, day)] = prof_day.get((prof, day), 0) + 1
            if prof_day[(prof, day)] > max_per_day:
                violations.append(f"🔴 {prof} تجاوز {max_per_day} في {day}")

    # ── قيد صارم: لا أكثر من يوم منعزل واحد (إذا المجموع >= 3) ──
    prof_day_slots = {}
    for mid, sv in sched.items():
        if not sv: continue
        day = sv[0]
        for prof in memo_members.get(mid, set()):
            prof_day_slots.setdefault(prof, {})
            prof_day_slots[prof][day] = prof_day_slots[prof].get(day, 0) + 1

    for prof, day_counts in prof_day_slots.items():
        total_p = sum(day_counts.values())
        if total_p < 3: continue  # استثناء: أقل من 3 مناقشات
        lonely_days = [d for d, c in day_counts.items() if c == 1]
        if len(lonely_days) > 2:
            violations.append(f"🔴 {prof}: {len(lonely_days)} أيام منعزلة — الحد المسموح يومان")

    return list(set(violations))


def _compute_soft_score(sched, memo_members, slot_to_idx, days, slots_per_day, rooms):
    """حساب النقاط للقيود الناعمة"""
    import math
    from datetime import datetime
    prof_slots_by_day = {}
    prof_days = {}
    total_by_prof = {}
    for mid, sv in sched.items():
        if not sv: continue
        day, slot, room = sv
        for prof in memo_members.get(mid, set()):
            prof_slots_by_day.setdefault(prof, {}).setdefault(day, []).append(slot_to_idx.get(slot, 0))
            prof_days.setdefault(prof, set()).add(day)
            total_by_prof[prof] = total_by_prof.get(prof, 0) + 1
    score = 0.0
    if total_by_prof:
        vals = list(total_by_prof.values())
        mean = sum(vals) / len(vals)
        std = math.sqrt(sum((v-mean)**2 for v in vals)/len(vals)) if len(vals)>1 else 0
        score += max(0, 100 - std * 8)
    for prof, days_dict in prof_slots_by_day.items():
        total_p = total_by_prof.get(prof, 0)
        for day, slot_indices in days_dict.items():
            if len(slot_indices) > 1:
                si = sorted(slot_indices)
                gaps = sum(si[i+1]-si[i]-1 for i in range(len(si)-1))
                score -= gaps * 5
        sorted_d = sorted(prof_days.get(prof, set()))
        consec = 1; max_c = 1
        for i in range(1, len(sorted_d)):
            try:
                d1 = datetime.strptime(sorted_d[i-1], "%Y-%m-%d")
                d2 = datetime.strptime(sorted_d[i], "%Y-%m-%d")
                if (d2-d1).days <= 3: consec += 1; max_c = max(max_c, consec)
                else: consec = 1
            except: pass
        if max_c > 3: score -= (max_c-3) * 15
        if total_p >= 3:
            lonely = sum(1 for d, si in days_dict.items() if len(si)==1)
            if lonely > 1: score -= (lonely-1) * 10
        score += max(0, 20 - len(prof_days.get(prof, set())) * 2)
    used = sum(1 for sv in sched.values() if sv)
    cap = len(days) * len(slots_per_day) * len(rooms)
    if cap > 0: score += (used/cap) * 30
    return score


# ================================================================
# 🧬 GA + Tabu Search Optimizer
# ================================================================
def ga_tabu_scheduler(df_memos, days, slots_per_day, rooms, constraints, streamlit_progress=None):
    """
    Genetic Algorithm + Tabu Search
    المرحلة 1: GA — توليد وتطوير أجيال من الجداول
    المرحلة 2: Tabu Search — تحسين أفضل جدول
    """
    import random, math, time
    from datetime import datetime

    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after, \
    prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split, \
    memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    prof_memos_map, memo_members = build_prof_memo_map(df_memos)
    memo_ids = df_memos["رقم المذكرة"].astype(str).tolist()
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}

    # ── دالة توليد جدول صالح (Individual) ──
    def generate_individual(seed=None):
        if seed is not None: random.seed(seed)
        schedule, occupied, prof_busy, prof_day_count, prof_first_phase_count = {}, {}, {}, {}, {}
        rej = {}
        scheduled = set()

        can_place, _ = make_can_place(
            occupied, prof_busy, prof_day_count, memo_members,
            fixed_slots, memo_date_limits, prof_banned_days,
            prof_not_before, prof_not_after, prof_allowed_days,
            profs_accept_18, memo_alt_days, prof_phase_split,
            prof_first_phase_count, slots_per_day, rej
        )
        place = make_place_fn(schedule, occupied, prof_busy, prof_day_count,
                              memo_members, prof_first_phase_count, prof_phase_split)

        apply_fixed_slots(fixed_slots, days, slots_per_day, rooms, can_place, place, scheduled)

        # ترتيب عشوائي للمذكرات
        order = memo_ids[:]
        random.shuffle(order)

        # الفترات الزمنية — نبدأ بالأيام المبكرة
        _cutoff = "2026-06-07"
        _early = [d for d in days if d < _cutoff]
        _late  = [d for d in days if d >= _cutoff]
        # 85% احتمال اختيار يوم مبكر
        if _use_intensive and _early:
            _weighted_days = _early * 3 + _late
        else:
            _weighted_days = days[:]  # كل الأيام سواسية

        for memo in order:
            if memo in scheduled: continue
            # ترتيب يفضل الأيام المبكرة
            if random.random() < 0.85 and _early:
                d_order = _early[:] + _late[:]
            else:
                d_order = days[:]
            random.shuffle(d_order[:len(_early)])  # خلط الأيام المبكرة فيما بينها
            placed = False
            for day in d_order:
                if placed: break
                s_order = slots_per_day[:]
                random.shuffle(s_order)
                for slot in s_order:
                    if placed: break
                    r_order = rooms[:]
                    random.shuffle(r_order)
                    for room in r_order:
                        if can_place(memo, day, slot, room, log=False):
                            place(memo, day, slot, room)
                            scheduled.add(memo)
                            placed = True; break
            if not placed:
                schedule[memo] = None

        for memo in memo_ids:
            if memo not in schedule: schedule[memo] = None

        # تحقق من الأيام المنعزلة
        _lv = len([v for v in _validate_hard_constraints(schedule, memo_members) if "أيام منعزلة" in v and "🔴" in v])
        schedule["__lonely_violations__"] = _lv
        return schedule

    # الفترات الزمنية
    # تاريخ النهاية والنسبة من session_state إذا حُددا من الواجهة
    import streamlit as _st_c
    _use_intensive = _st_c.session_state.get("j_use_intensive", True)
    _cutoff_obj = _st_c.session_state.get("j_cutoff_date", None)
    _target_r   = _st_c.session_state.get("j_target_ratio", 70) / 100
    if _use_intensive:
        CUTOFF_DATE = _cutoff_obj.strftime("%Y-%m-%d") if _cutoff_obj else "2026-06-08"
        early_days = [d for d in days if d < CUTOFF_DATE]
        late_days  = [d for d in days if d >= CUTOFF_DATE]
    else:
        CUTOFF_DATE = None
        early_days = days[:]
        late_days  = []

    # ── دالة الـ Fitness ──
    def fitness(schedule):
        violations = _validate_hard_constraints(schedule, memo_members)
        critical = len([v for v in violations if v.startswith("🔴")])
        if critical > 0:
            return -10000 * critical  # عقوبة شديدة

        placed = sum(1 for sv in schedule.values() if sv)
        total = len(schedule)
        placement_score = (placed / total) * 1000 if total > 0 else 0

        # Soft scores
        soft = _compute_soft_score(schedule, memo_members, slot_to_idx, days, slots_per_day, rooms)

        # ── مكافأة الجدولة في الفترة المبكرة (قبل 7 جوان) ──
        early_placed = sum(1 for sv in schedule.values() if sv and sv[0] < CUTOFF_DATE)
        early_ratio = early_placed / max(placed, 1)
        # نكافئ بقوة إذا 80%+ في الفترة المبكرة
        _tr = _target_r if "_target_r" in dir() else 0.70
        if not _use_intensive or not early_days:
            early_bonus = 0  # كل الأيام سواسية
        elif early_ratio >= _tr: early_bonus = 400
        elif early_ratio >= _tr - 0.05: early_bonus = 200
        elif early_ratio >= _tr - 0.15: early_bonus = 50
        else: early_bonus = -100
        soft += early_bonus

        # ── عقوبة الأيام المتتالية > 3 ──
        from datetime import datetime as _dt
        prof_days_sched = {}
        for mid, sv in schedule.items():
            if not sv: continue
            day = sv[0]
            for prof in memo_members.get(mid, set()):
                prof_days_sched.setdefault(prof, set()).add(day)
        for prof, pdays in prof_days_sched.items():
            sorted_d = sorted(pdays)
            consec = 1
            for i in range(1, len(sorted_d)):
                try:
                    d1 = _dt.strptime(sorted_d[i-1], "%Y-%m-%d")
                    d2 = _dt.strptime(sorted_d[i], "%Y-%m-%d")
                    if (d2-d1).days <= 3: consec += 1
                    else: consec = 1
                    if consec > 3: soft -= 50 * (consec - 3)  # عقوبة صارمة
                except: pass

        # ── عقوبة الأيام المنعزلة ──
        for prof, pdays in prof_days_sched.items():
            day_counts = {}
            for mid, sv in schedule.items():
                if not sv or sv[0] not in pdays: continue
                if prof not in memo_members.get(mid, set()): continue
                day_counts[sv[0]] = day_counts.get(sv[0], 0) + 1
            total_p = sum(day_counts.values())
            if total_p >= 3:
                lonely = sum(1 for d, c in day_counts.items() if c == 1)
                if lonely > 2:
                    soft -= 500 * (lonely - 2)  # عقوبة حرجة جداً
                elif lonely == 2:
                    soft -= 50  # تنبيه فقط

        return placement_score + soft

    # ── Crossover: دمج جدولين ──
    def crossover(parent1, parent2):
        child = {}
        occupied_c = {}
        prof_busy_c = {}
        prof_day_c = {}

        # لكل مذكرة — اختر من أي والد بناءً على الـ fitness المحلي
        memos_shuffled = memo_ids[:]
        random.shuffle(memos_shuffled)
        cutpoint = len(memos_shuffled) // 2

        for i, mid in enumerate(memos_shuffled):
            sv = parent1.get(mid) if i < cutpoint else parent2.get(mid)
            if not sv:
                child[mid] = None; continue

            day, slot, room = sv
            # تحقق من عدم التعارض
            conflict = False
            if (day, slot, room) in occupied_c:
                conflict = True
            else:
                for prof in memo_members.get(mid, set()):
                    if (day, slot, prof) in prof_busy_c:
                        conflict = True; break
                    if prof_day_c.get((prof, day), 0) >= 3:
                        conflict = True; break

            if not conflict:
                child[mid] = sv
                occupied_c[(day, slot, room)] = mid
                for prof in memo_members.get(mid, set()):
                    prof_busy_c[(day, slot, prof)] = mid
                    prof_day_c[(prof, day)] = prof_day_c.get((prof, day), 0) + 1
            else:
                child[mid] = None  # سيُعالج في repair

        # Repair: برمجة المذكرات غير المجدولة
        schedule_c = dict(child)
        occupied_r, prof_busy_r, prof_day_r, ppc_r = {}, {}, {}, {}
        for mid, sv in schedule_c.items():
            if not sv: continue
            day, slot, room = sv
            occupied_r[(day, slot, room)] = mid
            for prof in memo_members.get(mid, set()):
                prof_busy_r[(day, slot, prof)] = mid
                prof_day_r[(prof, day)] = prof_day_r.get((prof, day), 0) + 1

        can_r, _ = make_can_place(occupied_r, prof_busy_r, prof_day_r, memo_members,
            fixed_slots, memo_date_limits, prof_banned_days, prof_not_before,
            prof_not_after, prof_allowed_days, profs_accept_18, memo_alt_days,
            prof_phase_split, ppc_r, slots_per_day, None)
        place_r = make_place_fn(schedule_c, occupied_r, prof_busy_r, prof_day_r,
                                memo_members, ppc_r, prof_phase_split)

        for mid in memo_ids:
            if schedule_c.get(mid): continue
            for day in days:
                for slot in slots_per_day:
                    for room in rooms:
                        if can_r(mid, day, slot, room, log=False):
                            place_r(mid, day, slot, room); break
                    else: continue; break
                else: continue; break

        for mid in memo_ids:
            if mid not in schedule_c: schedule_c[mid] = None

        return schedule_c

    # ── Mutation: تغيير عشوائي ──
    def mutate(schedule, rate=0.05):
        mutated = dict(schedule)
        occupied_m = {sv: mid for mid, sv in mutated.items() if sv}
        prof_busy_m = {}
        prof_day_m = {}
        for mid, sv in mutated.items():
            if not sv: continue
            day, slot, room = sv
            for prof in memo_members.get(mid, set()):
                prof_busy_m[(day, slot, prof)] = mid
                prof_day_m[(prof, day)] = prof_day_m.get((prof, day), 0) + 1

        for mid in memo_ids:
            if random.random() > rate: continue
            if mid in fixed_slots: continue  # لا نغير المثبتة
            old_sv = mutated.get(mid)
            if old_sv:
                day, slot, room = old_sv
                del occupied_m[old_sv]
                for prof in memo_members.get(mid, set()):
                    if (day, slot, prof) in prof_busy_m:
                        del prof_busy_m[(day, slot, prof)]
                    prof_day_m[(prof, day)] = max(0, prof_day_m.get((prof, day), 1) - 1)

            can_m, _ = make_can_place(occupied_m, prof_busy_m, prof_day_m, memo_members,
                fixed_slots, memo_date_limits, prof_banned_days, prof_not_before,
                prof_not_after, prof_allowed_days, profs_accept_18, memo_alt_days,
                prof_phase_split, {}, slots_per_day, None)

            placed = False
            d_order = days[:]; random.shuffle(d_order)
            for day in d_order:
                if placed: break
                for slot in slots_per_day:
                    if placed: break
                    for room in rooms:
                        if can_m(mid, day, slot, room, log=False):
                            mutated[mid] = (day, slot, room)
                            occupied_m[(day, slot, room)] = mid
                            for prof in memo_members.get(mid, set()):
                                prof_busy_m[(day, slot, prof)] = mid
                                prof_day_m[(prof, day)] = prof_day_m.get((prof, day), 0) + 1
                            placed = True; break
            if not placed:
                mutated[mid] = None

        return mutated

    # ── المرحلة 1: Genetic Algorithm ──
    POPULATION_SIZE = 20
    GENERATIONS = 30
    ELITE_SIZE = 4
    MUTATION_RATE = 0.08

    if streamlit_progress:
        streamlit_progress.text("🧬 المرحلة 1: توليد الجيل الأول...")

    # توليد الجيل الأول
    population = []
    for i in range(POPULATION_SIZE):
        # حاول توليد جدول بلا أيام منعزلة زائدة
        best_attempt = None
        for attempt in range(5):
            ind = generate_individual(seed=i * 13 + 7 + attempt * 97)
            lonely_v = ind.pop("__lonely_violations__", 99)
            if lonely_v == 0:
                best_attempt = ind; break
            if best_attempt is None or lonely_v < best_attempt.get("__lv__", 99):
                ind["__lv__"] = lonely_v
                best_attempt = ind
        if "__lv__" in best_attempt: del best_attempt["__lv__"]
        population.append(best_attempt)
        if streamlit_progress:
            streamlit_progress.progress(int((i+1)/POPULATION_SIZE * 30))

    best_overall = max(population, key=fitness)
    best_fitness = fitness(best_overall)

    # تطوير الأجيال
    for gen in range(GENERATIONS):
        # تقييم الجيل
        scored = sorted(population, key=fitness, reverse=True)
        gen_best = scored[0]
        gen_best_fit = fitness(gen_best)

        if gen_best_fit > best_fitness:
            best_fitness = gen_best_fit
            best_overall = dict(gen_best)

        if streamlit_progress:
            placed_n = sum(1 for sv in best_overall.values() if sv)
            streamlit_progress.text(f"🧬 الجيل {gen+1}/{GENERATIONS} | أفضل: {placed_n}/{len(memo_ids)} مجدول | fitness: {best_fitness:.0f}")
            streamlit_progress.progress(30 + int((gen+1)/GENERATIONS * 40))

        # Elite selection
        new_pop = scored[:ELITE_SIZE]

        # Crossover + Mutation
        while len(new_pop) < POPULATION_SIZE:
            p1, p2 = random.choices(scored[:10], k=2)
            child = crossover(p1, p2)
            child = mutate(child, rate=MUTATION_RATE)
            new_pop.append(child)

        population = new_pop

    # ── المرحلة 2: Tabu Search على أفضل جدول ──
    if streamlit_progress:
        streamlit_progress.text("🔍 المرحلة 2: Tabu Search — تحسين أفضل جدول...")

    current = dict(best_overall)
    current_fit = best_fitness
    tabu_list = set()  # تعقب المبادلات المزارة
    TABU_SIZE = 50
    TABU_ITERATIONS = 200

    scheduled_list = [m for m, sv in current.items() if sv and m not in fixed_slots]

    for it in range(TABU_ITERATIONS):
        if streamlit_progress and it % 20 == 0:
            placed_n = sum(1 for sv in current.values() if sv)
            streamlit_progress.text(f"🔍 Tabu Iteration {it}/{TABU_ITERATIONS} | مجدول: {placed_n} | fitness: {current_fit:.0f}")
            streamlit_progress.progress(70 + int(it/TABU_ITERATIONS * 25))

        if len(scheduled_list) < 2: break

        best_neighbor = None
        best_neighbor_fit = -float("inf")
        best_move = None

        # استكشف جيرة المبادلة
        candidates = random.sample(scheduled_list, min(20, len(scheduled_list)))
        for i in range(len(candidates)):
            for j in range(i+1, len(candidates)):
                m1, m2 = candidates[i], candidates[j]
                move = (m1, m2) if m1 < m2 else (m2, m1)
                if move in tabu_list: continue

                sv1, sv2 = current[m1], current[m2]
                if not sv1 or not sv2: continue

                # جرب المبادلة
                current[m1] = sv2; current[m2] = sv1

                # تحقق من Hard Constraints
                v = _validate_hard_constraints(current, memo_members)
                if not any(x.startswith("🔴") for x in v):
                    nf = fitness(current)
                    if nf > best_neighbor_fit:
                        best_neighbor_fit = nf
                        best_neighbor = (dict(current), nf)
                        best_move = move

                # أعد الحال
                current[m1] = sv1; current[m2] = sv2

        if best_neighbor:
            current, current_fit = best_neighbor
            tabu_list.add(best_move)
            if len(tabu_list) > TABU_SIZE:
                tabu_list.pop()
            if current_fit > best_fitness:
                best_fitness = current_fit
                best_overall = dict(current)

    if streamlit_progress:
        placed_f = sum(1 for sv in best_overall.values() if sv)
        streamlit_progress.text(f"✅ اكتمل! مجدول: {placed_f}/{len(memo_ids)} | fitness نهائي: {best_fitness:.0f}")
        streamlit_progress.progress(100)

    # rej_log
    rej_log = {mid: {"لا توجد خانة متاحة مع القيود"} for mid, sv in best_overall.items() if not sv}

    return best_overall, memo_members, rej_log



def multi_start_best(df_memos, days, slots_per_day, rooms, constraints,
                     n_tries=20, algo_name="🧱 كتل الأساتذة",
                     progress_cb=None):
    """
    Multi-Start Optimizer — يشغل الخوارزمية n_tries مرة
    ويختار الجدول الأفضل بناءً على score مركّب
    """
    import random

    fixed_slots   = constraints[0]
    prof_banned   = constraints[2]
    prof_allowed  = constraints[6]

    # ── دالة التقييم ──
    def score_schedule(sched, memo_members):
        placed   = sum(1 for sv in sched.values() if sv)
        total    = len(sched)
        if total == 0: return -999999

        # انتهاكات حرجة
        violations = _validate_hard_constraints(sched, memo_members)
        critical   = len([v for v in violations if v.startswith("🔴")])
        warnings   = len([v for v in violations if v.startswith("🟡")])

        # توزيع يومي
        day_counts = {}
        for sv in sched.values():
            if sv: day_counts[sv[0]] = day_counts.get(sv[0], 0) + 1
        avg = sum(day_counts.values()) / max(len(day_counts), 1)
        imbalance = sum(abs(c - avg) for c in day_counts.values())

        score = (placed * 100
                 - critical * 500
                 - warnings  * 30
                 - imbalance * 5)
        return score

    best_schedule   = None
    best_score      = -999999
    best_memo_mbrs  = {}
    best_rej        = {}

    for i in range(n_tries):
        seed = i * 37 + 7
        try:
            sched, memo_mbrs, rej = run_algorithm(
                algo_name, df_memos, days, slots_per_day, rooms,
                constraints, improve=False, seed=seed
            )[0], None, {}
            # run_algorithm returns tuple — unpack properly
            result = run_algorithm(
                algo_name, df_memos, days, slots_per_day, rooms,
                constraints, improve=False, seed=seed
            )
            sched     = result[0]
            memo_mbrs = result[6]
            rej       = result[7]
        except Exception as _e:
            continue

        sc = score_schedule(sched, memo_mbrs)
        if sc > best_score:
            best_score    = sc
            best_schedule = dict(sched)
            best_memo_mbrs= memo_mbrs
            best_rej      = rej

        if progress_cb:
            progress_cb(i + 1, n_tries, best_score,
                        sum(1 for sv in best_schedule.values() if sv) if best_schedule else 0)

    # تحسين نهائي على أفضل جدول
    if best_schedule:
        _pbd  = constraints[2]
        _pad  = constraints[6]
        _pa18 = constraints[11]
        _fix  = constraints[0]
        best_schedule = improve_schedule(
            best_schedule, best_memo_mbrs, days, slots_per_day, rooms,
            iterations=2000, prof_banned_days=_pbd,
            prof_allowed_days=_pad, profs_accept_18=_pa18,
            fixed_slots=_fix
        )
        # أعد تطبيق المثبتات بالقوة
        for fmid, fsv in _fix.items():
            fd, fs, fr = fsv
            if fd and fs and str(fmid) in best_schedule:
                best_schedule[str(fmid)] = (fd, fs, fr if fr else rooms[0])

    return best_schedule, best_memo_mbrs, best_rej, best_score


def run_algorithm(algo_name, df_memos, days, slots_per_day, rooms, constraints, improve=True, seed=42):
    """تشغيل الخوارزمية المختارة"""
    import random
    random.seed(seed)

    # ── تصفية المذكرات المجمّدة قبل أي شيء ──
    fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after,     prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split,     memo_alt_days, profs_accept_18, profs_cluster_days = constraints

    if frozen_profs:
        _, _all_members = build_prof_memo_map(df_memos)
        _frozen_memos = {mid for mid, members in _all_members.items() if members & frozen_profs}
        if _frozen_memos:
            import streamlit as _st
            _st.warning(f"🔒 {len(_frozen_memos)} مذكرة استُثنيت بسبب أساتذة مجمّدين: {', '.join(sorted(_frozen_memos)[:10])}")
        df_memos = df_memos[~df_memos["رقم المذكرة"].astype(str).isin(_frozen_memos)].copy()

    if algo_name == "🧱 كتل الأساتذة":
        schedule, memo_members, rej_log = algo_blocks(df_memos, days, slots_per_day, rooms, constraints)
    elif algo_name == "📅 الجدول أولاً":
        schedule, memo_members, rej_log = algo_day_first(df_memos, days, slots_per_day, rooms, constraints)
    elif algo_name == "🧬 GA + Tabu Search":
        import streamlit as _st_ga
        _prog_container = _st_ga.empty()
        _prog_bar = _st_ga.progress(0)
        class _ProgHelper:
            def text(self, t): _prog_container.text(t)
            def progress(self, v): _prog_bar.progress(v)
        schedule, memo_members, rej_log = ga_tabu_scheduler(df_memos, days, slots_per_day, rooms, constraints, streamlit_progress=_ProgHelper())
        _prog_container.empty(); _prog_bar.empty()
        quality, placed, unplaced, idle, total_days, _ = calc_schedule_quality(schedule, memo_members, days, slots_per_day)
        return schedule, quality, placed, unplaced, idle, total_days, memo_members, rej_log
    else:  # Greedy
        schedule, memo_members, rej_log = algo_greedy(df_memos, days, slots_per_day, rooms, constraints)

    if improve:
        _pbd = constraints[2] if constraints else {}
        _pad = constraints[6] if constraints else {}
        _pa18 = constraints[11] if constraints else set()
        schedule = improve_schedule(schedule, memo_members, days, slots_per_day, rooms, iterations=2000,
                                   prof_banned_days=_pbd, prof_allowed_days=_pad, profs_accept_18=_pa18,
                                   fixed_slots=constraints[0] if constraints else {})
        # ── إعادة تطبيق المثبتات بالقوة بعد improve_schedule ──
        _fixed = constraints[0] if constraints else {}
        for _fmid, _fsv in _fixed.items():
            _fd, _fs, _fr = _fsv
            if not _fd or not _fs: continue
            # أزل من موقعه الحالي
            if str(_fmid) in schedule and schedule[str(_fmid)] != (_fd, _fs, _fr):
                schedule[str(_fmid)] = (_fd, _fs, _fr if _fr else (list(rooms)[0] if rooms else ""))

    # ── تحقق صارم ما بعد التوليد وإصلاح انتهاكات الأيام الممنوعة ──
    _pbd = constraints[2] if constraints else {}
    _pad = constraints[6] if constraints else {}
    if _pbd:
        _violations_found = []
        for mid, sv in list(schedule.items()):
            if not sv: continue
            day, slot, room = sv
            for prof in memo_members.get(mid, set()):
                if day in _pbd.get(prof, set()):
                    _violations_found.append((mid, prof, day))
                if _pad.get(prof) and day not in _pad[prof]:
                    _violations_found.append((mid, prof, day))
        
        # أصلح كل انتهاك — انقل المذكرة ليوم صالح
        if _violations_found:
            _fixed_count = 0
            for mid, viol_prof, bad_day in _violations_found:
                old_sv = schedule.get(mid)
                if not old_sv: continue
                # ابحث عن خانة بديلة
                _occ = {sv: m for m, sv in schedule.items() if sv}
                _pbusy = {}
                _pdc = {}
                for m2, sv2 in schedule.items():
                    if not sv2: continue
                    d2, s2, r2 = sv2
                    for p2 in memo_members.get(m2, set()):
                        _pbusy[(d2, s2, p2)] = m2
                        _pdc[(p2, d2)] = _pdc.get((p2, d2), 0) + 1
                
                placed_alt = False
                for day in days:
                    if placed_alt: break
                    # تحقق أن هذا اليوم مقبول لكل الأعضاء
                    ok_for_all = True
                    for p in memo_members.get(mid, set()):
                        if day in _pbd.get(p, set()):
                            ok_for_all = False; break
                        if _pad.get(p) and day not in _pad[p]:
                            ok_for_all = False; break
                    if not ok_for_all: continue
                    
                    for slot in slots_per_day:
                        if placed_alt: break
                        for room in rooms:
                            if (day, slot, room) in _occ: continue
                            conflict = False
                            for p in memo_members.get(mid, set()):
                                if (day, slot, p) in _pbusy: conflict = True; break
                                if _pdc.get((p, day), 0) >= 3: conflict = True; break
                            if not conflict:
                                schedule[mid] = (day, slot, room)
                                placed_alt = True
                                _fixed_count += 1
                                break

    quality, placed, unplaced, idle, total_days, _ = calc_schedule_quality(schedule, memo_members, days, slots_per_day)
    return schedule, quality, placed, unplaced, idle, total_days, memo_members, rej_log


def run_smart_schedule(df_memos, days, slots_per_day, rooms, max_per_day=3, max_consecutive=3, improve=True,
                       fixed_slots=None, memo_date_limits=None, prof_banned_days=None,
                       prof_not_before=None, prof_not_after=None, prof_one_day=None,
                       prof_allowed_days=None, prof_consecutive=None,
                       frozen_profs=None, prof_phase_split=None,
                       memo_alt_days=None, day_time_limits=None,
                       profs_accept_18=None, profs_cluster_days=None):
    """الدالة الرئيسية للجدولة الذكية"""
    fixed_slots = fixed_slots or {}
    memo_date_limits = memo_date_limits or {}
    prof_banned_days = prof_banned_days or {}
    prof_not_before = prof_not_before or {}
    prof_not_after = prof_not_after or {}
    prof_one_day = prof_one_day or set()
    prof_allowed_days = prof_allowed_days or {}
    prof_consecutive = prof_consecutive or set()
    frozen_profs = frozen_profs or set()
    prof_phase_split = prof_phase_split or {}
    memo_alt_days = memo_alt_days or {}
    day_time_limits = day_time_limits or {}
    profs_accept_18 = profs_accept_18 or set()
    profs_cluster_days = profs_cluster_days or set()

    # تصفية المذكرات التي تحتوي أساتذة مجمّدين
    if frozen_profs:
        _, _memo_mbrs = build_prof_memo_map(df_memos)
        _frozen_memos = {mid for mid, members in _memo_mbrs.items() if members & frozen_profs}
        df_memos = df_memos[~df_memos["رقم المذكرة"].astype(str).isin(_frozen_memos)].copy()
        if _frozen_memos:
            import streamlit as _st
            _st.warning(f"⚠️ {len(_frozen_memos)} مذكرة استُثنيت بسبب أستاذ مجمّد: {', '.join(sorted(_frozen_memos)[:10])}")

    # المرحلة 1: Professor-First Scheduling مع الاستثناءات
    schedule, memo_members, _rej_log = professor_first_schedule(
        df_memos, days, slots_per_day, rooms, max_per_day, max_consecutive,
        fixed_slots=fixed_slots, memo_date_limits=memo_date_limits,
        prof_banned_days=prof_banned_days, prof_not_before=prof_not_before,
        prof_not_after=prof_not_after, prof_one_day=prof_one_day,
        prof_allowed_days=prof_allowed_days, prof_consecutive=prof_consecutive,
        frozen_profs=frozen_profs, prof_phase_split=prof_phase_split,
        memo_alt_days=memo_alt_days, day_time_limits=day_time_limits or {},
        profs_accept_18=profs_accept_18,
        profs_cluster_days=profs_cluster_days
    )
    
    # المرحلة 2: تحسين الجدول
    if improve:
        schedule = improve_schedule(schedule, memo_members, days, slots_per_day, rooms, iterations=400)
    
    # حساب الجودة
    quality, placed, unplaced, idle, total_days, max_gap = calc_schedule_quality(
        schedule, memo_members, days, slots_per_day
    )
    
    return schedule, quality, placed, unplaced, idle, total_days, memo_members, _rej_log

def schedule_to_rows(schedule, df_memos):
    """تحويل الجدول لصفوف عرض"""
    memo_rows = {str(r["رقم المذكرة"]): r for _, r in df_memos.iterrows()}
    rows = []
    for mid, slot in schedule.items():
        row = memo_rows.get(mid, pd.Series())
        r = {
            "رقم المذكرة": mid,
            "العنوان": str(row.get("عنوان المذكرة",""))[:40] if hasattr(row,"get") else "",
            "اليوم": slot[0] if slot else "غير مجدول",
            "التوقيت": slot[1] if slot else "",
            "القاعة": slot[2] if slot else "",
            "المشرف": str(row.get("الأستاذ","")).strip() if hasattr(row,"get") else "",
            "الرئيس": str(row.get("الرئيس","")).strip() if hasattr(row,"get") else "",
            "المناقش1": str(row.get("المناقش1","")).strip() if hasattr(row,"get") else "",
            "المناقش2": str(row.get("المناقش2","")).strip() if hasattr(row,"get") else "",
            "رابط الملف": str(row.get("رابط الملف","")).strip() if hasattr(row,"get") else "",
        }
        rows.append(r)
    return rows

def save_full_schedule_to_sheets(schedule, df_memos):
    """حفظ الجدول الكامل في Google Sheets"""
    try:
        # أولاً: اقرأ الشيت الكامل للحصول على الصفوف الحقيقية
        all_memos = load_memos()
        # بناء خريطة رقم المذكرة → رقم الصف الحقيقي في الشيت
        row_map = {}
        for i, (_, row) in enumerate(all_memos.iterrows()):
            mid = str(row.get("رقم المذكرة","")).strip()
            if mid and mid not in ["","nan"]:
                row_map[mid] = i + 2  # +2 لأن الصف الأول هو الرأس

        updates = []
        for mid, slot in schedule.items():
            if not slot: continue
            row_idx = row_map.get(str(mid))
            if not row_idx: continue
            updates += [
                {"range": f"Feuille 1!W{row_idx}", "values": [[slot[0]]]},
                {"range": f"Feuille 1!X{row_idx}", "values": [[slot[1]]]},
                {"range": f"Feuille 1!Y{row_idx}", "values": [[slot[2] if slot[2] else ""]]},
            ]
        if updates:
            for i in range(0, len(updates), 100):
                sheets_service.spreadsheets().values().batchUpdate(
                    spreadsheetId=MEMOS_SHEET_ID,
                    body={"valueInputOption": "USER_ENTERED", "data": updates[i:i+100]}
                ).execute()
            clear_cache_and_reload()
            return True, f"✅ تم حفظ {len([s for s in schedule.values() if s])} مذكرة — لم يتم النشر بعد، الإدارة تتحكم في النشر من الشيت"
        return False, "لا شيء للحفظ"
    except Exception as e:
        return False, f"❌ {str(e)}"

def send_prof_schedule_email(prof_name, items, df_profs):
    """إرسال جدول المناقشات لأستاذ"""
    try:
        rows = df_profs[df_profs["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
        if rows.empty: return False, "غير موجود"
        email = get_email_smart(rows.iloc[0])
        if not email or "@" not in email: return False, "لا بريد"
        by_day = {}
        for it in sorted(items, key=lambda x: (x["اليوم"], x["التوقيت"])):
            by_day.setdefault(it["اليوم"], []).append(it)
        rows_html = ""
        for day, its in sorted(by_day.items()):
            rows_html += f'<tr style="background:#e8f4f8;"><td colspan="4" style="padding:8px;font-weight:900;color:#1e3a5f;">📅 {day}</td></tr>'
            for it in its:
                lnk = f'<a href="{it["رابط الملف"]}" style="color:#2F6F7E;">📄</a>' if it.get("رابط الملف","") not in ["","nan"] else "—"
                rows_html += f'<tr><td style="padding:8px;border-bottom:1px solid #eee;">{it["التوقيت"]}</td><td style="padding:8px;border-bottom:1px solid #eee;">{it["القاعة"]}</td><td style="padding:8px;border-bottom:1px solid #eee;">{it["رقم المذكرة"]}</td><td style="padding:8px;border-bottom:1px solid #eee;">{it["الصفة"]} {lnk}</td></tr>'
        body = f'''<html dir="rtl"><head><meta charset="UTF-8"><style>body{{font-family:Arial;direction:rtl;background:#f4f4f4;padding:20px}}.c{{background:#fff;padding:28px;border-radius:12px;max-width:700px;margin:auto}}.h{{background:linear-gradient(135deg,#0F2942,#2F6F7E);color:#fff;padding:20px;border-radius:8px;text-align:center;margin-bottom:18px}}table{{width:100%;border-collapse:collapse}}th{{background:#2F6F7E;color:#fff;padding:10px;text-align:right}}</style></head><body><div class="c"><div class="h"><h2>📋 برنامج مناقشاتك</h2><p style="opacity:.9">جامعة محمد البشير الإبراهيمي</p></div><p>الأستاذ(ة) <strong>{prof_name}</strong>، فيما يلي برنامج مناقشاتك: <strong>{len(items)} مناقشة</strong> في <strong>{len(by_day)} يوم</strong>.</p><table><thead><tr><th>التوقيت</th><th>القاعة</th><th>المذكرة</th><th>صفتك</th></tr></thead><tbody>{rows_html}</tbody></table><p style="background:#fff8e1;padding:12px;border-right:4px solid #F59E0B;border-radius:6px;margin-top:16px;">⚠️ يُرجى الحضور قبل الموعد بـ 15 دقيقة.</p><div style="text-align:center;margin:16px 0"><a href="https://memoires2026.streamlit.app" style="background:#2F6F7E;color:#fff;padding:12px 28px;border-radius:8px;text-decoration:none;font-weight:bold">🔗 الدخول للمنصة</a></div></div></body></html>'''
        msg = MIMEMultipart("alternative")
        msg["From"] = EMAIL_SENDER
        msg["To"] = email
        msg["Subject"] = f"📋 برنامج مناقشاتك — {len(items)} مناقشة في {len(by_day)} يوم"
        msg.attach(MIMEText(body, "html", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as s:
            s.starttls(); s.login(EMAIL_SENDER, EMAIL_PASSWORD); s.send_message(msg)
        return True, f"✅ {email}"
    except Exception as e:
        return False, f"❌ {str(e)}"

def send_student_schedule_email(student_data, memo_num, defense_date, defense_time, defense_room):
    """إرسال موعد المناقشة للطالب"""
    try:
        email = get_email_smart(student_data)
        if not email or "@" not in email: return False, "لا بريد"
        ln, fn = get_student_name_display(student_data)
        name = f"{ln} {fn}".strip()
        body = f'''<html dir="rtl"><head><meta charset="UTF-8"><style>body{{font-family:Arial;direction:rtl;background:#f4f4f4;padding:20px}}.c{{background:#fff;padding:28px;border-radius:12px;max-width:600px;margin:auto}}.h{{background:linear-gradient(135deg,#0F2942,#2F6F7E);color:#fff;padding:20px;border-radius:8px;text-align:center}}.s{{background:#f0fdf4;padding:14px;border-right:4px solid #10B981;margin:12px 0;border-radius:6px}}</style></head><body><div class="c"><div class="h"><h2>📅 موعد مناقشتك</h2></div><p>تحية طيبة، <strong>{name}</strong>،</p><p>تم تحديد موعد مناقشة مذكرتك رقم <strong>{memo_num}</strong>.</p><div class="s"><p>📆 <strong>التاريخ:</strong> {defense_date}</p><p>🕐 <strong>التوقيت:</strong> {defense_time}</p><p>🏛️ <strong>القاعة:</strong> {defense_room}</p></div><p style="background:#fff8e1;padding:12px;border-right:4px solid #F59E0B;border-radius:6px">⚠️ يُرجى الحضور قبل الموعد بـ 15 دقيقة مع جميع الوثائق.</p><div style="text-align:center;margin:16px 0"><a href="https://memoires2026.streamlit.app" style="background:#2F6F7E;color:#fff;padding:12px 28px;border-radius:8px;text-decoration:none;font-weight:bold">🔗 متابعة على المنصة</a></div></div></body></html>'''
        msg = MIMEMultipart("alternative")
        msg["From"] = EMAIL_SENDER
        msg["To"] = email
        msg["Subject"] = f"📅 موعد مناقشتك — {defense_date} الساعة {defense_time}"
        msg.attach(MIMEText(body, "html", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as s:
            s.starttls(); s.login(EMAIL_SENDER, EMAIL_PASSWORD); s.send_message(msg)
        return True, f"✅ {email}"
    except Exception as e:
        return False, f"❌ {str(e)}"

def save_member_observations(memo_number, prof_name, role, observations):
    """حفظ ملاحظات عضو اللجنة"""
    col_map = {"مشرف":"Z","رئيس لجنة":"AE","مناقش":"AF"}
    col = col_map.get(role, "AE")
    try:
        df_m = load_memos()
        row = df_m[df_m["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False, "❌ غير موجودة"
        row_idx = row.index[0] + 2
        ts = datetime.now().strftime("%Y-%m-%d %H:%M")
        obs_full = f"[{prof_name}—{role}][{ts}]: {observations}"
        sheets_service.spreadsheets().values().update(
            spreadsheetId=MEMOS_SHEET_ID,
            range=f"Feuille 1!{col}{row_idx}",
            valueInputOption="USER_ENTERED",
            body={"values": [[obs_full]]}
        ).execute()
        clear_cache_and_reload()
        return True, "✅ تم حفظ الملاحظات"
    except Exception as e:
        return False, f"❌ {str(e)}"

def clear_missing_flag(memo_number):
    """إزالة علامة المفقودة"""
    try:
        df_memos = load_memos()
        row = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==normalize_text(memo_number)]
        if row.empty: return False
        row_idx = row.index[0] + 2
        sheets_service.spreadsheets().values().update(
            spreadsheetId=MEMOS_SHEET_ID,
            range=f"Feuille 1!AH{row_idx}",
            valueInputOption="USER_ENTERED",
            body={"values": [["0"]]}
        ).execute()
        clear_cache_and_reload()
        return True
    except: return False

def send_recovery_email_to_admin(memo_number, memo_title, student1_name, student2_name=None):
    """إرسال إيميل للإدارة عند استرجاع مذكرة مفقودة"""
    try:
        s2 = f" / {student2_name}" if student2_name else ""
        body = f'''<html dir="rtl"><head><meta charset="UTF-8"><style>body{{font-family:Arial;direction:rtl;background:#f4f4f4;padding:20px}}.c{{background:#fff;padding:28px;border-radius:12px;max-width:600px;margin:auto}}.h{{background:linear-gradient(135deg,#0F2942,#1a472a);color:#fff;padding:20px;border-radius:8px;text-align:center}}.badge{{background:#10B981;color:#fff;padding:4px 14px;border-radius:20px;font-weight:700}}</style></head><body><div class="c"><div class="h"><h2>✅ استرجاع مذكرة مفقودة</h2></div><table style="width:100%;border-collapse:collapse;margin:16px 0"><tr><td style="padding:10px;border:1px solid #e2e8f0;font-weight:700">رقم المذكرة</td><td style="padding:10px;border:1px solid #e2e8f0"><span class="badge">{memo_number}</span></td></tr><tr><td style="padding:10px;border:1px solid #e2e8f0;font-weight:700">العنوان</td><td style="padding:10px;border:1px solid #e2e8f0">{memo_title}</td></tr><tr><td style="padding:10px;border:1px solid #e2e8f0;font-weight:700">الطالب</td><td style="padding:10px;border:1px solid #e2e8f0">{student1_name}{s2}</td></tr><tr><td style="padding:10px;border:1px solid #e2e8f0;font-weight:700">التاريخ</td><td style="padding:10px;border:1px solid #e2e8f0">{datetime.now().strftime("%Y-%m-%d %H:%M")}</td></tr></table></div></body></html>'''
        msg = MIMEMultipart("alternative")
        msg["From"] = EMAIL_SENDER
        msg["To"] = EMAIL_SENDER
        msg["Subject"] = f"✅ استرجاع مذكرة — رقم {memo_number}"
        msg.attach(MIMEText(body, "html", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls(); server.login(EMAIL_SENDER, EMAIL_PASSWORD); server.send_message(msg)
        return True, "✅"
    except Exception as e:
        return False, f"❌ {str(e)}"


def send_jury_notification_email(prof_row, has_pending_memos=False, pending_memo_list=None):
    """إرسال إيميل إعلام بلجان المناقشة للأستاذ"""
    try:
        email = get_email_smart(prof_row)
        if not email or "@" not in email:
            return False, "لا بريد"
        
        # اسم المستخدم وكلمة المرور من عمود M و N
        cols = list(prof_row.index) if hasattr(prof_row, 'index') else []
        username = str(prof_row.iloc[12]).strip() if len(prof_row) > 12 else ""
        password = str(prof_row.iloc[13]).strip() if len(prof_row) > 13 else ""
        
        # اسم الأستاذ
        prof_name_e = str(prof_row.get("الأستاذ", "")).strip() if hasattr(prof_row, 'get') else ""
        if not prof_name_e:
            prof_name_e = str(prof_row.iloc[0]).strip() if len(prof_row) > 0 else "الأستاذ(ة)"

        # فقرة المذكرات العالقة (للمشرف فقط)
        pending_section = ""
        if has_pending_memos and pending_memo_list:
            memo_items = "".join([f'<li style="margin:4px 0;"><strong>مذكرة رقم {m}</strong></li>' for m in pending_memo_list])
            pending_section = f"""
            <div style="background:#fff8e1;border-right:4px solid #F59E0B;border-radius:10px;padding:18px 22px;margin:20px 0;">
                <div style="font-weight:700;color:#92400E;margin-bottom:10px;font-size:1.05rem;">⚠️ تنبيه خاص — مذكرات تنتظر رأيكم</div>
                <div style="color:#78350F;line-height:1.9;margin-bottom:12px;">
                    لا تزال <strong>{len(pending_memo_list)} مذكرة</strong> مسندة إليكم في انتظار رأيكم النهائي:
                </div>
                <ul style="color:#78350F;padding-right:20px;margin:0 0 12px;">
                    {memo_items}
                </ul>
                <div style="color:#78350F;line-height:1.9;">
                    ندعوكم إلى الولوج إلى المنصة وإبداء رأيكم بشأنها، سواء بالموافقة أو بإرسال ملاحظاتكم للطالب المعني،
                    وذلك حتى يتسنّى إدراجها ضمن الجدول الرسمي للمناقشات.
                </div>
            </div>"""
        elif has_pending_memos:
            pending_section = """
            <div style="background:#fff8e1;border-right:4px solid #F59E0B;border-radius:10px;padding:18px 22px;margin:20px 0;">
                <div style="font-weight:700;color:#92400E;margin-bottom:8px;">⚠️ تنبيه خاص</div>
                <div style="color:#78350F;line-height:1.9;">
                    لا تزال بعض المذكرات المسندة إليكم في انتظار رأيكم النهائي. ندعوكم إلى الولوج إلى المنصة وإبداء رأيكم بشأنها.
                </div>
            </div>"""

        body = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
body {{ font-family: Arial, sans-serif; direction: rtl; text-align: right; background: #f4f6f8; margin: 0; padding: 20px; }} * {{ direction: rtl; text-align: right; }} td {{ text-align: right; }}
.container {{ background: #ffffff; max-width: 650px; margin: auto; border-radius: 14px; overflow: hidden; box-shadow: 0 4px 20px rgba(0,0,0,0.08); }}
.header {{ background: linear-gradient(135deg, #0F2942, #1A4A6E); padding: 28px 32px; text-align: center; }}
.header h2 {{ color: #FFD700; font-size: 1.3rem; margin: 0 0 6px; }}
.header p {{ color: rgba(255,255,255,0.85); font-size: 0.9rem; margin: 0; }}
.body {{ padding: 28px 32px; color: #1e293b; line-height: 1.9; font-size: 1.1rem; }}
.credentials {{ background: #f0f9ff; border: 1px solid #bae6fd; border-radius: 10px; padding: 16px 20px; margin: 20px 0; }}
.credentials table {{ width: 100%; border-collapse: collapse; }}
.credentials td {{ padding: 6px 10px; }}
.credentials .label {{ color: #0369a1; font-weight: 700; width: 40%; }}
.credentials .value {{ color: #0F2942; font-weight: 900; font-size: 1rem; }}
.btn {{ display: inline-block; background: linear-gradient(135deg, #0F2942, #2F6F7E); color: #FFD700 !important; padding: 12px 30px; border-radius: 25px; text-decoration: none; font-weight: 700; font-size: 0.95rem; margin-top: 16px; }}
.footer {{ background: #f8fafc; padding: 16px 32px; text-align: center; color: #64748b; font-size: 0.8rem; border-top: 1px solid #e2e8f0; }}
</style>
</head>
<body>
<div class="container">
    <div class="header" style="text-align:center;" dir="rtl">
        <h2 style="font-size:1.5rem;text-align:center;">🎓 منصة مذكرات الماستر 2025-2026</h2>
        <p style="font-size:1.05rem;text-align:center;">كلية الحقوق والعلوم السياسية — جامعة محمد البشير الإبراهيمي</p>
    </div>
    <div class="body">
        <p>الأستاذ(ة) الفاضل(ة) <strong>{prof_name_e}</strong>،</p>
        <p>تحية طيبة وبعد،</p>
        <p>
        نتقدم إليكم بجزيل الشكر على تفاعلكم مع منصة مذكرات الماستر ومساهمتكم في إنجاح مشروع رقمنة تسيير مذكرات التخرج بكلية الحقوق والعلوم السياسية.
        </p>
        <p>
        يسرّنا إعلامكم بانطلاق مناقشات مذكرات الماستر ابتداءً من يوم <strong style="color:#0F2942;">31 ماي 2026</strong>،
        كما نعلمكم بأنه تم ضبط لجان المناقشة وإدراجها على المنصة.
        </p>
        <p>
        يمكنكم الاطلاع عليها من خلال الولوج إلى المنصة، ثم اختيار <strong>«فضاء الأساتذة»</strong> وإدخال بياناتكم أدناه.
        وستجدون ضمن تبويب <strong>«برنامج المناقشات»</strong> جميع المذكرات التي أنتم أعضاء في لجان مناقشتها،
        مع إمكانية معاينة كل مذكرة مباشرةً عبر المنصة.
        </p>

        <div class="credentials">
            <div style="font-weight:700;color:#0369a1;margin-bottom:10px;">🔐 بيانات الدخول</div>
            <table>
                <tr><td class="label">اسم المستخدم:</td><td class="value">{username}</td></tr>
                <tr><td class="label">كلمة المرور:</td><td class="value">{password}</td></tr>
            </table>
        </div>

        {pending_section}

        <div style="text-align:center;margin-top:20px;">
            <a href="https://memoires2026.streamlit.app" class="btn">🔗 الدخول إلى المنصة</a>
        </div>

        <div style="text-align:center;margin:28px 0 8px;">
            <div style="display:inline-block;background:#f0f9ff;border:2px solid #2F6F7E;border-radius:12px;padding:14px 24px;max-width:90%;">
                <div style="font-weight:900;color:#0F2942;font-size:1.05rem;margin-bottom:4px;">⚠️ ملاحظة هامة</div>
                <div style="color:#1e3a5c;font-size:1rem;line-height:1.7;">
                    التكليفات الرسمية والبرنامج التفصيلي للمناقشات سيتم إرسالها إليكم عبر البريد الإلكتروني المهني.
                </div>
            </div>
        </div>

        <p style="margin-top:24px;">مع فائق التقدير والاحترام،<br>
        <strong>مسؤول الميدان</strong><br>
        <strong>البروفيسور رفاف لخضر</strong><br><br>
        <span style="color:#64748b;font-size:0.9rem;">للاستفسار: يمكنكم الاتصال بمكتب فريق التكوين في الطابق الأرضي بالكلية.</span></p>
    </div>
    <div class="footer">هذا البريد مُرسَل تلقائياً من منصة مذكرات الماستر — لا تردّ على هذه الرسالة</div>
</div>
</body>
</html>'''

        msg = MIMEMultipart("alternative")
        msg["From"] = EMAIL_SENDER
        msg["To"] = email
        msg["Subject"] = "🎓 لجان مناقشة مذكرات الماستر — انطلاق المناقشات 31 ماي 2026"
        msg.attach(MIMEText(body, "html", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_SENDER, EMAIL_PASSWORD)
            server.send_message(msg)
        return True, f"✅ {email}"
    except Exception as e:
        return False, f"❌ {str(e)}"


@st.cache_data(ttl=60)
def load_memo_exceptions():
    """قراءة استثناءات المذكرات من الشيت"""
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=MEMOS_SHEET_ID,
            range="استثناءات_مذكرات!A1:G1000"
        ).execute()
        values = result.get('values', [])
        if not values or len(values) < 2: return pd.DataFrame()
        headers = values[0]
        rows = values[1:]
        padded = [r + [''] * (len(headers) - len(r)) for r in rows]
        return pd.DataFrame(padded, columns=headers)
    except Exception as e:
        logger.error(f"خطأ استثناءات المذكرات: {e}")
        return pd.DataFrame()

@st.cache_data(ttl=60)
def load_prof_exceptions():
    """قراءة استثناءات الأساتذة من الشيت"""
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=MEMOS_SHEET_ID,
            range="استثناءات_أساتذة!A1:L1000"
        ).execute()
        values = result.get('values', [])
        if not values or len(values) < 2: return pd.DataFrame()
        headers = values[0]
        rows = values[1:]
        padded = [r + [''] * (len(headers) - len(r)) for r in rows]
        return pd.DataFrame(padded, columns=headers)
    except Exception as e:
        logger.error(f"خطأ استثناءات الأساتذة: {e}")
        return pd.DataFrame()

def save_memo_exception(row_data):
    """حفظ استثناء مذكرة في الشيت"""
    try:
        values = [[
            str(row_data.get("رقم المذكرة","")),
            str(row_data.get("يوم مثبت","")),
            str(row_data.get("توقيت مثبت","")),
            str(row_data.get("قاعة مثبتة","")),
            str(row_data.get("أقرب تاريخ","")),
            str(row_data.get("أبعد تاريخ","")),
            str(row_data.get("أيام بديلة",""))
        ]]
        sheets_service.spreadsheets().values().append(
            spreadsheetId=MEMOS_SHEET_ID,
            range="استثناءات_مذكرات!A:G",
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body={"values": values}
        ).execute()
        return True
    except Exception as e:
        return False

def save_prof_exception(row_data):
    """حفظ استثناء أستاذ في الشيت"""
    try:
        values = [[
            str(row_data.get("اسم الأستاذ","")),
            str(row_data.get("أيام ممنوعة","")),
            str(row_data.get("أيام مسموحة فقط","")),
            str(row_data.get("لا قبل","")),
            str(row_data.get("لا بعد","")),
            str(row_data.get("يوم واحد","")),
            str(row_data.get("أيام متتالية","")),
            str(row_data.get("مجمّد","")),
            str(row_data.get("عدد مناقشات الفترة الأولى","")),
            str(row_data.get("بداية الفترة الثانية","")),
            str(row_data.get("يقبل 18:00","")),
            str(row_data.get("تجميع الأيام",""))
        ]]
        sheets_service.spreadsheets().values().append(
            spreadsheetId=MEMOS_SHEET_ID,
            range="استثناءات_أساتذة!A:L",
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body={"values": values}
        ).execute()
        return True
    except Exception as e:
        return False

def delete_exception_row(sheet_name, row_idx):
    """حذف استثناء من الشيت"""
    try:
        # نحذف بكتابة صفوف فارغة
        sheets_service.spreadsheets().values().clear(
            spreadsheetId=MEMOS_SHEET_ID,
            range=f"{sheet_name}!A{row_idx}:F{row_idx}"
        ).execute()
        return True
    except: return False

def _norm_date(d):
    """تحويل التاريخ من أي صيغة إلى YYYY-MM-DD"""
    d = str(d).strip()
    if not d or d in ["","nan"]: return ""
    # صيغة DD/MM/YYYY
    import re as _re
    m = _re.match(r"(\d{1,2})/(\d{1,2})/(\d{4})", d)
    if m: return f"{m.group(3)}-{int(m.group(2)):02d}-{int(m.group(1)):02d}"
    # صيغة YYYY-MM-DD
    m2 = _re.match(r"(\d{4})-(\d{2})-(\d{2})", d)
    if m2: return d
    return d

def build_constraints(df_memo_exc, df_prof_exc, slots_per_day):
    """
    بناء قاموس القيود من الاستثناءات:
    - fixed_slots: {memo_id: (day, slot, room)} — مواعيد مثبتة
    - memo_date_limits: {memo_id: (earliest, latest)} — حدود التاريخ
    - prof_banned_days: {prof: set(days)} — أيام ممنوعة
    - prof_not_before: {prof: slot} — لا قبل هذا التوقيت
    - prof_not_after: {prof: slot} — لا بعد هذا التوقيت
    - prof_one_day: set(profs) — أساتذة يوم واحد فقط
    """
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    
    fixed_slots = {}
    memo_date_limits = {}
    prof_banned_days = {}
    prof_not_before = {}
    prof_not_after = {}
    prof_one_day = set()

    # استثناءات المذكرات
    if not df_memo_exc.empty:
        for _, row in df_memo_exc.iterrows():
            mid = str(row.get("رقم المذكرة","")).strip()
            if not mid or mid in ["","nan"]: continue
            
            day_f = _norm_date(str(row.get("يوم مثبت","")).strip())
            slot_f = str(row.get("توقيت مثبت","")).strip()
            room_f = str(row.get("قاعة مثبتة","")).strip()
            early = _norm_date(str(row.get("أقرب تاريخ","")).strip())
            late  = _norm_date(str(row.get("أبعد تاريخ","")).strip())
            
            if day_f and day_f not in ["","nan"] and slot_f and slot_f not in ["","nan"]:

                fixed_slots[mid] = (day_f, slot_f, room_f if room_f not in ["","nan"] else None)
            
            if early not in ["","nan"] or late not in ["","nan"]:
                memo_date_limits[mid] = (
                    early if early not in ["","nan"] else None,
                    late if late not in ["","nan"] else None
                )

    prof_allowed_days = {}   # أيام مسموحة فقط
    prof_consecutive = set() # أيام متتالية

    # استثناءات الأساتذة
    if not df_prof_exc.empty:
        for _, row in df_prof_exc.iterrows():
            prof = str(row.get("اسم الأستاذ","")).strip()
            if not prof or prof in ["","nan"]: continue
            
            banned    = str(row.get("أيام ممنوعة","")).strip()
            allowed   = str(row.get("أيام مسموحة فقط","")).strip()
            not_before= str(row.get("لا قبل","")).strip()
            not_after = str(row.get("لا بعد","")).strip()
            one_day   = str(row.get("يوم واحد","")).strip()
            consec    = str(row.get("أيام متتالية","")).strip()
            
            if banned and banned not in ["","nan"]:
                prof_banned_days.setdefault(prof, set()).update(
                    [_norm_date(d) for d in banned.split(",") if d.strip()])
            
            if allowed and allowed not in ["","nan"]:
                prof_allowed_days[prof] = set(
                    [_norm_date(d) for d in allowed.split(",") if d.strip()])
            
            if not_before and not_before not in ["","nan"]:
                prof_not_before[prof] = not_before
            
            if not_after and not_after not in ["","nan"]:
                prof_not_after[prof] = not_after
            
            if one_day.lower() in ["نعم","yes","1","true"]:
                prof_one_day.add(prof)

            if consec.lower() in ["نعم","yes","1","true"]:
                prof_consecutive.add(prof)
    
    frozen_profs = set()
    prof_phase_split = {}
    memo_alt_days = {}
    profs_accept_18 = set()
    profs_cluster_days = set()  # أساتذة يفضلون تجميع الأيام

    if not df_prof_exc.empty:
        for _, row in df_prof_exc.iterrows():
            prof = str(row.get("اسم الأستاذ","")).strip()
            if not prof or prof in ["","nan"]: continue
            frozen = str(row.get("مجمّد","")).strip()
            if frozen.lower() in ["نعم","yes","1","true"]:
                frozen_profs.add(prof)
            n_first = str(row.get("عدد مناقشات الفترة الأولى","")).strip()
            start_second = str(row.get("بداية الفترة الثانية","")).strip()
            if n_first and n_first not in ["","nan"] and start_second and start_second not in ["","nan"]:
                try:
                    prof_phase_split[prof] = (int(n_first), _norm_date(start_second))
                except: pass
            late = str(row.get("يقبل 18:00","")).strip()
            if late.lower() in ["نعم","yes","1","true"]:
                profs_accept_18.add(prof)
            cluster = str(row.get("تجميع الأيام","")).strip()
            if cluster.lower() in ["نعم","yes","1","true"]:
                profs_cluster_days.add(prof)

    # استثناءات الأيام البديلة للمذكرات
    if not df_memo_exc.empty:
        for _, row in df_memo_exc.iterrows():
            mid = str(row.get("رقم المذكرة","")).strip()
            alt = str(row.get("أيام بديلة","")).strip()
            if mid and alt and alt not in ["","nan"]:
                memo_alt_days[mid] = set([_norm_date(d) for d in alt.split(",") if d.strip()])

    return fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after, prof_one_day, prof_allowed_days, prof_consecutive, frozen_profs, prof_phase_split, memo_alt_days, profs_accept_18, profs_cluster_days


def detect_constraint_conflicts(df_memos, fixed_slots, memo_date_limits, prof_banned_days, prof_not_before, prof_not_after, slots_per_day):
    """
    فحص التعارضات بين الاستثناءات قبل التوليد
    يُرجع قائمة التعارضات
    """
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    conflicts = []
    
    _, memo_members = build_prof_memo_map(df_memos)
    
    for memo_id, (fixed_day, fixed_slot, fixed_room) in fixed_slots.items():
        members = memo_members.get(str(memo_id), set())
        for prof in members:
            # تعارض: يوم مثبت ممنوع للأستاذ
            if fixed_day in prof_banned_days.get(prof, set()):
                conflicts.append(f"⚠️ المذكرة **{memo_id}** مثبتة يوم **{fixed_day}** لكنه ممنوع على الأستاذ **{prof}**")
            
            # تعارض: توقيت مثبت قبل حد الأستاذ
            if prof in prof_not_before:
                if slot_to_idx.get(fixed_slot, 0) < slot_to_idx.get(prof_not_before[prof], 0):
                    conflicts.append(f"⚠️ المذكرة **{memo_id}** مثبتة في **{fixed_slot}** لكن الأستاذ **{prof}** لا يحضر قبل **{prof_not_before[prof]}**")
            
            # تعارض: توقيت مثبت بعد حد الأستاذ
            if prof in prof_not_after:
                if slot_to_idx.get(fixed_slot, 99) > slot_to_idx.get(prof_not_after[prof], 99):
                    conflicts.append(f"⚠️ المذكرة **{memo_id}** مثبتة في **{fixed_slot}** لكن الأستاذ **{prof}** لا يحضر بعد **{prof_not_after[prof]}**")
    
    # تعارض: حدود تاريخ المذكرة متناقضة
    for memo_id, (earliest, latest) in memo_date_limits.items():
        if earliest and latest and earliest > latest:
            conflicts.append(f"⚠️ المذكرة **{memo_id}** لها أقرب تاريخ **{earliest}** أكبر من أبعد تاريخ **{latest}**")

    # تعارض: أيام ممنوعة وأيام مسموحة متناقضة لنفس الأستاذ
    if not df_prof_exc.empty:
        _dc = list(build_constraints(pd.DataFrame(), df_prof_exc, slots_per_day))
        prof_banned_days2 = _dc[2] if len(_dc) > 2 else {}
        prof_allowed_days2 = _dc[6] if len(_dc) > 6 else {}
    else:
        prof_banned_days2 = {}
        prof_allowed_days2 = {}
    for prof, allowed in prof_allowed_days2.items():
        banned = prof_banned_days2.get(prof, set())
        overlap = allowed & banned
        if overlap:
            conflicts.append(f"⚠️ الأستاذ **{prof}**: الأيام {overlap} موجودة في الممنوعة والمسموحة معاً")

    return conflicts


@st.cache_data(ttl=60)
def load_schedule_days():
    """قراءة أيام الجدولة من الشيت"""
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=MEMOS_SHEET_ID,
            range="الأيام!A1:D1000"
        ).execute()
        values = result.get('values', [])
        if not values or len(values) < 2: return [], {}
        headers = values[0]
        days = []
        day_time_limits = {}  # day -> (from_slot, to_slot)
        for row in values[1:]:
            if not row or not row[0].strip(): continue
            date = _norm_date(row[0].strip())  # تحويل لـ YYYY-MM-DD
            note = row[1].strip() if len(row) > 1 else ""
            from_t = row[2].strip() if len(row) > 2 else ""
            to_t = row[3].strip() if len(row) > 3 else ""
            if note:  # أي ملاحظة = يوم محذوف
                continue
            days.append(date)
            if from_t or to_t:
                day_time_limits[date] = (from_t or None, to_t or None)
        return days, day_time_limits
    except Exception as e:
        logger.error(f"خطأ أيام الجدولة: {e}")
        return [], {}

@st.cache_data(ttl=60)
def load_schedule_slots():
    """قراءة التوقيتات من الشيت"""
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=MEMOS_SHEET_ID,
            range="التوقيت!A1:A100"
        ).execute()
        values = result.get('values', [])
        slots = [row[0].strip() for row in values[1:] if row and row[0].strip()]
        return slots
    except Exception as e:
        logger.error(f"خطأ التوقيتات: {e}")
        return []

@st.cache_data(ttl=60)
def load_schedule_rooms():
    """قراءة القاعات من الشيت"""
    try:
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=MEMOS_SHEET_ID,
            range="القاعات!A1:A100"
        ).execute()
        values = result.get('values', [])
        rooms = [row[0].strip() for row in values[1:] if row and row[0].strip()]
        return rooms
    except Exception as e:
        logger.error(f"خطأ القاعات: {e}")
        return []


def validate_schedule(schedule, memo_members, days, slots_per_day):
    """
    تقرير تحقق شامل من صحة الجدول
    يكتشف كل الانتهاكات
    """
    violations = []
    slot_to_idx = {s: i for i, s in enumerate(slots_per_day)}
    
    # بناء برنامج كل أستاذ
    prof_program = {}  # prof -> {day -> [slots]}
    day_slot_room = {}  # (day,slot,room) -> [memo_ids]
    
    for mid, slot_val in schedule.items():
        if not slot_val: continue
        day, slot, room = slot_val
        # تعارض القاعة
        key = (day, slot, room)
        day_slot_room.setdefault(key, []).append(mid)
        # برنامج الأستاذ
        for prof in memo_members.get(mid, set()):
            prof_program.setdefault(prof, {}).setdefault(day, []).append(slot)
    
    # 1. تعارض القاعة
    for (day, slot, room), memos in day_slot_room.items():
        if len(memos) > 1:
            violations.append(f"🔴 تعارض قاعة: {room} في {day} {slot} — مذكرات: {', '.join(memos)}")
    
    for prof, days_dict in prof_program.items():
        day_counts = {d: len(slots) for d, slots in days_dict.items()}
        
        # 2. تجاوز 3 مذكرات/يوم
        for day, count in day_counts.items():
            if count > 3:
                violations.append(f"🔴 {prof}: {count} مذكرات في {day} (الحد 3)")
        
        # 3. تعارض في نفس التوقيت
        for day, slots in days_dict.items():
            if len(slots) != len(set(slots)):
                violations.append(f"🔴 {prof}: تعارض توقيتات في {day}")
        
        # 4. أكثر من يوم منعزل — فقط إذا المجموع >= 3
        total_sessions = sum(day_counts.values())
        lonely_days = [d for d, cnt in day_counts.items() if cnt == 1]
        if total_sessions >= 3 and len(lonely_days) > 2:
            violations.append(f"🔴 {prof}: {len(lonely_days)} أيام منعزلة ({', '.join(sorted(lonely_days))}) — الحد المسموح يومان")
        elif total_sessions >= 3 and len(lonely_days) == 2:
            violations.append(f"🟡 {prof}: 2 أيام منعزلة ({', '.join(sorted(lonely_days))})")
        if False and len(lonely_days) > 1 and total_sessions > 2 and len(lonely_days) < total_sessions:
            violations.append(f"🟡 {prof}: {len(lonely_days)} أيام منعزلة ({', '.join(sorted(lonely_days))})")
        
        # 5. أكثر من 3 أيام متتالية
        sorted_days = sorted(days_dict.keys())
        consec = 1
        max_consec = 1
        for i in range(1, len(sorted_days)):
            d1, d2 = sorted_days[i-1], sorted_days[i]
            # تحقق التتالي (يومان متجاوران)
            try:
                from datetime import datetime, timedelta
                dt1 = datetime.strptime(d1, "%Y-%m-%d")
                dt2 = datetime.strptime(d2, "%Y-%m-%d")
                diff = (dt2 - dt1).days
                # تجاهل عطلة الجمعة والسبت
                if diff <= 3:
                    consec += 1
                    max_consec = max(max_consec, consec)
                else:
                    consec = 1
            except: pass
        if max_consec > 3:
            violations.append(f"🟡 {prof}: {max_consec} أيام متتالية")
    
    return violations



import io, zipfile, re as _re_mahdar, tempfile as _tempfile

def generate_qr_png(url):
    """QR Code معطّل"""
    return b""
def generate_mahdar(memo_data, seq_num, template_bytes):
    """توليد محضر مناقشة من القالب بـ placeholders"""
    import io as _io
    from docx import Document as _Doc
    from docx.oxml import OxmlElement as _Elem
    from docx.oxml.ns import qn as _qn
    from docx.shared import Pt as _Pt, Cm as _Cm

    memo_num     = str(memo_data.get("رقم المذكرة","")).strip()
    title        = str(memo_data.get("عنوان المذكرة","")).strip()
    specialty    = str(memo_data.get("التخصص","")).strip()
    dept         = str(memo_data.get("القسم","")).strip()
    def_date     = str(memo_data.get("تاريخ المناقشة","")).strip()
    student_name = str(memo_data.get("الطالب","")).strip()
    student_id   = str(memo_data.get("رقم ملف الطالب","")).strip()
    student2_name= str(memo_data.get("الطالب2","")).strip()
    student2_id  = str(memo_data.get("رقم ملف الطالب2","")).strip()
    drive_link   = str(memo_data.get("رابط الملف","")).strip()
    pres    = str(memo_data.get("الرئيس","")).strip()
    sup     = str(memo_data.get("الأستاذ","")).strip()
    ex1     = str(memo_data.get("المناقش1","")).strip()
    ex2     = str(memo_data.get("المناقش2","")).strip()
    rank_p  = str(memo_data.get("رتبة_الرئيس","")).strip()
    rank_s  = str(memo_data.get("رتبة_المشرف","")).strip()
    rank_e1 = str(memo_data.get("رتبة_المناقش1","")).strip()
    rank_e2 = str(memo_data.get("رتبة_المناقش2","")).strip()

    has_student2 = bool(student2_name and student2_name not in ["","nan","-"])

    doc = _Doc(_io.BytesIO(template_bytes))
    seq_str = str(seq_num).zfill(3)

    # ── Placeholders ──
    # إضافة علامة RTL لمنع عكس الأقواس
    seq_str_rtl = "‏" + seq_str + "‏"
    replacements = {
        "{{SEQ}}":         seq_str_rtl,
        "{{DATE}}":        def_date,
        "{{MEMO_NUM}}":    memo_num,
        "{{TITLE}}":       title,
        "{{SPECIALTY}}":   specialty,
        "{{DEPT}}":        dept,
        "{{Dept}}":        dept,
        "{{Dep}}":         dept,
        "{{STUDENT1}}":    student_name,
        # طالب واحد → لا رقم ملف، طالبان → رقم ملف
        "{{STUDENT1_ID}}": student_id,  # يظهر دائماً
        "{{STUDENT2}}":    student2_name if has_student2 else "",
        "{{STUDENT2_ID}}": student2_id if has_student2 else "",
        "{{STUDENTS_LABEL}}": memo_data.get("STUDENTS_LABEL", "للطالبين" if has_student2 else "للطالب(ة)"),
        "...../......./......": def_date,
    }

    def replace_in_para(p, mapping):
        """يدمج كل runs الفقرة ثم يستبدل مع الحفاظ على التنسيق"""
        if not p.runs: return
        full = "".join(r.text for r in p.runs)
        changed = any(old in full for old in mapping)
        if changed:
            for old, new in mapping.items():
                full = full.replace(old, new)
            target_run = next((r for r in p.runs if r.bold), None) or p.runs[0]
            for r in p.runs: r.text = ""
            target_run.text = full

    def remove_para(p):
        """حذف فقرة كاملة"""
        p._p.getparent().remove(p._p)

    # حذف فقرة الطالب الثاني قبل الاستبدال
    paras_to_remove = []
    for p in doc.paragraphs:
        raw = "".join(r.text for r in p.runs)
        if not has_student2 and ("{{STUDENT2}}" in raw or "{{STUDENT2_ID}}" in raw):
            paras_to_remove.append(p)
    for p in paras_to_remove:
        p._p.getparent().remove(p._p)

    for p in doc.paragraphs:
        replace_in_para(p, replacements)

    # ── Tab Stop ثابت لمحاذاة رقم الملف ──
    from docx.oxml import OxmlElement as _OE2
    from docx.oxml.ns import qn as _qn2
    from docx.shared import Cm as _Cm2
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if "(رقم " in full and "- " in full:
            pPr = p._p.get_or_add_pPr()
            # احذف tab stops الموجودة
            old_tabs = pPr.find(_qn2("w:tabs"))
            if old_tabs is not None: pPr.remove(old_tabs)
            # أضف tab stop ثابت عند 9 سم
            tabs = _OE2("w:tabs")
            tab = _OE2("w:tab")
            tab.set(_qn2("w:val"), "left")
            tab.set(_qn2("w:pos"), str(int(_Cm2(6).pt * 20)))
            tabs.append(tab)
            pPr.append(tabs)

    # ── إزالة {{QR}} ──
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if "{{QR}}" in full:
            for r in p.runs: r.text = ""
            break

    # ── جدول اللجنة ──
    def fill_cell(cell, text, size=14, bold=True):
        for p in cell.paragraphs:
            for r in p.runs: r.text = ""
        p = cell.paragraphs[0]
        pPr = p._p.get_or_add_pPr()
        for tag in ["w:bidi","w:jc"]:
            ex = pPr.find(_qn(tag))
            if ex is not None: pPr.remove(ex)
        bidi = _Elem("w:bidi"); pPr.append(bidi)
        jc = _Elem("w:jc"); jc.set(_qn("w:val"), "center"); pPr.append(jc)
        run = p.add_run(text)
        rPr = run._r.get_or_add_rPr()
        rF = _Elem("w:rFonts")
        for a in ["w:ascii","w:hAnsi","w:cs","w:eastAsia"]:
            rF.set(_qn(a), "Sakkal Majalla")
        ex = rPr.find(_qn("w:rFonts"))
        if ex is not None: rPr.remove(ex)
        rPr.insert(0, rF)
        run.font.size = _Pt(size)
        run.font.bold = bold

    t = doc.tables[0]
    members = [
        ("01", pres,  rank_p,  "جامعة برج بوعريريج", "رئيساً"),
        ("02", sup,   rank_s,  "جامعة برج بوعريريج", "مشرفاً ومقرراً"),
        ("03", ex1,   rank_e1, "جامعة برج بوعريريج", "ممتحناً"),
        ("04", ex2,   rank_e2, "جامعة برج بوعريريج", "ممتحناً"),
    ]
    for ri, (num, name, rank, univ, role) in enumerate(members, 1):
        if ri >= len(t.rows): break
        row = t.rows[ri]
        fill_cell(row.cells[0], num)
        fill_cell(row.cells[1], name, size=13)
        fill_cell(row.cells[2], rank, size=13)
        fill_cell(row.cells[3], univ, size=13)
        fill_cell(row.cells[4], role, size=13)

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def lookup_student(username):
    if df_students.empty: return None
    s = df_students[df_students["اسم المستخدم"].astype(str).apply(normalize_text)==normalize_text(username)]
    return s.iloc[0].to_dict() if not s.empty else None

def restore_session_from_url():
    if st.session_state.get('logged_in', False): return
    qp = st.query_params
    if 'ut' not in qp or 'un' not in qp: return
    user_type = qp['ut'] if isinstance(qp['ut'],str) else (qp['ut'][0] if qp['ut'] else "")
    username_enc = qp['un'] if isinstance(qp['un'],str) else (qp['un'][0] if qp['un'] else "")
    username = decode_str(username_enc)
    if not username: return
    if user_type=='student':
        s_data = lookup_student(username)
        if s_data:
            ph_ok,_ = is_phone_valid(str(s_data.get('الهاتف','')))
            nin_ok,_ = is_nin_valid(normalize_text(s_data.get('NIN','')))
            if ph_ok and nin_ok:
                st.session_state.user_type='student'; st.session_state.student1=s_data
                note_num = normalize_text(s_data.get('رقم المذكرة',''))
                st.session_state.mode = "view" if note_num else "register"
                st.session_state.logged_in = True
                if note_num:
                    _df_m_r = load_memos()
                    mr = _df_m_r[_df_m_r["رقم المذكرة"].astype(str).apply(normalize_text)==note_num]
                    if not mr.empty:
                        s2_name = str(mr.iloc[0].get("الطالب الثاني","")).strip()
                        if s2_name and s2_name!="--":
                            _df_st_r = load_students()
                            s2 = load_student2_for_memo(mr.iloc[0], normalize_text(s_data.get('رقم التسجيل','')), _df_st_r)
                            if s2: st.session_state.student2=s2
    elif user_type=='professor':
        _df_pm_r = load_prof_memos()
        p = _df_pm_r[_df_pm_r["إسم المستخدم"].astype(str).apply(normalize_text)==normalize_text(username)]
        if not p.empty: st.session_state.professor=p.iloc[0].to_dict(); st.session_state.logged_in=True
    elif user_type=='admin':
        if username in ADMIN_CREDENTIALS: st.session_state.admin_user=username; st.session_state.logged_in=True
    if user_type: st.session_state.user_type=user_type

restore_session_from_url()

required_state = {
    'user_type':None,'logged_in':False,'student1':None,'student2':None,
    'professor':None,'admin_user':None,'memo_type':"فردية",
    'mode':"register",'note_number':"",'prof_password':"",
    'show_confirmation':False,'selected_memo_id':None,
    'admin_edit_student_user':None,'s2_phone_input':"",'s2_nin_input':"",
    'profile_incomplete':False,'profile_user_temp':None,'profile_error_msg':None,
    'admin_mode':None,'wizard_step':1,'generated_schedule':None,
    'prof_action':None
}
for k,v in required_state.items():
    if k not in st.session_state: st.session_state[k]=v

def logout():
    st.query_params.clear()
    for k in list(st.session_state.keys()): del st.session_state[k]
    for k,v in required_state.items(): st.session_state[k]=v
    st.rerun()


# ================================================================
# الصفحة الرئيسية
# ================================================================
if st.session_state.user_type is None:
    render_countdown_banner()
    st.markdown("<p style='text-align:center;color:#ffffff;font-size:1.05rem;'>جامعة محمد البشير الإبراهيمي — كلية الحقوق والعلوم السياسية</p>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align:center;font-size:2rem;margin-bottom:2rem;'>📘 منصة مذكرات الماستر</h1>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    for col, icon, title, desc, utype, btn_label in [
        (col1,"🎓","فضاء الطلبة","تسجيل المذكرة، إيداع النسخة النهائية، متابعة الملف","student","دخول الطلبة"),
        (col2,"📚","فضاء الأساتذة","متابعة المذكرات، مراجعة الإيداعات، تحديث التقدم","professor","دخول الأساتذة"),
        (col3,"⚙️","فضاء الإدارة","لوحة التحكم الكاملة، برنامج المناقشات","admin","دخول الإدارة"),
    ]:
        with col:
            st.markdown(f"""<div class="card" style="text-align:center;min-height:190px;display:flex;flex-direction:column;align-items:center;justify-content:center;"><div style="font-size:2.8rem;margin-bottom:10px;">{icon}</div><h3 style="margin:0 0 7px;">{title}</h3><p style="color:#ffffff!important;font-size:0.83rem;">{desc}</p></div>""", unsafe_allow_html=True)
            if st.button(btn_label, key=f"btn_{utype}", use_container_width=True):
                st.session_state.user_type=utype; st.rerun()

# ================================================================
# فضاء الطلبة
# ================================================================
elif st.session_state.user_type == "student":

    if st.session_state.get('profile_incomplete', False):
        render_countdown_banner()
        st.markdown("<h2>⚠️ استكمال الملف الشخصي</h2>", unsafe_allow_html=True)
        temp = st.session_state.profile_user_temp
        with st.form("complete_profile"):
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            dp = str(temp.get('الهاتف','')).strip()
            if dp in ['0','0.0','nan','-','']: dp=""
            dn = normalize_text(temp.get('NIN',''))
            if dn in ['0','nan','-','']: dn=""
            new_phone = st.text_input("📞 رقم الهاتف (10 أرقام)", value=dp)
            new_nin = st.text_input("🆔 رقم التعريف الوطني", value=dn)
            if st.form_submit_button("💾 حفظ والمتابعة", type="primary", use_container_width=True):
                ph_ok,_=is_phone_valid(new_phone); nin_ok,_=is_nin_valid(new_nin)
                if not ph_ok: st.error("❌ رقم الهاتف غير صالح")
                elif not nin_ok: st.error("❌ رقم التعريف غير صحيح")
                else:
                    username = normalize_text(temp.get('اسم المستخدم',''))
                    ok, msg = update_student_profile(username, new_phone, new_nin)
                    if ok:
                        st.session_state.profile_incomplete=False; st.session_state.logged_in=True
                        df_up = load_students()
                        st.session_state.student1 = df_up[df_up["اسم المستخدم"].astype(str).apply(normalize_text)==username].iloc[0].to_dict()
                        note_num = normalize_text(st.session_state.student1.get('رقم المذكرة',''))
                        st.session_state.mode = "view" if note_num else "register"
                        st.rerun()
                    else: st.error(msg)
            st.markdown("</div>", unsafe_allow_html=True)

    elif not st.session_state.logged_in:
        render_countdown_banner()
        c1,c2=st.columns([4,1])
        with c2:
            if st.button("رجوع"): st.session_state.user_type=None; st.rerun()
        st.markdown("<h2>🎓 فضاء الطلبة</h2>", unsafe_allow_html=True)
        with st.form("student_login"):
            st.write("### 🔐 تسجيل الدخول")
            username1 = st.text_input("اسم المستخدم")
            password1 = st.text_input("كلمة السر", type="password")
            if st.form_submit_button("دخول"):
                df_students = load_students()
                valid, result = verify_student(username1, password1, df_students)
                if not valid: st.error(result)
                else:
                    ph_ok,_=is_phone_valid(str(result.get('الهاتف','')))
                    nin_ok,_=is_nin_valid(normalize_text(result.get('NIN','')))
                    if ph_ok and nin_ok:
                        st.session_state.student1=result
                        note_num = normalize_text(result.get('رقم المذكرة',''))
                        if note_num:
                            st.session_state.mode="view"
                            df_memos = load_memos()
                            mr = df_memos[df_memos["رقم المذكرة"].astype(str).apply(normalize_text)==note_num]
                            if not mr.empty:
                                s2_name = str(mr.iloc[0].get("الطالب الثاني","")).strip()
                                if s2_name and s2_name!="--":
                                    s2 = load_student2_for_memo(mr.iloc[0], normalize_text(result.get('رقم التسجيل','')), df_students)
                                    if s2: st.session_state.student2=s2
                        else: st.session_state.mode="register"
                        st.session_state.logged_in=True
                        st.query_params['ut']='student'
                        st.query_params['un']=encode_str(normalize_text(result.get('اسم المستخدم','')))
                        st.rerun()
                    else:
                        st.session_state.profile_user_temp=result
                        st.session_state.profile_incomplete=True
                        st.rerun()

    else:
        render_countdown_banner()
        s1 = st.session_state.student1; s2 = st.session_state.student2

        if st.session_state.mode == "register":
            c1,c2=st.columns([4,1])
            with c2:
                if st.button("خروج"): logout()
            st.markdown("<h2>📝 تسجيل مذكرة ماستر</h2>", unsafe_allow_html=True)
            registration_type = st.radio("نوع المذكرة:", ["فردية","ثنائية"], horizontal=True)
            username2=password2=None; student2_obj=None; s2_missing_info=False
            s2_new_phone_val=""; s2_new_nin_val=""
            if registration_type=="ثنائية":
                st.markdown("### 👥 بيانات الطالب الثاني")
                username2=st.text_input("اسم المستخدم الثاني")
                password2=st.text_input("كلمة السر الثاني", type="password")
                if username2 and password2:
                    v2,r2=verify_student(username2,password2,df_students)
                    if v2:
                        student2_obj=r2
                        ph_ok,_=is_phone_valid(str(r2.get('الهاتف','')))
                        nin_ok,_=is_nin_valid(normalize_text(r2.get('NIN','')))
                        if not ph_ok or not nin_ok:
                            s2_missing_info=True
                            st.warning("⚠️ بيانات الطالب الثاني ناقصة.")
                            s2_new_phone_val=st.text_input("📞 هاتف الطالب الثاني", key="s2_ph")
                            s2_new_nin_val=st.text_input("🆔 NIN الطالب الثاني", key="s2_nin")
                        else: st.success(f"✅ {r2.get('لقب','')} {r2.get('إسم','')} — بياناته مكتملة")
                    else: st.error(r2)
            st.markdown("### 🔍 المذكرات المتاحة")
            student_specialty = str(s1.get("التخصص","")).strip()
            available_memos_df = df_memos[(df_memos["تم التسجيل"].astype(str).str.strip()!="نعم") & (df_memos["التخصص"].astype(str).str.strip()==student_specialty)]
            def clean_text(val):
                v=str(val).strip(); return '' if v in ['','nan','None','NaN','-','0','0.0'] else v
            prof_counts = df_prof_memos.groupby("الأستاذ")["رقم المذكرة"].apply(lambda x: sum(1 for v in x if clean_text(v)!='')).to_dict()
            available_profs=[]
            if not available_memos_df.empty:
                for p in available_memos_df["الأستاذ"].unique():
                    if prof_counts.get(str(p).strip(),0)<4: available_profs.append(str(p).strip())
            available_profs=sorted(set(available_profs))
            if available_profs:
                sel_prof=st.selectbox("الأستاذ المشرف:", [""]+available_profs)
                if sel_prof:
                    pm=available_memos_df[available_memos_df["الأستاذ"].astype(str).str.strip()==sel_prof.strip()]
                    if not pm.empty:
                        st.success(f"✅ {len(pm)} عنوان متاح:")
                        for _,row in pm.iterrows():
                            st.markdown(f"""<div style="background:rgba(47,111,126,0.1);border:1px solid #2F6F7E;padding:10px;border-radius:8px;margin-bottom:5px;"><strong style="color:#FFD700;">{row['رقم المذكرة']}</strong> — <span style="color:#E2E8F0;">{row['عنوان المذكرة']}</span></div>""", unsafe_allow_html=True)
            else:
                st.warning("🔒 لا توجد أماكن شاغرة في تخصصك حالياً.")
            st.markdown("### ✍️ تسجيل المذكرة")
            c1,c2=st.columns([3,1])
            with c1: st.session_state.note_number=st.text_input("رقم المذكرة", value=st.session_state.note_number)
            with c2: st.session_state.prof_password=st.text_input("كلمة سر المشرف", type="password")
            if not st.session_state.show_confirmation:
                if st.button("المتابعة للتأكيد"):
                    if not st.session_state.note_number or not st.session_state.prof_password:
                        st.error("⚠️ يرجى إدخال بيانات المذكرة")
                    elif registration_type=="ثنائية" and not student2_obj:
                        st.error("❌ يرجى إدخال بيانات الطالب الثاني")
                    elif s2_missing_info and (not s2_new_phone_val or not s2_new_nin_val):
                        st.error("❌ يرجى إدخال بيانات الطالب الثاني الناقصة")
                    else:
                        s1_perm=str(s1.get('التسجيل','')).strip()
                        if registration_type=="ثنائية":
                            s2_temp=student2_obj; s2_perm=str(s2_temp.get('التسجيل','')).strip()
                            if s1_perm!='1' and s2_perm!='1': st.error("⛔ لم يتم السماح لك بالتسجيل."); st.stop()
                        else:
                            fardiya=str(s1.get('فردية','')).strip()
                            if fardiya not in ["1","نعم"]: st.error("❌ لا يمكنك تسجيل مذكرة فردية"); st.stop()
                            if s1_perm!='1': st.error("⛔ لم يتم السماح لك."); st.stop()
                        st.session_state.show_confirmation=True
                        st.session_state.s2_phone_input=s2_new_phone_val if s2_missing_info else ""
                        st.session_state.s2_nin_input=s2_new_nin_val if s2_missing_info else ""
                        st.rerun()
            else:
                st.warning(f"⚠️ تأكيد التسجيل — المذكرة رقم: {st.session_state.note_number}")
                c1,c2=st.columns(2)
                with c1:
                    if st.button("تأكيد نهائي", type="primary"):
                        valid,_,err=verify_professor_password(st.session_state.note_number, st.session_state.prof_password, df_memos, df_prof_memos)
                        if not valid: st.error(err); st.session_state.show_confirmation=False
                        else:
                            with st.spinner("⏳ جاري التسجيل..."):
                                s2_pass=student2_obj if registration_type=="ثنائية" else None
                                ok,msg=update_registration(st.session_state.note_number,s1,s2_pass,st.session_state.s2_phone_input,st.session_state.s2_nin_input)
                            if ok:
                                st.success(msg); st.balloons(); clear_cache_and_reload()
                                st.session_state.mode="view"; st.session_state.show_confirmation=False
                                time_module.sleep(2); st.rerun()
                            else: st.error(msg); st.session_state.show_confirmation=False
                with c2:
                    if st.button("إلغاء"): st.session_state.show_confirmation=False; st.rerun()

        elif st.session_state.mode == "view":
            c1,c2=st.columns([4,1])
            with c2:
                if st.button("خروج"): logout()
            note_num = normalize_text(s1.get('رقم المذكرة',''))
            df_memos_fresh = load_memos()
            memo_info = df_memos_fresh[df_memos_fresh["رقم المذكرة"].astype(str).apply(normalize_text)==note_num].iloc[0]
            deposit_status = str(memo_info.get("حالة الإيداع","")).strip()
            deposit_link   = str(memo_info.get("رابط الملف","")).strip()
            deposit_date   = str(memo_info.get("تاريخ إيداع المذكرة","")).strip()
            def_date_m     = str(memo_info.get("تاريخ المناقشة","")).strip()
            def_time_m     = str(memo_info.get("توقيت المناقشة","")).strip()
            def_room_m     = str(memo_info.get("القاعة","")).strip()
            is_published   = str(memo_info.get("AF","")).strip()=="نعم" if "AF" in memo_info.index else False
            prof_name_m    = str(memo_info.get("الأستاذ","")).strip()

            # ── تحقق من حالة الإيداع الاستثنائية ──
            ah_val = str(memo_info.get("مفقودة","")).strip() or str(memo_info.get("AH","")).strip()
            is_missing = ah_val == "1"      # مذكرة مفقودة → رسالة + إعادة رفع
            is_extended = ah_val == "2"     # تمديد استثنائي → إيداع صامت بدون رسالة

            # ── موعد المناقشة في الصدارة ──
            # عمود نشر البرنامج — يتحكم في ظهور الموعد للطالب
            _pub_status = str(memo_info.get("نشر البرنامج","")).strip()
            _show_schedule = _pub_status.lower() in ["نعم","yes","1","true"]
            # إذا تمت المناقشة لا نعرض بطاقة الموعد
            _hal_done = str(memo_info.get("الحالة","") if memo_info is not None else "").strip()
            if _hal_done == "تمت المناقشة":
                _show_schedule = False

            if _show_schedule and def_date_m and def_date_m not in ["","nan"]:
                st.markdown(f'''<div style="background:linear-gradient(135deg,#0a1f12,#0f2d1a);
                    border:2px solid rgba(16,185,129,0.5);border-radius:16px;
                    padding:20px 24px;margin-bottom:18px;text-align:center;">
                    <div style="font-size:0.85rem;color:#6EE7B7;margin-bottom:6px;font-weight:700;">📅 موعد مناقشتك</div>
                    <div style="font-size:1.8rem;font-weight:900;color:#10B981;">{def_date_m}</div>
                    <div style="display:flex;justify-content:center;gap:28px;margin-top:12px;flex-wrap:wrap;">
                        <div><span style="color:#6EE7B7;font-size:0.82rem;">🕐 التوقيت</span><br>
                             <span style="color:#ffffff;font-weight:700;font-size:1.1rem;">{def_time_m if def_time_m and def_time_m not in ["","nan"] else "—"}</span></div>
                        <div><span style="color:#6EE7B7;font-size:0.82rem;">🏛️ القاعة</span><br>
                             <span style="color:#ffffff;font-weight:700;font-size:1.1rem;">{def_room_m if def_room_m and def_room_m not in ["","nan"] else "—"}</span></div>
                    </div>
                </div>''', unsafe_allow_html=True)
            elif not _show_schedule and def_date_m and def_date_m not in ["","nan"]:
                st.markdown('''<div style="background:linear-gradient(135deg,#0f1f2d,#1a2f3d);
                    border:2px solid rgba(245,158,11,0.4);border-radius:16px;
                    padding:20px 24px;margin-bottom:18px;text-align:center;">
                    <div style="font-size:1.4rem;margin-bottom:10px;">📅</div>
                    <div style="font-size:1rem;font-weight:700;color:#F59E0B;margin-bottom:8px;">المناقشات ستنطلق ابتداءً من 31 ماي 2026</div>
                    <div style="font-size:0.88rem;color:#CBD5E1;line-height:1.7;">
                        يُرجى متابعة فضائك على المنصة باستمرار للاطلاع على موعد مناقشتك فور تحديده.
                    </div>
                </div>''', unsafe_allow_html=True)
            else:
                st.markdown('''<div style="background:linear-gradient(135deg,#0f1f2d,#1a2f3d);
                    border:2px solid rgba(245,158,11,0.4);border-radius:16px;
                    padding:20px 24px;margin-bottom:18px;text-align:center;">
                    <div style="font-size:1.4rem;margin-bottom:10px;">📅</div>
                    <div style="font-size:1rem;font-weight:700;color:#F59E0B;margin-bottom:8px;">المناقشات ستنطلق ابتداءً من 31 ماي 2026</div>
                    <div style="font-size:0.88rem;color:#CBD5E1;line-height:1.7;">
                        يُرجى متابعة فضائك على المنصة باستمرار للاطلاع على موعد مناقشتك فور تحديده.
                    </div>
                </div>''', unsafe_allow_html=True)


            # ── ملاحظات المشرف (تظهر دائماً) ──
          # العمود Z يظهر فقط عند الرفض في الإشعارات
            if is_missing:
                st.markdown("""
                <div style="background:linear-gradient(135deg,#1a0a0a,#3d0f0f);border:4px solid #EF4444;
                            border-radius:24px;padding:48px 36px;margin-bottom:28px;text-align:center;
                            box-shadow:0 0 60px rgba(239,68,68,0.35);">
                    <div style="font-size:5rem;margin-bottom:20px;">🚨</div>
                    <div style="font-size:2.2rem;font-weight:900;color:#EF4444;margin-bottom:18px;line-height:1.4;">
                        تعذّر الوصول إلى ملف مذكرتك
                    </div>
                    <div style="background:rgba(239,68,68,0.12);border-radius:14px;padding:22px;margin-bottom:20px;">
                        <div style="font-size:1.25rem;color:#FCA5A5;line-height:2.2;">
                            نتيجة لخلل تقني طرأ على المنظومة،<br>
                            <strong style="color:#ffffff;font-size:1.35rem;">لم يعد بالإمكان الوصول إلى الملف الذي رفعته سابقاً.</strong>
                        </div>
                    </div>
                    <div style="background:rgba(255,215,0,0.12);border:3px solid rgba(255,215,0,0.5);
                                border-radius:14px;padding:22px;">
                        <div style="font-size:1.5rem;font-weight:900;color:#FFD700;line-height:1.8;">
                            ⬇️ يُرجى إعادة رفع مذكرتك فوراً<br>من هذه الصفحة
                        </div>
                    </div>
                </div>""", unsafe_allow_html=True)

                uploaded_recovery = st.file_uploader("📁 أعد رفع ملف المذكرة (PDF فقط)", type=["pdf"], key="upload_recovery_pdf")
                if uploaded_recovery:
                    rec_bytes = uploaded_recovery.read()
                    size_mb = len(rec_bytes) / (1024*1024)
                    uploaded_recovery.seek(0)
                    st.info(f"📊 حجم الملف: {size_mb:.1f} MB")
                    if size_mb > 20:
                        st.error("❌ الحجم يتجاوز 20 MB")
                    elif rec_bytes[:4] != b'%PDF':
                        st.error("❌ الملف ليس PDF حقيقياً")
                    else:
                        if st.button("📤 إعادة رفع المذكرة", type="primary", use_container_width=True, key="btn_recovery_upload"):
                            with st.spinner("⏳ جاري رفع الملف..."):
                                ok, link, msg = upload_memo_to_drive(rec_bytes, note_num, memo_info["عنوان المذكرة"])
                                if ok:
                                    s, m = save_memo_deposit(note_num, link)
                                    if s:
                                        clear_missing_flag(note_num)
                                        # إيميل للإدارة فقط
                                        s1_ln, s1_fn = get_student_name_display(st.session_state.student1)
                                        s1_display = f"{s1_ln} {s1_fn}".strip()
                                        s2_display = ""
                                        s2_obj = load_student2_for_memo(memo_info, normalize_text(st.session_state.student1.get("رقم التسجيل","")), load_students())
                                        if s2_obj:
                                            s2l, s2f = get_student_name_display(s2_obj)
                                            s2_display = f"{s2l} {s2f}".strip()
                                        send_recovery_email_to_admin(note_num, memo_info["عنوان المذكرة"], s1_display, s2_display)
                                        st.success("✅ تمت إعادة رفع مذكرتك بنجاح!")
                                        st.balloons()
                                        clear_cache_and_reload()
                                        time_module.sleep(2)
                                        st.rerun()
                                    else:
                                        st.error(m)
                                else:
                                    st.error(msg)

            elif deposit_status == "مرفوضة":
                rejection_raw = str(memo_info.get("توقيع المشرف","")).strip()
                reason_display = rejection_raw.split("السبب:")[-1].strip() if "السبب:" in rejection_raw else "يرجى مراجعة المشرف."
                st.markdown(f"""<div class="notif-card notif-card-rejected"><div class="notif-icon">🔴</div><div><div class="notif-title notif-title-rejected">المذكرة بحاجة لمراجعة</div><div class="notif-desc"><strong>ملاحظات المشرف:</strong><br>{reason_display}</div></div></div>""", unsafe_allow_html=True)

            if deposit_status in ["", "nan", "مرفوضة"] or not deposit_status:
                deadline_passed = datetime.now() > DEPOSIT_DEADLINE
                if deadline_passed and not is_extended:
                    st.markdown("""
                    <div style="background:linear-gradient(135deg,#1a0a0a,#2d0f0f);border:2px solid rgba(239,68,68,0.5);
                                border-radius:18px;padding:28px 32px;margin-bottom:18px;text-align:center;">
                        <div style="font-size:2.8rem;margin-bottom:12px;">🔒</div>
                        <div style="font-size:1.2rem;font-weight:900;color:#EF4444;margin-bottom:8px;">
                            انتهى أجل إيداع المذكرات
                        </div>
                        <div style="font-size:0.9rem;color:#E2E8F0;line-height:1.7;">
                            انتهى الأجل الرسمي لإيداع المذكرات في
                            <strong style="color:#FFD700;">23 ماي 2026 الساعة 23:59</strong>
                            <br>لم يعد بإمكانك إيداع أي ملف عبر المنصة.
                            <br><br>
                            <span style="color:#F59E0B;">للاستفسار أو في حالة الضرورة القصوى، راجع الإدارة مباشرة.</span>
                        </div>
                    </div>""", unsafe_allow_html=True)
                else:
                    days_left = get_days_remaining()
                    st.markdown(f"""<div class="deposit-hero"><span class="deposit-hero-icon">📤</span><div class="deposit-hero-title">يمكنك الآن إيداع مذكرتك النهائية</div><div class="deposit-hero-sub" style="color:#E2E8F0!important;">آخر أجل: <strong style="color:#FFD700;">23 ماي 2026</strong> — تبقى <strong style="color:{'#EF4444' if days_left<=7 else '#FFD700'};">{days_left} يوم</strong><br>ارفع نسخة PDF من مذكرتك. سيراجعها المشرف ويوافق أو يرسل ملاحظاته.</div></div>""", unsafe_allow_html=True)
                    uploaded_pdf = st.file_uploader("📁 اختر ملف المذكرة (PDF فقط)", type=["pdf"], key="upload_pdf")
                    if uploaded_pdf:
                        pdf_bytes = uploaded_pdf.read(); size_mb = len(pdf_bytes)/(1024*1024); uploaded_pdf.seek(0)
                        st.info(f"📊 حجم الملف: {size_mb:.1f} MB")
                        if size_mb > 20: st.error("❌ الحجم يتجاوز 20 MB")
                        elif pdf_bytes[:4] != b'%PDF': st.error("❌ الملف ليس PDF حقيقياً — تأكد من الملف")
                        else:
                            if st.button("📤 إيداع المذكرة الآن", type="primary", use_container_width=True):
                                with st.spinner("⏳ جاري رفع الملف..."):
                                    ok, link, msg = upload_memo_to_drive(pdf_bytes, note_num, memo_info['عنوان المذكرة'])
                                    if ok:
                                        s, m = save_memo_deposit(note_num, link)
                                        if s:
                                            s1_ln,s1_fn = get_student_name_display(st.session_state.student1)
                                            s1_display = f"{s1_ln} {s1_fn}".strip()
                                            s2_display = ""
                                            s2_obj_dep = load_student2_for_memo(memo_info, normalize_text(st.session_state.student1.get('رقم التسجيل','')), load_students())
                                            if s2_obj_dep:
                                                s2l,s2f = get_student_name_display(s2_obj_dep); s2_display=f"{s2l} {s2f}".strip()
                                            email_ok, email_msg = send_deposit_email_to_professor(prof_name_m, note_num, memo_info['عنوان المذكرة'], s1_display, s2_display)
                                            st.success("✅ تم إيداع مذكرتك! سيراجعها المشرف قريباً.")
                                            if email_ok: st.info("📧 تم إرسال إشعار للمشرف والإدارة.")
                                            else: st.warning(f"⚠️ فشل إرسال الإيميل للمشرف: {email_msg}")
                                            st.balloons(); clear_cache_and_reload(); time_module.sleep(2); st.rerun()
                                        else: st.error(m)
                                    else: st.error(msg)
            elif deposit_status == "مودعة" and not is_missing and not is_extended:
                st.markdown("""<div class="notif-card notif-card-waiting"><div class="notif-icon">🟡</div><div><div class="notif-title notif-title-waiting">مذكرتك مودعة — في انتظار مراجعة المشرف</div><div class="notif-desc">تم استلام ملفك. سيراجعه المشرف ويوافق أو يرسل ملاحظاته. ستتلقى إشعاراً فور اتخاذ القرار.</div></div></div>""", unsafe_allow_html=True)
                if deposit_date and deposit_date not in ["","nan"]: st.caption(f"📅 تاريخ الإيداع: {deposit_date}")
                if deposit_link and deposit_link not in ["","nan"]: st.markdown(f"📎 [عرض الملف المودع]({deposit_link})")
                st.markdown("""<div style="background:rgba(239,68,68,0.07);border:1px solid rgba(239,68,68,0.2);border-radius:10px;padding:11px 15px;margin-top:9px;"><p style="color:#EF4444!important;margin:0;font-size:0.86rem;">⛔ لقد قمت بإيداع مذكرتك؛ لا يمكنك إيداع نسخة أخرى.</p></div>""", unsafe_allow_html=True)
            elif deposit_status == "قابلة للمناقشة":
                # تحقق من حالة المناقشة
                _hal_check = str(memo_info.get("الحالة","") if memo_info is not None else "").strip()
                if _hal_check == "تمت المناقشة":
                    pass  # البطاقة الرئيسية تظهر أعلى
                else:
                    if _hal_done != "تمت المناقشة":
                        st.markdown("""<div class="notif-card notif-card-approved"><div class="notif-icon">🟢</div><div><div class="notif-title notif-title-approved">مذكرتك معتمدة — قابلة للمناقشة ✓</div><div class="notif-desc">وافق المشرف على مذكرتك رسمياً. ستتلقى إشعاراً من الإدارة بموعد المناقشة قريباً.</div></div></div>""", unsafe_allow_html=True)
                        if deposit_link and deposit_link not in ["","nan"]: st.markdown(f"📎 [عرض الملف المودع]({deposit_link})")

            _hal_check2 = str(memo_info.get("الحالة","") if memo_info is not None else "").strip()
            if _hal_check2 == "تمت المناقشة":
                # تمت المناقشة — أظهر بطاقة إعلامية فقط
                _def_date_done = str(memo_info.get("تاريخ المناقشة","")).strip() if memo_info is not None else ""
                _def_slot_done = str(memo_info.get("توقيت المناقشة","")).strip() if memo_info is not None else ""
                _def_room_done = str(memo_info.get("القاعة","")).strip() if memo_info is not None else ""
                st.markdown(f'''<div style="background:linear-gradient(135deg,rgba(16,185,129,0.12),rgba(16,185,129,0.04));border:2px solid rgba(16,185,129,0.4);border-radius:16px;padding:22px;text-align:center;margin-bottom:16px;">
                    <div style="font-size:2.5rem;">🎓</div>
                    <h3 style="color:#10B981!important;margin:8px 0;">تمت مناقشة مذكرتك</h3>
                    <p style="color:#E2E8F0!important;font-size:0.9rem;margin:4px 0;">بتاريخ <strong style="color:#FFD700;">{_def_date_done}</strong> الساعة <strong style="color:#FFD700;">{_def_slot_done}</strong> في <strong style="color:#FFD700;">{_def_room_done}</strong></p>
                    <p style="color:#94A3B8!important;font-size:0.85rem;margin-top:10px;">يجب عليك القيام بالإيداع النهائي للحصول على تبرئة المكتبة الضرورية للحصول على الشهادة.</p>
                </div>''', unsafe_allow_html=True)
            elif is_published and def_date_m and def_date_m not in ["","nan"]:
                st.markdown(f"""<div class="defense-schedule-card"><h4 style="color:#818CF8!important;margin:0 0 6px;">📅 موعد مناقشتك</h4><div class="defense-info-grid"><div class="defense-info-item"><div class="defense-info-label">📆 التاريخ</div><div class="defense-info-value">{def_date_m}</div></div><div class="defense-info-item"><div class="defense-info-label">🕐 التوقيت</div><div class="defense-info-value">{def_time_m if def_time_m and def_time_m!='nan' else '—'}</div></div><div class="defense-info-item"><div class="defense-info-label">🏛️ القاعة</div><div class="defense-info-value">{def_room_m if def_room_m and def_room_m!='nan' else '—'}</div></div></div></div>""", unsafe_allow_html=True)
                president_s = str(memo_info.get("AE","")).strip() if "AE" in memo_info.index else ""
                exam1_s     = str(memo_info.get("AD","")).strip() if "AD" in memo_info.index else ""
                exam2_s     = str(memo_info.get("AE","")).strip() if "AE" in memo_info.index else ""
                if president_s and president_s not in ["","nan"]:
                    members_html = f"""<div class="jury-member-card"><div class="jury-member-avatar avatar-president">🏛️</div><div class="jury-member-role role-president">رئيس اللجنة</div><div class="jury-member-name">{president_s}</div></div><div class="jury-member-card"><div class="jury-member-avatar avatar-supervisor">👨‍🏫</div><div class="jury-member-role role-supervisor">المشرف</div><div class="jury-member-name">{prof_name_m}</div></div>"""
                    if exam1_s and exam1_s!='nan': members_html += f"""<div class="jury-member-card"><div class="jury-member-avatar avatar-examiner">📋</div><div class="jury-member-role role-examiner">مناقش 1</div><div class="jury-member-name">{exam1_s}</div></div>"""
                    if exam2_s and exam2_s not in ['','nan']: members_html += f"""<div class="jury-member-card"><div class="jury-member-avatar avatar-examiner">📋</div><div class="jury-member-role role-examiner">مناقش 2</div><div class="jury-member-name">{exam2_s}</div></div>"""
                    st.markdown(f"""<div class="jury-card"><div class="jury-header"><div class="jury-header-icon">⚖️</div><div><div class="jury-header-title">لجنة مناقشة المذكرة رقم {note_num}</div><div class="jury-header-sub" style="color:rgba(255,255,255,0.85)!important;">{str(memo_info.get('عنوان المذكرة',''))[:60]}</div></div></div><div class="jury-members-grid">{members_html}</div></div>""", unsafe_allow_html=True)

            st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.07);margin:22px 0;'>", unsafe_allow_html=True)

            tab_memo, tab_track, tab_notify = st.tabs(["📄 مذكرتي", "📂 ملف الشهادة", "🔔 الإشعارات"])

            with tab_memo:
                import datetime as _dti_idaa
                _memo_row_idaa = memo_info
                _hal_idaa = str(_memo_row_idaa.get("الحالة","")).strip()
                _aq_idaa  = str(_memo_row_idaa.get("رابط المذكرة النهائية","")).strip()
                _ar_idaa  = str(_memo_row_idaa.get("رابط الملخص النهائي","")).strip()
                _as_idaa  = str(_memo_row_idaa.get("موافقة المشرف","")).strip()
                _at_idaa  = str(_memo_row_idaa.get("موافقة رئيس القسم","")).strip()
                _au_idaa  = str(_memo_row_idaa.get("تبرئة المكتبة","")).strip()

                if _hal_idaa == "تمت المناقشة":
                    # الإيداع النهائي فقط
                    st.markdown("### 📥 الإيداع النهائي")
                    def _step(_l,_d): return f'<span style="color:{"#10B981" if _d else "#94A3B8"};font-size:0.88rem;">{"✅" if _d else "⏳"} {_l}</span>'
                    st.markdown(" &nbsp;→&nbsp; ".join([
                        _step("رُفعت الملفات", bool(_aq_idaa and _aq_idaa not in ["","nan"])),
                        _step("موافقة المشرف", _as_idaa=="نعم"),
                        _step("موافقة رئيس القسم", _at_idaa=="نعم"),
                        _step("تبرئة المكتبة", _au_idaa=="نعم"),
                    ]), unsafe_allow_html=True)
                    if _au_idaa == "نعم":
                        st.success("🎉 اكتمل الإيداع النهائي!")
                    elif not (_aq_idaa and _aq_idaa not in ["","nan"]) or _as_idaa == "لا":
                        _mid_i = str(_memo_row_idaa.get("رقم المذكرة","")).strip()
                        st.info("📤 يُرجى رفع ملفَي الإيداع النهائي:")
                        _upl_pdf_i = st.file_uploader("📄 المذكرة النهائية (PDF):", type=["pdf"], key=f"idaa_pdf_{_mid_i}")
                        _upl_wrd_i = st.file_uploader("📝 ملف الملخص (عربي+إنجليزي):", type=["docx","doc"], key=f"idaa_word_{_mid_i}")
                        if st.button("📤 إيداع الملفات", type="primary", use_container_width=True, key=f"do_idaa_{_mid_i}"):
                            if not _upl_pdf_i or not _upl_wrd_i:
                                st.error("❌ يجب رفع الملفين معاً")
                            else:
                                with st.spinner("⏳ جاري الرفع..."):
                                    _ok1,_l1 = upload_to_drive(_upl_pdf_i.read(), f"مذكرة_نهائية_{_mid_i}.pdf", IDAA_FOLDER_ID, "application/pdf")
                                    _ok2,_l2 = upload_to_drive(_upl_wrd_i.read(), f"ملخص_{_mid_i}.docx", IDAA_FOLDER_ID, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                    if _ok1 and _ok2:
                                        _all_mi=load_memos(); _rni=None
                                        for _ii,(_, _rri) in enumerate(_all_mi.iterrows()):
                                            _mki=str(_rri.get("رقم المذكرة","")).strip()
                                            try: _mki=str(int(float(_mki)))
                                            except: pass
                                            if _mki==str(int(float(_mid_i))): _rni=_ii+2; break
                                        if _rni:
                                            sheets_service.spreadsheets().values().batchUpdate(
                                                spreadsheetId=MEMOS_SHEET_ID,
                                                body={"valueInputOption":"USER_ENTERED","data":[
                                                    {"range":f"Feuille 1!AQ{_rni}","values":[[_l1]]},
                                                    {"range":f"Feuille 1!AR{_rni}","values":[[_l2]]},
                                                    {"range":f"Feuille 1!AV{_rni}","values":[[_dti_idaa.datetime.now().strftime("%Y-%m-%d")]]},
                                                ]}
                                            ).execute()
                                            clear_cache_and_reload()
                                            st.success("✅ تم الإيداع — في انتظار موافقة المشرف"); st.rerun()
                                    else:
                                        if not _ok1: st.error(f"❌ {_l1}")
                                        if not _ok2: st.error(f"❌ {_l2}")
                    else:
                        st.info("✅ تم رفع الملفات — في انتظار الموافقات")
                        if _aq_idaa not in ["","nan"]: st.markdown(f"[📄 تحميل المذكرة النهائية]({_aq_idaa})")
                        if _ar_idaa not in ["","nan"]: st.markdown(f"[📝 تحميل الملخص]({_ar_idaa})")
                    if _as_idaa == "لا":
                        st.error("❌ رفض المشرف الإيداع — يُرجى تصحيح الملاحظات وإعادة الرفع")

                else:
                    # ما قبل المناقشة — المعلومات العادية
                    session_date = memo_info.get("موعد الجلسة القادمة","")
                    session_html = f"<p>📅 <b>موعد الجلسة القادمة:</b> {session_date}</p>" if session_date and str(session_date) not in ["","nan"] else ""
                    st.markdown(f"""<div class="card" style="border-top:3px solid #FFD700;"><h3>✅ مذكرتك المسجلة</h3><p><b>رقم المذكرة:</b> {memo_info['رقم المذكرة']}</p><p><b>العنوان:</b> {memo_info['عنوان المذكرة']}</p><p><b>المشرف:</b> {prof_name_m}</p><p><b>التخصص</b></p>{session_html}</div>""", unsafe_allow_html=True)
                    s1_ln,s1_fn=get_student_name_display(s1); s1_email=get_email_smart(s1)
                    st.markdown(f"""<div class="card"><h4 style="color:#FFD700;">👤 الطالب الأول</h4><p><b>{s1_ln} {s1_fn}</b></p></div>""", unsafe_allow_html=True)
                    if s2:
                        s2_ln,s2_fn=get_student_name_display(s2); s2_email=get_email_smart(s2)
                        st.markdown(f"""<div class="card"><h4 style="color:#FFD700;">👤 الطالب الثاني</h4><p><b>{s2_ln} {s2_fn}</b></p></div>""", unsafe_allow_html=True)

            with tab_track:
                st.subheader("📂 حالة ملف التخرج")
                df_students = load_students()
                def render_diploma(student_data, title):
                    cols_s = df_students.columns.tolist()
                    def gv(idx): return student_data.get(cols_s[idx],"غير محدد") if isinstance(student_data,dict) and len(cols_s)>idx else "غير محدد"
                    def badge(val):
                        v=str(val).strip()
                        if v in ["متوفر","مكتملة","كامل","جاهز","تم التسليم"]: return "status-available"
                        elif v in ["غير متوفر","ناقص","غير جاهز"]: return "status-unavailable"
                        return "status-pending"
                    items=[("📄 شهادة الميلاد",gv(14)),("📊 كشف M1",gv(15)),("📊 كشف M2",gv(16)),("📜 محضر المناقشة",gv(17)),("📂 حالة الملف",gv(18)),("🎓 حالة الشهادة",gv(19))]
                    rows_html="".join([f'<div class="diploma-item"><span>{label}</span><span class="status-badge {badge(val)}">{val}</span></div>' for label,val in items])
                    return f'<div class="card" style="border-right:4px solid #2F6F7E;"><h4 style="color:#FFD700;border-bottom:1px solid #334155;padding-bottom:9px;">{title}</h4>{rows_html}</div>'
                s1_ln,s1_fn=get_student_name_display(s1)
                st.markdown(render_diploma(s1, f"👤 {s1_ln} {s1_fn}"), unsafe_allow_html=True)

            with tab_notify:
                st.subheader("🔔 الإشعارات")
                df_reqs=load_requests()
                if note_num:
                    df_mn=load_memos()
                    my_row=df_mn[df_mn["رقم المذكرة"].astype(str).apply(normalize_text)==note_num]
                    sessions=pd.DataFrame()
                    if not my_row.empty:
                        my_prof=str(my_row.iloc[0]["الأستاذ"]).strip()
                        sessions=df_reqs[(df_reqs["النوع"]=="ملاحظة من المشرف") & (df_reqs["الأستاذ"].astype(str).str.strip()==my_prof)]
                        if not sessions.empty:
                            last=sessions.iloc[-1]; details=""
                            try:
                                if len(last)>8:
                                    rv=str(last.iloc[8]).strip()
                                    if rv and rv.lower() not in ['nan','none']:
                                        dm=re.search(r'(\d{4}-\d{2}-\d{2})',rv)
                                        if dm:
                                            try: dt_obj=datetime.strptime(dm.group(0),'%Y-%m-%d'); details=rv.replace(dm.group(0), format_arabic_date(dt_obj))
                                            except: details=rv
                                        else: details=rv
                            except: pass
                            if details: st.markdown(f"""<div class="notif-card notif-card-scheduled"><div class="notif-icon">📅</div><div><div class="notif-title notif-title-scheduled">ملاحظة من المشرف</div><div class="notif-desc">{details}</div></div></div>""", unsafe_allow_html=True)
                    my_reqs=df_reqs[df_reqs["رقم المذكرة"].astype(str).apply(normalize_text)==note_num]
                    if my_reqs.empty and sessions.empty: st.info("لا توجد إشعارات جديدة.")
                    else:
                        for _,r in my_reqs.iterrows():
                            rt=r['النوع']; details=""
                            if len(r)>8:
                                rv=str(r.iloc[8]).strip()
                                if rv and rv.lower() not in ['nan','none']: details=rv
                            st.markdown(f"""<div class="card" style="border-right:4px solid #F59E0B;"><h4>{rt}</h4><p>الحالة: <b>{r.get('الحالة','—')}</b></p>{'<p>'+details+'</p>' if details else ''}</div>""", unsafe_allow_html=True)


# ================================================================
# فضاء الأساتذة
# ================================================================
elif st.session_state.user_type == "professor":
    if not st.session_state.logged_in:
        render_countdown_banner()
        c1,c2=st.columns([4,1])
        with c2:
            if st.button("رجوع"): st.session_state.user_type=None; st.rerun()
        st.markdown("<h2>📚 فضاء الأساتذة</h2>", unsafe_allow_html=True)
        with st.form("prof_login"):
            c1,c2=st.columns(2)
            with c1: u=st.text_input("اسم المستخدم")
            with c2: p=st.text_input("كلمة المرور", type="password")
            if st.form_submit_button("تسجيل الدخول"):
                df_prof_memos = load_prof_memos()
                v,r=verify_professor(u,p,df_prof_memos)
                if not v: st.error(r)
                else:
                    st.session_state.professor=r; st.session_state.logged_in=True
                    st.query_params['ut']='professor'
                    st.query_params['un']=encode_str(normalize_text(r.get('إسم المستخدم','')))
                    st.rerun()
    else:
        render_countdown_banner()
        prof=st.session_state.professor; prof_name=prof["الأستاذ"]

        if st.session_state.get('selected_memo_id'):
            memo_id=st.session_state.selected_memo_id
            df_m_fresh=load_memos()
            df_students=load_students()
            current_memo=df_m_fresh[df_m_fresh["رقم المذكرة"].astype(str).apply(normalize_text)==memo_id].iloc[0]
            student_info=get_student_info_from_memo(current_memo, df_students)
            c_back,_=st.columns([1,6])
            with c_back:
                if st.button("⬅️ العودة"): st.session_state.selected_memo_id=None; st.session_state.prof_action=None; st.rerun()

            deposit_status=str(current_memo.get("حالة الإيداع","")).strip()
            deposit_link=str(current_memo.get("رابط الملف","")).strip()
            deposit_date=str(current_memo.get("تاريخ إيداع المذكرة","")).strip()

            prog_val=str(current_memo.get('نسبة التقدم','0')).strip()
            try: prog_int=int(prog_val) if prog_val else 0
            except: prog_int=0

            dep_color={"مودعة":"#F59E0B","قابلة للمناقشة":"#10B981","مرفوضة":"#EF4444"}.get(deposit_status,"#ffffff")
            dep_label={"مودعة":"📤 مودعة","قابلة للمناقشة":"🟢 معتمدة","مرفوضة":"🔴 معادة"}.get(deposit_status,"⏳ لم تودَع")
            st.markdown(f"""<div style="background:linear-gradient(135deg,#0F2942,#1A3A5C);border-radius:16px;padding:20px 24px;margin-bottom:22px;border:1px solid rgba(47,111,126,0.3);display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:12px;"><div><p style="font-size:2.4rem;font-weight:900;color:#FFD700;margin:0;line-height:1;">{current_memo['رقم المذكرة']}</p><p style="font-size:1.05rem;font-weight:700;color:#ffffff;margin:4px 0;">{current_memo['عنوان المذكرة']}</p><p style="font-size:0.83rem;color:#ffffff;margin:0;">{current_memo['التخصص']} | نسبة الإنجاز: {prog_int}%</p></div><div style="background:rgba(0,0,0,0.25);color:#ffffff;padding:7px 16px;border-radius:20px;font-weight:700;font-size:0.88rem;border:1px solid {dep_color};">{dep_label}</div></div>""", unsafe_allow_html=True)

            cards_html = f"""<div class="student-card"><h4 style="color:#FFD700;margin-top:0;">الطالب الأول</h4><p style="font-size:1.1rem;font-weight:700;margin:9px 0 3px;">{student_info['s1_name']}</p><p style="color:#E2E8F0;font-size:0.83rem;">رقم التسجيل: {student_info['s1_reg'] or '—'}</p><div style="background:rgba(16,185,129,0.07);border-radius:8px;padding:7px;margin-top:10px;color:#10B981;font-size:0.83rem;">📧 {student_info['s1_email'] or 'غير متوفر'}</div></div>"""
            if student_info['s2_name']:
                cards_html += f"""<div class="student-card"><h4 style="color:#FFD700;margin-top:0;">الطالب الثاني</h4><p style="font-size:1.1rem;font-weight:700;margin:9px 0 3px;">{student_info['s2_name']}</p><p style="color:#E2E8F0;font-size:0.83rem;">رقم التسجيل: {student_info['s2_reg'] or '—'}</p><div style="background:rgba(16,185,129,0.07);border-radius:8px;padding:7px;margin-top:10px;color:#10B981;font-size:0.83rem;">📧 {student_info['s2_email'] or 'غير متوفر'}</div></div>"""
            st.markdown(f'<div class="students-grid">{cards_html}</div>', unsafe_allow_html=True)

            c1,c2=st.columns(2)
            with c1:
                new_prog=st.selectbox("📊 تحديث نسبة التقدم:",["0%","10% - ضبط المقدمة","30% - الفصل الأول","60% - الفصل الثاني","80% - الخاتمة","100% - مكتملة"],key=f"prog_{memo_id}")
                if st.button("حفظ التقدم",key=f"save_prog_{memo_id}"):
                    mapping={"0%":0,"10% - ضبط المقدمة":10,"30% - الفصل الأول":30,"60% - الفصل الثاني":60,"80% - الخاتمة":80,"100% - مكتملة":100}
                    ok,msg=update_progress(memo_id,mapping[new_prog])
                    if ok: st.success(msg); time_module.sleep(1); st.rerun()
                    else: st.error(msg)

            st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.07);margin:20px 0;'>", unsafe_allow_html=True)
            st.markdown("<h3 style='text-align:center;margin-bottom:16px;'>📥 حالة إيداع المذكرة</h3>", unsafe_allow_html=True)

            _ah_val_p = str(current_memo.get("مفقودة","")).strip() or str(current_memo.get("AH","")).strip()
            is_extended = _ah_val_p == "2"
            is_missing = _ah_val_p == "1"

            if not deposit_status or deposit_status in ["nan",""]:
                st.markdown("""<div style="background:rgba(47,111,126,0.08);border:1px solid rgba(47,111,126,0.3);border-radius:14px;padding:22px;text-align:center;"><div style="font-size:2.3rem;">⏳</div><p style="color:#ffffff!important;font-size:0.95rem;margin:7px 0;">لم يودع الطالب المذكرة بعد.</p></div>""", unsafe_allow_html=True)

            elif deposit_status == "مودعة" and not is_extended:
                st.markdown(f"""<div style="background:linear-gradient(135deg,rgba(245,158,11,0.1),rgba(245,158,11,0.04));border:2px solid rgba(245,158,11,0.42);border-radius:16px;padding:20px;margin-bottom:16px;"><h4 style="color:#F59E0B!important;margin:0 0 5px;">📥 مذكرة بانتظار مراجعتك</h4><p style="color:#E2E8F0!important;margin:0;font-size:0.88rem;">تاريخ الإيداع: <strong style="color:#FFD700;">{deposit_date if deposit_date and deposit_date!='nan' else '—'}</strong></p></div>""", unsafe_allow_html=True)
                if deposit_link and deposit_link not in ["","nan"]:
                    st.markdown(f"""<div style="text-align:center;margin:14px 0 20px;"><a href="{deposit_link}" target="_blank" style="display:inline-block;background:linear-gradient(135deg,#1E3A5F,#2F6F7E);color:#ffffff;padding:15px 38px;border-radius:13px;text-decoration:none;font-size:1.05rem;font-weight:700;box-shadow:0 8px 22px rgba(47,111,126,0.38);">📄 الاطلاع على المذكرة المودعة</a></div>""", unsafe_allow_html=True)
                if not st.session_state.get('prof_action'):
                    st.markdown("#### اتخذ قرارك:")
                    ca,cb,cc=st.columns(3)
                    with ca:
                        if st.button("✅ موافقة على المذكرة",use_container_width=True,key=f"btn_approve_{memo_id}"):
                            st.session_state['prof_action']='approve'; st.rerun()
                    with cb:
                        if st.button("🔴 إعادة للمراجعة",use_container_width=True,key=f"btn_reject_{memo_id}"):
                            st.session_state['prof_action']='reject'; st.rerun()
                    with cc:
                        if st.button("💬 إرسال ملاحظات",use_container_width=True,key=f"btn_notes_{memo_id}"):
                            st.session_state['prof_action']='notes'; st.rerun()

                elif st.session_state.get('prof_action') == 'approve':
                    s1_name_ap=str(current_memo.get("الطالب الأول","")).strip()
                    s2_name_ap=str(current_memo.get("الطالب الثاني","")).strip()
                    students_str_ap=s1_name_ap
                    if s2_name_ap and s2_name_ap not in ["","nan","--"]: students_str_ap+=f" و {s2_name_ap}"
                    st.markdown("""<div class="declaration-card"><div class="declaration-card-header"><div class="declaration-card-title">📋 التصريح الرسمي بالموافقة</div><div class="declaration-card-sub">يُرجى إدخال عدد الصفحات ثم الضغط على تأكيد الموافقة.</div></div><div class="declaration-card-body">""", unsafe_allow_html=True)
                    st.markdown('<div class="declaration-step-label">① عدد صفحات المذكرة (دليل على اطلاعك الكامل)</div>', unsafe_allow_html=True)
                    page_count = st.number_input("عدد الصفحات",min_value=0,max_value=999,value=0,step=1,key=f"pages_{memo_id}")
                    st.markdown(f"""<div class="declaration-step-label" style="margin-top:14px;">② نص التصريح الذي سيُحفظ</div><div class="declaration-preview">أنا الأستاذ <strong>{prof_name}</strong>، أصرّح للطالب(ين) {students_str_ap} بإيداع المذكرة رقم {memo_id} عدد الصفحات: <strong>{page_count}</strong> بعنوان «{current_memo.get('عنوان المذكرة','')}»، وأصرح بأن النسخة المودعة هي التي ستُعرض على لجنة المناقشة.</div>""", unsafe_allow_html=True)
                    st.markdown('<div class="declaration-step-label" style="margin-top:14px;">③ الإقرار والتوقيع الإلكتروني</div>', unsafe_allow_html=True)
                    st.markdown(f"""<div style="background:rgba(47,111,126,0.07);border-radius:8px;padding:8px 13px;margin:7px 0 5px;"><p style="color:#E2E8F0!important;font-size:0.78rem;margin:0;">سيتم توثيق موافقتك باسم: <strong style="color:#FFD700;">{prof_name}</strong></p></div>""", unsafe_allow_html=True)
                    signature = prof_name
                    st.markdown("</div></div>", unsafe_allow_html=True)
                    col_ok,col_cancel=st.columns(2)
                    with col_ok:
                        if not st.session_state.get(f"confirm_step_{memo_id}",False):
                            if st.button("📋 متابعة للتأكيد النهائي",type="primary",use_container_width=True,key=f"pre_approve_{memo_id}"):
                                errs=[]
                                if page_count < 60: errs.append("❌ عدد الصفحات يجب أن يكون 60 على الأقل")
                                if not signature.strip(): errs.append("❌ التوقيع فارغ")
                                elif signature.strip()!=prof_name.strip(): errs.append(f"❌ التوقيع غير مطابق — اكتب: «{prof_name}»")
                                if errs:
                                    for e in errs: st.error(e)
                                else:
                                    st.session_state[f"confirm_step_{memo_id}"]=True
                                    st.session_state[f"sig_value_{memo_id}"]=signature.strip()
                                    st.session_state[f"pages_value_{memo_id}"]=page_count
                                    st.rerun()
                        else:
                            st.markdown("""<div style="background:rgba(239,68,68,0.09);border:2px solid #EF4444;border-radius:13px;padding:18px;text-align:center;margin-bottom:13px;"><div style="font-size:1.7rem;">⚠️</div><h4 style="color:#EF4444!important;margin:7px 0;">تأكيد نهائي</h4><p style="color:#E2E8F0!important;font-size:0.86rem;">ستُحفظ موافقتك وتُرسل الإشعارات للطلبة.</p></div>""", unsafe_allow_html=True)
                            if st.button("✅ نعم، أوافق نهائياً",type="primary",use_container_width=True,key=f"final_approve_{memo_id}"):
                                sig_saved=st.session_state.get(f"sig_value_{memo_id}","")
                                pages_saved=st.session_state.get(f"pages_value_{memo_id}",page_count)
                                with st.spinner("⏳ جاري حفظ التصريح..."):
                                    timestamp_ap=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                                    full_declaration=f"تصريح: {prof_name} | {timestamp_ap} | أنا الأستاذ {prof_name}، أصرّح للطالب(ين) {students_str_ap} بإيداع المذكرة رقم {memo_id} عدد الصفحات: {pages_saved} بعنوان «{current_memo.get('عنوان المذكرة','')}»، وأصرح بأن النسخة المودعة هي التي ستُعرض على لجنة المناقشة."
                                    save_approval_declaration(memo_id, prof_name, sig_saved, full_declaration)
                                    ok,msg=approve_memo_for_defense(memo_id)
                                    if ok:
                                        df_sf=load_students(); df_sf['رقم التسجيل_norm']=df_sf['رقم التسجيل'].astype(str).apply(normalize_text)
                                        memo_vals=current_memo.tolist() if hasattr(current_memo,'tolist') else list(current_memo.values())
                                        r1=normalize_text(current_memo.get("رقم تسجيل الطالب 1",memo_vals[18] if len(memo_vals)>18 else ""))
                                        r2=normalize_text(current_memo.get("رقم تسجيل الطالب 2",memo_vals[19] if len(memo_vals)>19 else ""))
                                        s1d=s2d=None
                                        if r1 and r1 not in ["","nan"]:
                                            rr=df_sf[df_sf["رقم التسجيل_norm"]==r1]
                                            if not rr.empty: s1d=rr.iloc[0].to_dict()
                                        if r2 and r2 not in ["","nan"]:
                                            rr=df_sf[df_sf["رقم التسجيل_norm"]==r2]
                                            if not rr.empty: s2d=rr.iloc[0].to_dict()
                                        send_approval_email_to_students(memo_id,str(current_memo.get('عنوان المذكرة','')),prof_name,s1d,s2d)
                                        for k in [f"confirm_step_{memo_id}",f"sig_value_{memo_id}",f"pages_value_{memo_id}"]: st.session_state.pop(k,None)
                                        st.session_state['prof_action']=None
                                        st.success("✅ تمت الموافقة وحُفظ التصريح. تم إشعار الطلبة.")
                                        st.balloons(); clear_cache_and_reload(); time_module.sleep(2); st.rerun()
                                    else: st.error(msg)
                    with col_cancel:
                        if st.button("إلغاء",use_container_width=True,key=f"cancel_ap_{memo_id}"):
                            st.session_state['prof_action']=None; st.session_state.pop(f"confirm_step_{memo_id}",None); st.rerun()

                elif st.session_state.get('prof_action') == 'reject':
                    st.markdown("""<div style="background:rgba(239,68,68,0.07);border:2px solid rgba(239,68,68,0.32);border-radius:15px;padding:20px;margin-bottom:14px;"><h4 style="color:#EF4444!important;margin:0 0 9px;">🔴 إعادة المذكرة للمراجعة</h4><p style="color:#E2E8F0!important;font-size:0.86rem;margin:0;">سيتم إرسال ملاحظاتك للطالب وفتح الإيداع مجدداً.</p></div>""", unsafe_allow_html=True)
                    rejection_reason=st.text_area("📋 سبب الإعادة وملاحظاتك التفصيلية:",height=140,placeholder="اكتب ملاحظاتك بدقة...",key=f"rej_reason_{memo_id}")
                    cr1,cr2=st.columns(2)
                    with cr1:
                        if st.button("🔴 تأكيد الإعادة للمراجعة",type="primary",use_container_width=True,key=f"confirm_rej_{memo_id}"):
                            if not rejection_reason.strip(): st.error("❌ يجب كتابة سبب الإعادة")
                            else:
                                with st.spinner("⏳ جاري الحفظ..."):
                                    ok,msg=reject_memo_and_reopen(memo_id,prof_name,rejection_reason)
                                    if ok:
                                        df_sf=load_students(); df_sf['رقم التسجيل_norm']=df_sf['رقم التسجيل'].astype(str).apply(normalize_text)
                                        memo_vals=current_memo.tolist() if hasattr(current_memo,'tolist') else list(current_memo.values())
                                        r1=normalize_text(current_memo.get("رقم تسجيل الطالب 1",memo_vals[18] if len(memo_vals)>18 else ""))
                                        r2=normalize_text(current_memo.get("رقم تسجيل الطالب 2",memo_vals[19] if len(memo_vals)>19 else ""))
                                        s1d=s2d=None
                                        if r1: rr=df_sf[df_sf["رقم التسجيل_norm"]==r1]; s1d=rr.iloc[0].to_dict() if not rr.empty else None
                                        if r2: rr=df_sf[df_sf["رقم التسجيل_norm"]==r2]; s2d=rr.iloc[0].to_dict() if not rr.empty else None
                                        send_rejection_email_to_students(memo_id,str(current_memo.get('عنوان المذكرة','')),prof_name,rejection_reason,s1d,s2d)
                                        st.session_state['prof_action']=None
                                        st.success("✅ تم تسجيل الإعادة. أُرسل إشعار للطلبة وفُتح الإيداع من جديد.")
                                        time_module.sleep(2); st.rerun()
                                    else: st.error(msg)
                    with cr2:
                        if st.button("إلغاء",use_container_width=True,key=f"cancel_rej_{memo_id}"):
                            st.session_state['prof_action']=None; st.rerun()

                elif st.session_state.get('prof_action') == 'notes':
                    st.markdown("""<div style="background:rgba(99,102,241,0.08);border:2px solid rgba(99,102,241,0.32);border-radius:15px;padding:20px;margin-bottom:14px;"><h4 style="color:#818CF8!important;margin:0 0 9px;">💬 إرسال ملاحظات للطالب</h4><p style="color:#E2E8F0!important;font-size:0.86rem;margin:0;">ستُرسل ملاحظاتك للطالب. تبقى حالة الإيداع كما هي.</p></div>""", unsafe_allow_html=True)
                    notes_text=st.text_area("📝 ملاحظاتك:",height=130,placeholder="اكتب ملاحظاتك هنا...",key=f"notes_txt_{memo_id}")
                    cn1,cn2=st.columns(2)
                    with cn1:
                        if st.button("💬 إرسال الملاحظات",type="primary",use_container_width=True,key=f"send_notes_{memo_id}"):
                            if not notes_text.strip(): st.error("❌ يجب كتابة الملاحظات")
                            else:
                                ok,msg=save_prof_notes(memo_id,prof_name,notes_text)
                                if ok:
                                    df_sf=load_students(); df_sf['رقم التسجيل_norm']=df_sf['رقم التسجيل'].astype(str).apply(normalize_text)
                                    memo_vals=current_memo.tolist() if hasattr(current_memo,'tolist') else list(current_memo.values())
                                    r1=normalize_text(current_memo.get("رقم تسجيل الطالب 1",memo_vals[18] if len(memo_vals)>18 else ""))
                                    r2=normalize_text(current_memo.get("رقم تسجيل الطالب 2",memo_vals[19] if len(memo_vals)>19 else ""))
                                    s1d=s2d=None
                                    if r1: rr=df_sf[df_sf["رقم التسجيل_norm"]==r1]; s1d=rr.iloc[0].to_dict() if not rr.empty else None
                                    if r2: rr=df_sf[df_sf["رقم التسجيل_norm"]==r2]; s2d=rr.iloc[0].to_dict() if not rr.empty else None
                                    send_notes_email_to_students(memo_id,str(current_memo.get('عنوان المذكرة','')),prof_name,notes_text,s1d,s2d)
                                    st.session_state['prof_action']=None
                                    st.success("✅ تم إرسال الملاحظات للطالب."); time_module.sleep(1); st.rerun()
                                else: st.error(msg)
                    with cn2:
                        if st.button("إلغاء",use_container_width=True,key=f"cancel_notes_{memo_id}"):
                            st.session_state['prof_action']=None; st.rerun()

            elif deposit_status == "قابلة للمناقشة":
                st.markdown(f"""<div style="background:rgba(16,185,129,0.09);border:2px solid #10B981;border-radius:15px;padding:20px;text-align:center;"><div style="font-size:2.3rem;">🟢</div><h4 style="color:#10B981!important;margin:9px 0;">المذكرة معتمدة — قابلة للمناقشة</h4><p style="color:#E2E8F0!important;">تم التصريح رسمياً وأُرسلت الإشعارات للطلبة.</p>{f'<a href="{deposit_link}" target="_blank" style="color:#10B981;font-size:0.88rem;">📄 عرض المذكرة المودعة</a>' if deposit_link and deposit_link!='nan' else ''}</div>""", unsafe_allow_html=True)

            st.markdown("<hr style='border:none;border-top:1px solid rgba(255,255,255,0.07);margin:20px 0;'>", unsafe_allow_html=True)

        else:
            # ================================================================
            # الصفحة الرئيسية لفضاء الأستاذ
            # ================================================================
            c1,c2=st.columns([4,1])
            with c2:
                if st.button("خروج"): logout()
            st.markdown(f"<h2>📚 فضاء الأستاذ <span style='color:#FFD700;'>{prof_name}</span></h2>", unsafe_allow_html=True)
            df_m_fresh=load_memos()
            prof_memos=df_m_fresh[df_m_fresh["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
            total=len(prof_memos); registered=len(prof_memos[prof_memos["تم التسجيل"].astype(str).str.strip()=="نعم"]); available=total-registered; is_exhausted=registered>=4
            col_dep="حالة الإيداع"
            deposited_memos_pr=prof_memos[prof_memos["تم التسجيل"].astype(str).str.strip()=="نعم"]

            # ================================================================
            # ✅ التعديل 1: بطاقات المذكرات المودعة قابلة للنقر المباشر
            # ================================================================
            if col_dep in deposited_memos_pr.columns:
                pending_deposit=deposited_memos_pr[deposited_memos_pr[col_dep].astype(str).str.strip()=="مودعة"]
                if not pending_deposit.empty:
                    items_html=""
                    for _,dep_memo in pending_deposit.iterrows():
                        dep_date=str(dep_memo.get("تاريخ إيداع المذكرة","")).strip()
                        items_html+=f"""<div class="prof-deposit-item">
                            <div>
                                <div class="prof-deposit-memo-num">{dep_memo['رقم المذكرة']}</div>
                                <div class="prof-deposit-memo-title">{str(dep_memo['عنوان المذكرة'])[:52]}{'...' if len(str(dep_memo['عنوان المذكرة']))>52 else ''}</div>
                                <div class="prof-deposit-memo-date">📅 {dep_date if dep_date and dep_date!='nan' else '—'}</div>
                            </div>
                            <div style="background:rgba(245,158,11,0.1);color:#F59E0B;padding:4px 12px;border-radius:20px;font-size:0.78rem;font-weight:700;border:1px solid rgba(245,158,11,0.28);">⏳ بانتظار مراجعتك</div>
                        </div>"""

                    st.markdown(f"""<div class="prof-deposit-alert">
                        <div class="prof-deposit-alert-header">
                            <div class="prof-deposit-alert-icon">📥</div>
                            <div>
                                <div class="prof-deposit-alert-title">لديك {len(pending_deposit)} مذكرة مودعة — اضغط للمراجعة المباشرة</div>
                                <div class="prof-deposit-alert-sub">اضغط على زر المذكرة المعنية للانتقال إليها مباشرةً.</div>
                            </div>
                        </div>
                        <div class="prof-deposit-list">{items_html}</div>
                    </div>""", unsafe_allow_html=True)

                    # أزرار النقر المباشر — واحد لكل مذكرة مودعة
                    for _,dep_memo in pending_deposit.iterrows():
                        memo_num_btn   = str(dep_memo['رقم المذكرة'])
                        memo_title_btn = str(dep_memo['عنوان المذكرة'])
                        dep_date_btn   = str(dep_memo.get("تاريخ إيداع المذكرة","")).strip()
                        col_info, col_btn = st.columns([3, 1])
                        with col_info:
                            st.markdown(f"""
                            <div style="background:linear-gradient(135deg,rgba(245,158,11,0.1),rgba(245,158,11,0.04));
                                        border:1.5px solid rgba(245,158,11,0.38);border-radius:12px;
                                        padding:12px 16px;margin-bottom:6px;">
                                <span style="font-size:1.3rem;font-weight:900;color:#FFD700;">{memo_num_btn}</span>
                                &nbsp;—&nbsp;
                                <span style="color:#E2E8F0;font-size:0.88rem;">
                                    {memo_title_btn[:45]}{'...' if len(memo_title_btn)>45 else ''}
                                </span>
                                <div style="font-size:0.76rem;color:#F59E0B;margin-top:4px;">
                                    📅 {dep_date_btn if dep_date_btn and dep_date_btn!='nan' else '—'}
                                </div>
                            </div>""", unsafe_allow_html=True)
                        with col_btn:
                            if st.button(
                                f"📂 مراجعة ←",
                                key=f"quick_open_{memo_num_btn}",
                                use_container_width=True
                            ):
                                st.session_state.selected_memo_id = memo_num_btn
                                st.session_state.prof_action = None
                                st.rerun()

            # إحصائيات عضوية الأستاذ في اللجان
            df_jury_kpi = load_memos()
            sup_count = len(df_jury_kpi[df_jury_kpi["الأستاذ"].astype(str).str.strip()==prof_name.strip()]) if "الأستاذ" in df_jury_kpi.columns else 0
            pres_count = len(df_jury_kpi[df_jury_kpi["الرئيس"].astype(str).str.strip()==prof_name.strip()]) if "الرئيس" in df_jury_kpi.columns else 0
            ex_count = len(df_jury_kpi[df_jury_kpi["المناقش1"].astype(str).str.strip()==prof_name.strip()]) + len(df_jury_kpi[df_jury_kpi["المناقش2"].astype(str).str.strip()==prof_name.strip()]) if "المناقش1" in df_jury_kpi.columns else 0
            total_roles = sup_count + pres_count + ex_count
            st.markdown(f'''<div class="kpi-grid">
                <div class="kpi-card" style="border-top:3px solid #2F9EA0;"><div class="kpi-value" style="color:#2F9EA0;">{sup_count}</div><div class="kpi-label">📚 مشرف</div></div>
                <div class="kpi-card" style="border-top:3px solid #FFD700;"><div class="kpi-value" style="color:#FFD700;">{pres_count}</div><div class="kpi-label">🏛️ رئيس لجنة</div></div>
                <div class="kpi-card" style="border-top:3px solid #818CF8;"><div class="kpi-value" style="color:#818CF8;">{ex_count}</div><div class="kpi-label">📋 مناقش</div></div>
                <div class="kpi-card" style="border-top:3px solid #10B981;"><div class="kpi-value" style="color:#10B981;">{total_roles}</div><div class="kpi-label">🔢 المجموع</div></div>
            </div>''', unsafe_allow_html=True)

            st.markdown('''<div style="text-align:center;padding:18px 24px;margin-bottom:4px;">
                <div style="font-size:1.15rem;font-weight:900;color:#FFD700;margin-bottom:4px;">⚠️ ملاحظة هامة</div>
                <div style="font-size:1.05rem;font-weight:700;color:#ffffff;">سيتم إعلامكم ببرمجة أي مذكرة عبر الإيميل المهني.</div>
            </div>
            <hr style="border:none;border-top:1px solid rgba(255,255,255,0.12);margin:8px 0 16px 0;">''', unsafe_allow_html=True)

            tab5, = st.tabs(["📅 برنامج المناقشات"])
            with tab5:
                st.markdown('<h3 style="text-align:center;color:#FFD700;">📅 برنامج المناقشات</h3>', unsafe_allow_html=True)
                df_m_jury = load_memos()
                jury_memos = pd.DataFrame()

                if not df_m_jury.empty:
                    # تحقق من صلاحية "مسؤول" في شيت الأساتذة
                    _is_massoul = False
                    _prof_row_ms = df_p_tk[df_p_tk["الأستاذ"].astype(str).str.strip()==prof_name.strip()] if "df_p_tk" in dir() else pd.DataFrame()
                    if _prof_row_ms.empty:
                        _df_profs_ms = load_prof_memos()
                        _prof_row_ms = _df_profs_ms[_df_profs_ms["الأستاذ"].astype(str).str.strip()==prof_name.strip()]
                    if not _prof_row_ms.empty:
                        _massoul_val = str(_prof_row_ms.iloc[0].get("مسؤول","")).strip()
                        _is_massoul = _massoul_val == "نعم"
                    st.session_state["_prof_is_massoul"] = _is_massoul

                    masks = []
                    # كل الأساتذة يرون برنامجهم كاملاً (مشرف + رئيس + مناقش)
                    _roles_to_show = [
                        ("الأستاذ","مشرف"),
                        ("الرئيس","رئيس لجنة"),
                        ("المناقش1","مناقش"),
                        ("المناقش2","مناقش")
                    ]

                    for col_j, role_j in _roles_to_show:
                        if col_j in df_m_jury.columns:
                            mm = df_m_jury[df_m_jury[col_j].astype(str).str.strip() == prof_name.strip()].copy()
                            if not mm.empty:
                                mm["الصفة"] = role_j
                                masks.append(mm)
                    if masks:
                        jury_memos = pd.concat(masks).drop_duplicates(subset=["رقم المذكرة"])
                        # البحث عن عمود نشر البرنامج ثم AD ثم منشور
                        # AD = نشر المذكرة (يظهر البطاقة)
                        # AI = نشر البرنامج (يظهر الموعد)
                        col_pub = "منشور" if "منشور" in jury_memos.columns else ("AD" if "AD" in jury_memos.columns else None)
                        if col_pub:
                            is_pub = jury_memos[col_pub].astype(str).str.strip() == "نعم"
                            is_sup = jury_memos["الصفة"] == "مشرف"
                            # المشرف يرى مذكراته دائماً
                            # الرئيس والمناقش: فقط إذا نشر البرنامج = نعم
                            jury_memos = jury_memos[is_pub | is_sup]
                        else:
                            # إذا لم يوجد عمود نشر → لا يرى الرئيس والمناقش شيئاً
                            is_sup = jury_memos["الصفة"] == "مشرف"
                            jury_memos = jury_memos[is_sup]

                if jury_memos.empty:
                    st.info("⏳ لا توجد مناقشات منشورة تخصك حالياً.")
                else:
                    role_icons = {"مشرف":"👨‍🏫","رئيس لجنة":"🏛️","مناقش":"📋"}
                    role_colors = {"مشرف":"#2F9EA0","رئيس لجنة":"#FFD700","مناقش":"#94A3B8"}
                    role_counts = jury_memos["الصفة"].value_counts().to_dict()

                    # ملخص - وسط الصفحة
                    summary_parts = " &nbsp;&nbsp;|&nbsp;&nbsp; ".join([f"{role_icons.get(r,'')} {r}: <strong>{n}</strong>" for r,n in role_counts.items()])
                    st.markdown(f'<div style="text-align:center;padding:10px 0;font-size:1rem;color:#E2E8F0;">{summary_parts}</div>', unsafe_allow_html=True)

                    # فلتر
                    roles_list = ["الكل"] + [r for r in role_icons if r in role_counts]
                    selected_role = st.radio("", roles_list, key="jury_role_filter", horizontal=True)
                    filtered = jury_memos if selected_role == "الكل" else jury_memos[jury_memos["الصفة"] == selected_role]

                    # ترتيب
                    filtered = filtered.copy()
                    filtered["_has_date"] = filtered["تاريخ المناقشة"].astype(str).str.strip().apply(lambda x: 0 if x and x not in ["","nan"] else 1)
                    filtered = filtered.sort_values(["_has_date","تاريخ المناقشة"]).drop(columns=["_has_date"]).reset_index(drop=True)

                    st.markdown(f"**{len(filtered)} مذكرة**")

                    # البطاقات
                    cards_html = ""
                    for _, jm in filtered.iterrows():
                        jmid  = str(jm.get("رقم المذكرة","")).strip()
                        jtitle= str(jm.get("عنوان المذكرة","")).strip()
                        jrole = str(jm.get("الصفة","")).strip()
                        jlink = str(jm.get("رابط الملف","")).strip()
                        jdate = str(jm.get("تاريخ المناقشة","")).strip()
                        # تحقق من نشر البرنامج لهذا الصف
                        _col_pub_r = next((c for c in ["نشر البرنامج","AI","منشور"] if c in filtered.columns), None)
                        is_pub_row = str(jm.get(_col_pub_r,"")).strip() == "نعم" if _col_pub_r else False
                        jtime = str(jm.get("توقيت المناقشة","")).strip()
                        jroom = str(jm.get("القاعة","")).strip()
                        jsup       = str(jm.get("الأستاذ","")).strip()
                        # العمود T = حالة الإيداع
                        jdeposit = ""
                        _cols_list = list(jm.index)
                        for _cn in _cols_list:
                            if "حالة" in str(_cn) and "إيداع" in str(_cn):
                                _v = str(jm[_cn]).strip()
                                if _v and _v not in ["nan","None",""]:
                                    jdeposit = _v
                                    break

                        has_date = jdate and jdate not in ["","nan"]
                        j_published = str(jm.get("نشر البرنامج","")).strip() == "نعم"
                        has_time = jtime and jtime not in ["","nan"]
                        has_room = jroom and jroom not in ["","nan"]
                        has_link = jlink and jlink not in ["","nan"]

                        r_color = role_colors.get(jrole,"#94A3B8")
                        r_icon  = role_icons.get(jrole,"📄")

                        # زر المعاينة حسب حالة الإيداع والصفة
                        is_supervisor_role = (jrole == "مشرف")
                        # تنظيف jdeposit
                        jdeposit = jdeposit.strip()


                        if not jdeposit or jdeposit in ["","nan"]:
                            # لم يُودع بعد
                            preview_btn = '<span style="background:rgba(239,68,68,0.12);color:#EF4444;padding:4px 10px;border-radius:6px;font-size:0.78rem;font-weight:600;">⚠️ لم يتم الإيداع بعد</span>'
                        elif jdeposit == "قابلة للمناقشة":
                            # وافق المشرف — الجميع يرى
                            if has_link:
                                preview_btn = f'<a href="{jlink}" target="_blank" style="background:#1E3A5F;color:#fff;padding:5px 14px;border-radius:8px;text-decoration:none;font-size:0.82rem;font-weight:700;">👁️ معاينة</a>'
                            else:
                                preview_btn = '<span style="color:#64748B;font-size:0.78rem;">لا ملف</span>'
                        elif jdeposit in ["مودعة","مرفوضة"]:
                            # لم يوافق المشرف بعد — المشرف فقط
                            if is_supervisor_role and has_link:
                                preview_btn = f'<a href="{jlink}" target="_blank" style="background:#1E3A5F;color:#fff;padding:5px 14px;border-radius:8px;text-decoration:none;font-size:0.82rem;font-weight:700;">👁️ معاينة</a>'
                            else:
                                preview_btn = '<span style="background:rgba(245,158,11,0.15);color:#F59E0B;padding:4px 10px;border-radius:6px;font-size:0.78rem;font-weight:600;">⏳ في انتظار موافقة المشرف</span>'
                        else:
                            preview_btn = '<span style="background:rgba(245,158,11,0.15);color:#F59E0B;padding:4px 10px;border-radius:6px;font-size:0.78rem;font-weight:600;">⏳ في انتظار موافقة المشرف</span>'

                        # حالة المناقشة من الشيت
                        _jhal = str(jm.get("الحالة","")).strip() if "الحالة" in filtered.columns else ""
                        _jpdf = str(jm.get("رابط المحضر","")).strip() if "رابط المحضر" in filtered.columns else ""
                        _jnote = str(jm.get("ملاحظات","")).strip() if "ملاحظات" in filtered.columns else ""
                        _hal_html = ""
                        if _jhal == "تمت":
                            _hal_html = '<div style="margin-top:6px;">'
                            _hal_html += '<span style="background:#10B981;color:#fff;padding:2px 10px;border-radius:12px;font-size:0.75rem;">✅ تمت المناقشة</span>'
                            if _jpdf and _jpdf not in ["","nan"]:
                                _hal_html += f' <a href="{_jpdf}" target="_blank" style="background:#1a3a6b;color:#FFD700;padding:2px 10px;border-radius:12px;font-size:0.75rem;text-decoration:none;margin-right:6px;">📄 تحميل المحضر</a>'
                            else:
                                _hal_html += ' <span style="color:#94A3B8;font-size:0.75rem;">⏳ في انتظار رفع المحضر</span>'
                            _hal_html += '</div>'
                        elif _jhal == "مؤجلة":
                            _hal_html = f'<div style="margin-top:6px;"><span style="background:#F59E0B;color:#fff;padding:2px 10px;border-radius:12px;font-size:0.75rem;">🔄 مؤجلة{" — "+_jnote if _jnote else ""}</span></div>'
                        elif _jhal == "ملغاة":
                            _hal_html = '<div style="margin-top:6px;"><span style="background:#EF4444;color:#fff;padding:2px 10px;border-radius:12px;font-size:0.75rem;">❌ ملغاة</span></div>'

                        if has_date and j_published:
                            schedule_line = f'''<div style="display:flex;gap:16px;flex-wrap:wrap;margin-top:8px;">
                                <span style="color:#10B981;font-weight:700;font-size:0.85rem;">📅 {jdate}</span>
                                <span style="color:#94A3B8;font-size:0.85rem;">🕐 {jtime if has_time else "—"}</span>
                                <span style="color:#94A3B8;font-size:0.85rem;">🏛️ {jroom if has_room else "—"}</span>
                            </div>{_hal_html}'''
                        else:
                            schedule_line = '<div style="margin-top:8px;"><span style="color:#F59E0B;font-size:0.82rem;">⏳ لم يُحدد موعد المناقشة بعد</span></div>'

                        sup_html = f'<div style="color:#94A3B8;font-size:0.8rem;margin-top:4px;">👨‍🏫 المشرف: <span style="color:#CBD5E1;">{jsup}</span></div>' if jrole != "مشرف" and jsup and jsup not in ["","nan"] else ""

                        # أعضاء اللجنة
                        jury_members_html = ""
                        _prof_data_m = st.session_state.get("professor", {})
                        _is_massoul = str(_prof_data_m.get("مسؤول","")).strip() == "نعم"
                        _jrow_full2 = jury_memos[jury_memos["رقم المذكرة"].astype(str)==str(jmid)]
                        if len(_jrow_full2) > 0:
                            _jrf = _jrow_full2.iloc[0]
                            _members2 = []
                            _pres2 = str(_jrf.get("الرئيس","")).strip()
                            _ex12  = str(_jrf.get("المناقش1","")).strip()
                            _ex22  = str(_jrf.get("المناقش2","")).strip()
                            _sup2  = str(_jrf.get("الأستاذ","")).strip()
                            if jrole == "مشرف" or _is_massoul:
                                # مشرف أو مسؤول → اللجنة كاملة
                                if _pres2 and _pres2 not in ["","nan"]: _members2.append(f"🎓 رئيس: {_pres2}")
                                if _ex12  and _ex12  not in ["","nan"]: _members2.append(f"🔍 مناقش: {_ex12}")
                                if _ex22  and _ex22  not in ["","nan"]: _members2.append(f"🔍 مناقش: {_ex22}")
                            else:
                                # رئيس أو مناقش → المشرف فقط
                                if _sup2 and _sup2 not in ["","nan"]: _members2.append(f"👨‍🏫 مشرف: {_sup2}")
                            if _members2:
                                jury_members_html = '<div style="margin-top:6px;padding:6px 10px;background:rgba(201,162,39,0.1);border-radius:8px;font-size:0.75rem;color:#94A3B8;">' + " &nbsp;|&nbsp; ".join(_members2) + '</div>'

                        cards_html += f'''<div style="background:#1E293B;border:1px solid rgba(255,255,255,0.07);
                                    border-right:4px solid {r_color};border-radius:12px;
                                    padding:14px 16px;margin-bottom:10px;">
                            <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px;gap:8px;">
                                <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
                                    <span style="color:#FFD700;font-size:1rem;font-weight:900;">{jmid}</span>
                                    <span style="color:{r_color};background:rgba(255,255,255,0.05);padding:2px 10px;
                                        border-radius:20px;font-size:0.75rem;font-weight:700;
                                        border:1px solid {r_color};">{r_icon} {jrole}</span>
                                </div>
                                {preview_btn}
                            </div>
                            <div style="color:#F1F5F9;font-size:0.9rem;font-weight:700;line-height:1.5;">
                                {jtitle[:70]}{"..." if len(jtitle)>70 else ""}
                            </div>
                            {sup_html}
                            {jury_members_html}
                            {schedule_line}
                        </div>'''

                    import streamlit.components.v1 as _cv1
                    _cv1.html(
                        f'''<!DOCTYPE html><html dir="rtl"><body style="font-family:Arial,sans-serif;direction:rtl;background:transparent;margin:0;">{cards_html}</body></html>''',
                        height=min(900, len(filtered)*130 + 100),
                        scrolling=True
                    )

                    # ── موافقة المشرف على الإيداع النهائي ──
                    _massoul_val_ap = str(st.session_state.get("professor",{}).get("مسؤول","")).strip()
                    _is_dept_head = _massoul_val_ap in ["رئيس ق.ع","رئيس ق.خ","رئيس ق.ع.","رئيس ق.خ."]
                    _prof_dept = "القانون العام" if "ع" in _massoul_val_ap else ("القانون الخاص" if "خ" in _massoul_val_ap else "")

                    # موافقة رئيس القسم
                    if _is_dept_head:
                        for _, _jm_dh in filtered.iterrows():
                            _mid_dh = str(_jm_dh.get("رقم المذكرة","")).strip()
                            _hal_dh = str(_jm_dh.get("الحالة","")).strip() if "الحالة" in filtered.columns else ""
                            _aq_dh  = str(_jm_dh.get("رابط المذكرة النهائية","")).strip() if "رابط المذكرة النهائية" in filtered.columns else ""
                            _as_dh  = str(_jm_dh.get("موافقة المشرف","")).strip() if "موافقة المشرف" in filtered.columns else ""
                            _at_dh  = str(_jm_dh.get("موافقة رئيس القسم","")).strip() if "موافقة رئيس القسم" in filtered.columns else ""
                            _dept_dh = str(_jm_dh.get("القسم","")).strip() if "القسم" in filtered.columns else ""
                            if _hal_dh == "تمت المناقشة" and _as_dh == "نعم" and _aq_dh not in ["","nan"] and _prof_dept in _dept_dh:
                                st.markdown(f"**📋 موافقة رئيس القسم — مذكرة {_mid_dh}:**")
                                _c1_dh,_c2_dh = st.columns([3,1])
                                with _c1_dh:
                                    st.caption(f"{'✅ وافقت' if _at_dh=='نعم' else '⏳ في انتظار موافقتك كرئيس قسم'}")
                                with _c2_dh:
                                    if _at_dh != "نعم":
                                        if st.button("✅ أوافق", key=f"ap_dh_{_mid_dh}", use_container_width=True):
                                            _all_dh=load_memos(); _rn_dh=None
                                            for _ii_dh,(_, _rr_dh) in enumerate(_all_dh.iterrows()):
                                                _mk_dh=str(_rr_dh.get("رقم المذكرة","")).strip()
                                                try: _mk_dh=str(int(float(_mk_dh)))
                                                except: pass
                                                if _mk_dh==str(int(float(_mid_dh))): _rn_dh=_ii_dh+2; break
                                            if _rn_dh:
                                                sheets_service.spreadsheets().values().batchUpdate(
                                                    spreadsheetId=MEMOS_SHEET_ID,
                                                    body={"valueInputOption":"USER_ENTERED","data":[
                                                        {"range":f"Feuille 1!AT{_rn_dh}","values":[["نعم"]]}
                                                    ]}
                                                ).execute()
                                                clear_cache_and_reload()
                                                st.success(f"✅ موافقة رئيس القسم على مذكرة {_mid_dh}"); st.rerun()

                    if jrole == "مشرف":
                        for _, _jm_ap in filtered.iterrows():
                            _mid_ap = str(_jm_ap.get("رقم المذكرة","")).strip()
                            _hal_ap = str(_jm_ap.get("الحالة","")).strip() if "الحالة" in filtered.columns else ""
                            _aq_ap  = str(_jm_ap.get("رابط المذكرة النهائية","")).strip() if "رابط المذكرة النهائية" in filtered.columns else ""
                            _as_ap  = str(_jm_ap.get("موافقة المشرف","")).strip() if "موافقة المشرف" in filtered.columns else ""
                            if _hal_ap == "تمت المناقشة" and _aq_ap and _aq_ap not in ["","nan"]:
                                st.markdown(f"**📥 إيداع نهائي — مذكرة {_mid_ap}:**")
                                _c1_ap,_c2_ap,_c3_ap = st.columns([2,1,1])
                                with _c1_ap:
                                    st.caption(f"{'✅ وافقت' if _as_ap=='نعم' else '⏳ في انتظار موافقتك'}")
                                with _c2_ap:
                                    if _aq_ap not in ["","nan"]: st.markdown(f"[📄 المذكرة]({_aq_ap})")
                                with _c3_ap:
                                    if _as_ap != "نعم":
                                        _bc1,_bc2 = st.columns(2)
                                        with _bc1:
                                            if st.button("✅ أوافق", key=f"ap_sup_{_mid_ap}", use_container_width=True):
                                                _all_map=load_memos(); _rn_ap=None
                                                for _ii_ap,(_, _rr_ap) in enumerate(_all_map.iterrows()):
                                                    _mk_ap=str(_rr_ap.get("رقم المذكرة","")).strip()
                                                    try: _mk_ap=str(int(float(_mk_ap)))
                                                    except: pass
                                                    if _mk_ap==str(int(float(_mid_ap))): _rn_ap=_ii_ap+2; break
                                                if _rn_ap:
                                                    sheets_service.spreadsheets().values().batchUpdate(
                                                        spreadsheetId=MEMOS_SHEET_ID,
                                                        body={"valueInputOption":"USER_ENTERED","data":[
                                                            {"range":f"Feuille 1!AS{_rn_ap}","values":[["نعم"]]}
                                                        ]}
                                                    ).execute()
                                                    clear_cache_and_reload()
                                                    st.success(f"✅ تمت الموافقة على مذكرة {_mid_ap}")
                                                    st.rerun()
                                        with _bc2:
                                            if st.button("❌ أرفض", key=f"rj_sup_{_mid_ap}", use_container_width=True):
                                                _all_map2=load_memos(); _rn_ap2=None
                                                for _ii_ap2,(_, _rr_ap2) in enumerate(_all_map2.iterrows()):
                                                    _mk_ap2=str(_rr_ap2.get("رقم المذكرة","")).strip()
                                                    try: _mk_ap2=str(int(float(_mk_ap2)))
                                                    except: pass
                                                    if _mk_ap2==str(int(float(_mid_ap))): _rn_ap2=_ii_ap2+2; break
                                                if _rn_ap2:
                                                    sheets_service.spreadsheets().values().batchUpdate(
                                                        spreadsheetId=MEMOS_SHEET_ID,
                                                        body={"valueInputOption":"USER_ENTERED","data":[
                                                            {"range":f"Feuille 1!AS{_rn_ap2}","values":[["لا"]]},
                                                            {"range":f"Feuille 1!AQ{_rn_ap2}","values":[[""]]},
                                                            {"range":f"Feuille 1!AR{_rn_ap2}","values":[[""]]}
                                                        ]}
                                                    ).execute()
                                                    clear_cache_and_reload()
                                                    st.warning(f"❌ رُفض إيداع مذكرة {_mid_ap} — سيُعاد فتح الرفع للطالب")
                                                    st.rerun()

                    # تصدير HTML
                    st.markdown("---")
                    rows_html = ""
                    # فلتر نشر البرنامج — HTML يظهر فقط المذكرات المنشورة
                    _col_ai_h = next((c for c in ["نشر البرنامج","AI","منشور"] if c in jury_memos.columns), None)
                    if _col_ai_h:
                        _all_for_html = jury_memos[jury_memos[_col_ai_h].astype(str).str.strip() == "نعم"].copy()
                    else:
                        _all_for_html = pd.DataFrame()  # لا نشر → لا HTML
                    for idx_r, jm_r in _all_for_html.iterrows():
                        jmid_r  = str(jm_r.get("رقم المذكرة","")).strip()
                        jtitle_r= str(jm_r.get("عنوان المذكرة","")).strip()
                        jrole_r = str(jm_r.get("الصفة","")).strip()
                        jdate_r = str(jm_r.get("تاريخ المناقشة","")).strip()
                        jtime_r = str(jm_r.get("توقيت المناقشة","")).strip()
                        jroom_r = str(jm_r.get("القاعة","")).strip()
                        jdate_r = jdate_r if jdate_r not in ["","nan"] else "—"
                        jtime_r = jtime_r if jtime_r not in ["","nan"] else "—"
                        jroom_r = jroom_r if jroom_r not in ["","nan"] else "—"
                        jlink_r    = str(jm_r.get("رابط الملف","")).strip()
                        jdeposit_r = str(jm_r.get("حالة الإيداع","")).strip()
                        jrole_r_v  = str(jm_r.get("الصفة","")).strip()
                        is_sup_r   = (jrole_r_v == "مشرف")
                        if not jdeposit_r or jdeposit_r in ["","nan"]:
                            link_cell = "⚠️ لم يتم الإيداع"
                        elif jdeposit_r == "قابلة للمناقشة" and jlink_r and jlink_r not in ["","nan"]:
                            link_cell = f'<a href="{jlink_r}" target="_blank" style="color:#2F6F7E;font-weight:700;text-decoration:none;">👁️ معاينة</a>'
                        elif is_sup_r and jlink_r and jlink_r not in ["","nan"]:
                            link_cell = f'<a href="{jlink_r}" target="_blank" style="color:#2F6F7E;font-weight:700;text-decoration:none;">👁️ معاينة</a>'
                        else:
                            link_cell = "⏳ في انتظار موافقة المشرف"
                        bg = "#f8fafc" if idx_r % 2 == 0 else "#ffffff"
                        rows_html += f'''<tr style="background:{bg};">
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;font-weight:700;color:#0F2942;">{jmid_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:right;">{jtitle_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;color:#2F6F7E;font-weight:600;">{jrole_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;color:#059669;font-weight:600;">{jdate_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;">{jtime_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;">{jroom_r}</td>
                            <td style="padding:10px 14px;border:1px solid #e2e8f0;text-align:center;">{link_cell}</td>
                        </tr>'''

                    html_export = f'''<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>برنامج مناقشة مذكرات الماستر 2025-2026</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;900&display=swap');
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family:'Cairo', Arial, sans-serif; direction:rtl; background:#fff; color:#1e293b; padding:30px; }}
  .header {{ text-align:center; margin-bottom:30px; padding-bottom:20px; border-bottom:3px solid #0F2942; }}
  .univ {{ font-size:14px; font-weight:600; color:#64748b; margin-bottom:4px; }}
  .faculty {{ font-size:13px; color:#94a3b8; margin-bottom:16px; }}
  .title {{ font-size:22px; font-weight:900; color:#0F2942; margin-bottom:8px; }}
  .prof-name {{ font-size:14px; color:#2F6F7E; font-weight:600; }}
  .meta {{ font-size:12px; color:#94a3b8; margin-top:6px; }}
  table {{ width:100%; border-collapse:collapse; margin-top:20px; }}
  thead tr {{ background:#0F2942; color:#ffffff; }}
  thead th {{ padding:12px 14px; text-align:center; font-size:13px; font-weight:700; border:1px solid #1e3a5c; }}
  thead th:nth-child(2) {{ text-align:right; }}
  tbody td {{ font-size:12px; color:#1e293b; }}
  .footer {{ text-align:center; margin-top:30px; font-size:11px; color:#94a3b8; border-top:1px solid #e2e8f0; padding-top:16px; }}
  @media print {{
    body {{ padding:15px; }}
    .no-print {{ display:none; }}
  }}
</style>
</head>
<body>
<div class="header">
  <div class="univ">🎓 جامعة محمد البشير الإبراهيمي — برج بوعريريج</div>
  <div class="faculty">كلية الحقوق والعلوم السياسية</div>
  <div class="title">برنامج مناقشة مذكرات الماستر 2025-2026</div>
  <div class="prof-name">الأستاذ(ة): {prof_name}</div>
  <div class="meta">عدد المذكرات: {len(jury_memos)} | التاريخ: {datetime.now().strftime("%Y-%m-%d %H:%M")}</div>
</div>

<table>
  <thead>
    <tr>
      <th style="width:60px;">رقم</th>
      <th style="text-align:right;">عنوان المذكرة</th>
      <th style="width:110px;">الصفة</th>
      <th style="width:120px;">تاريخ المناقشة</th>
      <th style="width:80px;">التوقيت</th>
      <th style="width:80px;">القاعة</th>
      <th style="width:80px;">المذكرة</th>
    </tr>
  </thead>
  <tbody>
    {rows_html}
  </tbody>
</table>

<div class="footer">
  <p></p>
</div>
</body>
</html>'''

                    if not _all_for_html.empty:
                        import base64 as _b64
                        _html_b64 = _b64.b64encode(html_export.encode("utf-8")).decode()
                        _fname = f"programme_{prof_name.replace(' ','_')}.html"
                        st.markdown(f'''<div style="text-align:center;margin-top:12px;">
                        <a href="data:text/html;base64,{_html_b64}" download="{_fname}"
                           style="display:inline-block;background:linear-gradient(135deg,#0F2942,#1A3A5C);
                                  color:#FFD700;padding:8px 22px;border-radius:20px;
                                  text-decoration:none;font-weight:700;font-size:0.88rem;
                                  border:1px solid rgba(255,215,0,0.4);">
                            📄 تحميل البرنامج
                        </a>
                    </div>''', unsafe_allow_html=True)
                    else:
                        st.info("⏳ لم يُنشر البرنامج بعد — التحميل غير متاح")



# ================================================================
# فضاء الإدارة
# ================================================================
elif st.session_state.user_type == "admin":
    if not st.session_state.logged_in:
        render_countdown_banner()
        c1,c2=st.columns([4,1])
        with c2:
            if st.button("رجوع"): st.session_state.user_type=None; st.rerun()
        st.markdown("<h2>⚙️ فضاء الإدارة</h2>", unsafe_allow_html=True)
        with st.form("admin_login"):
            u=st.text_input("اسم المستخدم"); p=st.text_input("كلمة المرور",type="password")
            if st.form_submit_button("دخول"):
                # تحقق من فضاء المحاضر أولاً
                if u in PRINTER_CREDENTIALS and PRINTER_CREDENTIALS[u] == p:
                    st.session_state.admin_user = u
                    st.session_state.logged_in = True
                    st.session_state.is_printer = True
                    st.session_state.is_library = False
                    st.query_params['ut'] = 'admin'
                    st.rerun()
                elif u in LIBRARY_CREDENTIALS and LIBRARY_CREDENTIALS[u] == p:
                    st.session_state.admin_user = u
                    st.session_state.logged_in = True
                    st.session_state.is_library = True
                    st.session_state.is_printer = False
                    st.query_params['ut'] = 'admin'
                    st.rerun()
                else:
                    st.session_state.is_printer = False
                    st.session_state.is_library = False
                v,r=verify_admin(u,p)
                if not v: st.error(r)
                else:
                    st.session_state.admin_user=r; st.session_state.logged_in=True
                    st.query_params['ut']='admin'; st.query_params['un']=encode_str(st.session_state.admin_user); st.rerun()
    else:
        render_countdown_banner()
        c1,c2=st.columns([4,1])
        with c2:
            if st.button("خروج"): logout()
        if not st.session_state.get("is_library", False):
            st.header("📊 لوحة تحكم الإدارة")
    
            st.markdown("<br>", unsafe_allow_html=True)
            df_prof_memos = load_prof_memos()
            df_students = load_students()
            df_memos = load_memos()
            df_students = load_students()
            st_s=len(df_students); t_m=len(df_memos); r_m=len(df_memos[df_memos["تم التسجيل"].astype(str).str.strip()=="نعم"])
            df_prof_memos = load_prof_memos()
            a_m=t_m-r_m; t_p=len(df_prof_memos["الأستاذ"].unique()) if "الأستاذ" in df_prof_memos.columns else 0
            memo_col=df_students["رقم المذكرة"].astype(str).str.strip() if "رقم المذكرة" in df_students.columns else pd.Series(dtype=str)
            reg_st=(memo_col!="").sum(); unreg_st=(memo_col=="").sum()
            st.markdown('<div class="kpi-grid">', unsafe_allow_html=True)
            # إحصائيات المناقشات
            if not st.session_state.get("is_library", False):
                df_kpi = load_memos()
                total_memos = len(df_kpi)
                def _has_jury(r):
                    def f(v): return str(v).strip() not in ["","nan","—","None"]
                    return f(r.get("الرئيس","")) and f(r.get("المناقش1","")) and f(r.get("المناقش2",""))
                has_jury_mask = df_kpi.apply(_has_jury, axis=1)
                no_jury = (~has_jury_mask).sum()
                scheduled = df_kpi["تاريخ المناقشة"].astype(str).str.strip().apply(lambda x: x not in ["","nan"]).sum() if "تاريخ المناقشة" in df_kpi.columns else 0
                defended = df_kpi[df_kpi.get("حالة المناقشة", pd.Series(dtype=str)).astype(str).str.strip()=="تمت"].shape[0] if "حالة المناقشة" in df_kpi.columns else 0
                st.markdown(f'''<div class="kpi-grid">
                    <div class="kpi-card"><div class="kpi-value">{total_memos}</div><div class="kpi-label">📚 إجمالي المذكرات</div></div>
                    <div class="kpi-card" style="border-top:3px solid #10B981;"><div class="kpi-value" style="color:#10B981;">{int(scheduled)}</div><div class="kpi-label">📅 مبرمجة</div></div>
                    <div class="kpi-card" style="border-top:3px solid #EF4444;"><div class="kpi-value" style="color:#EF4444;">{int(no_jury)}</div><div class="kpi-label">⚠️ ناقصة اللجان</div></div>
                    <div class="kpi-card" style="border-top:3px solid #FFD700;"><div class="kpi-value" style="color:#FFD700;">{total_memos - int(scheduled)}</div><div class="kpi-label">⏳ غير مبرمجة</div></div>
                </div>''', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
        _is_printer_user = st.session_state.get("is_printer", False)
        _is_library_user = st.session_state.get("is_library", False)
        if _is_library_user:
            st.markdown('<h2 style="color:#FFD700;">📚 فضاء المكتبة</h2>', unsafe_allow_html=True)
        if _is_printer_user:
            tab_mahdar, = st.tabs(["📄 المحاضر"])
            tab_takleef = None; tab_archive = None; tab_siyar = None; tab_seq = None; tab_stud = None
        elif _is_library_user:
            tab_maktaba, = st.tabs(["📚 الإيداع النهائي"])
            tab_takleef = None; tab_mahdar = None; tab_archive = None; tab_siyar = None; tab_seq = None; tab_stud = None
        else:
            tab_takleef,tab_mahdar,tab_siyar,tab_seq,tab_stud,tab_archive=st.tabs(["📋 التكاليف والتحقق","📄 المحاضر","📊 سير المناقشات","📥 استيراد أرقام المحاضر","👥 استيراد أرقام ملفات الطلبة","🗂️ أرشيف (جدولة ذكية)"])
            tab_maktaba = None

        if _is_library_user and tab_maktaba:
         with tab_maktaba:
            st.subheader("📚 الإيداع النهائي — مسؤول المكتبة")
            _df_lib = load_memos()
            _lib_col_as = "موافقة المشرف" if "موافقة المشرف" in _df_lib.columns else None
            _lib_col_aq = "رابط المذكرة النهائية" if "رابط المذكرة النهائية" in _df_lib.columns else None
            _lib_col_ar = "رابط الملخص النهائي" if "رابط الملخص النهائي" in _df_lib.columns else None
            _lib_col_au = "تبرئة المكتبة" if "تبرئة المكتبة" in _df_lib.columns else None

            if _lib_col_as and _lib_col_aq:
                _lib_memos = _df_lib[
                    (_df_lib[_lib_col_as].astype(str).str.strip() == "نعم") &
                    (_df_lib[_lib_col_aq].astype(str).str.strip().apply(lambda x: x not in ["","nan"]))
                ].copy()
            else:
                _lib_memos = pd.DataFrame()

            if _lib_memos.empty:
                st.info("⏳ لا توجد مذكرات في انتظار تبرئة المكتبة بعد")
            else:
                st.success(f"📋 {len(_lib_memos)} مذكرة في انتظار المراجعة")
                for _, _lrow in _lib_memos.iterrows():
                    _lmid = str(_lrow.get("رقم المذكرة","")).strip()
                    _ltitle = str(_lrow.get("عنوان المذكرة","")).strip()[:60]
                    _ls1 = str(_lrow.get("الطالب الأول","")).strip()
                    _laq = str(_lrow.get(_lib_col_aq,"")).strip() if _lib_col_aq else ""
                    _lar = str(_lrow.get(_lib_col_ar,"")).strip() if _lib_col_ar else ""
                    _lau = str(_lrow.get(_lib_col_au,"")).strip() if _lib_col_au else ""
                    with st.expander(f"مذكرة {_lmid} — {_ls1}"):
                        st.markdown(f"**{_ltitle}**")
                        _lc1,_lc2 = st.columns(2)
                        with _lc1:
                            if _laq not in ["","nan"]: st.markdown(f"[📄 تحميل المذكرة النهائية]({_laq})")
                        with _lc2:
                            if _lar not in ["","nan"]: st.markdown(f"[📝 تحميل الملخص]({_lar})")
                        if _lau == "نعم":
                            st.success("✅ تمت تبرئة المكتبة")
                        else:
                            if st.button("✅ منح تبرئة المكتبة", type="primary", key=f"lib_{_lmid}", use_container_width=True):
                                with st.spinner("⏳ جاري الحفظ..."):
                                    _all_lib=load_memos(); _rn_lib=None
                                    for _il,(_, _rl) in enumerate(_all_lib.iterrows()):
                                        _mkl=str(_rl.get("رقم المذكرة","")).strip()
                                        try: _mkl=str(int(float(_mkl)))
                                        except: pass
                                        if _mkl==str(int(float(_lmid))): _rn_lib=_il+2; break
                                    if _rn_lib:
                                        sheets_service.spreadsheets().values().batchUpdate(
                                            spreadsheetId=MEMOS_SHEET_ID,
                                            body={"valueInputOption":"USER_ENTERED","data":[
                                                {"range":f"Feuille 1!AU{_rn_lib}","values":[["نعم"]]}
                                            ]}
                                        ).execute()
                                        clear_cache_and_reload()
                                        st.success(f"✅ تبرئة المكتبة لمذكرة {_lmid}"); st.rerun()


        # ================================================================
        # TAB التكاليف والتحقق
        # ================================================================
        if not _is_printer_user and not _is_library_user and tab_takleef:
         with tab_takleef:
            st.subheader("📋 التكاليف والتحقق من التعارضات")

            df_m_tk = load_memos()
            df_p_tk = load_prof_memos()

            _col_ai_tk = next((c for c in ["نشر البرنامج","AI"] if c in df_m_tk.columns), None)
            _col_w = "تاريخ المناقشة"; _col_x = "توقيت المناقشة"; _col_y = "القاعة"
            _col_w_tk = _col_w; _col_x_tk = _col_x; _col_y_tk = _col_y

            sched_tk = df_m_tk[
                df_m_tk[_col_w].astype(str).str.strip().apply(lambda v: v not in ["","nan"])
            ].copy()

            if sched_tk.empty:
                st.warning("⚠️ لا توجد مذكرات مبرمجة في الشيت بعد.")
            else:
                st.success(f"✅ {len(sched_tk)} مذكرة مبرمجة")

                st.markdown("---")
                st.markdown("### 🔍 التحقق من التعارضات")
                _room_slot = {}; _prof_slot = {}; _prof_day = {}
                _conflicts = []

                for _, row in sched_tk.iterrows():
                    _mid = str(row.get("رقم المذكرة","")).strip()
                    _day = str(row.get(_col_w,"")).strip()
                    _slot = str(row.get(_col_x,"")).strip()
                    _room = str(row.get(_col_y,"")).strip()
                    _profs = [str(row.get(c,"")).strip() for c in ["الأستاذ","الرئيس","المناقش1","المناقش2"]
                              if str(row.get(c,"")).strip() not in ["","nan"]]

                    _key_r = (_day, _slot, _room)
                    if _room and _key_r in _room_slot:
                        _conflicts.append(f"🔴 تعارض قاعة: **{_room}** | {_day} {_slot} | مذكرة {_mid} ↔ {_room_slot[_key_r]}")
                    elif _room: _room_slot[_key_r] = _mid

                    for _p in _profs:
                        _key_p = (_day, _slot, _p)
                        if _key_p in _prof_slot:
                            _conflicts.append(f"🔴 تعارض أستاذ: **{_p}** | {_day} {_slot} | مذكرة {_mid} ↔ {_prof_slot[_key_p]}")
                        else: _prof_slot[_key_p] = _mid
                        _prof_day[(_p,_day)] = _prof_day.get((_p,_day),0) + 1
                        if _prof_day[(_p,_day)] > 3:
                            _conflicts.append(f"🟡 تجاوز 3/يوم: **{_p}** | {_day} ({_prof_day[(_p,_day)]} مناقشات)")

                if _conflicts:
                    st.error(f"⚠️ {len(_conflicts)} تعارض:")
                    for _c in _conflicts: st.markdown(f"- {_c}")
                else:
                    st.success("✅ لا تعارضات — الجدول سليم 100%")

                st.markdown("---")
                st.markdown("### 📧 إرسال التكاليف للأساتذة")

                _email_dict = {}
                if not df_p_tk.empty:
                    _email_col = "الإيميل" if "الإيميل" in df_p_tk.columns else "البريد الإلكتروني"
                    _name_col  = "الأستاذ" if "الأستاذ" in df_p_tk.columns else None
                    if _name_col and _email_col in df_p_tk.columns:
                        for _, _pr in df_p_tk.iterrows():
                            _pn = str(_pr.get(_name_col,"")).strip()
                            _pe = str(_pr.get(_email_col,"")).strip()
                            if _pn and _pe and "@" in _pe: _email_dict[_pn] = _pe

                st.info(f"📬 {len(_email_dict)} أستاذ لديه بريد إلكتروني")

                _all_profs_tk = sorted(set(
                    str(row.get(c,"")).strip()
                    for _, row in sched_tk.iterrows()
                    for c in ["الأستاذ","الرئيس","المناقش1","المناقش2"]
                    if str(row.get(c,"")).strip() not in ["","nan"]
                ))

                _send_mode = st.radio("إرسال إلى:", ["أستاذ واحد","الكل"], horizontal=True, key="send_mode_tk")

                if _send_mode == "أستاذ واحد":
                    _sel_prof = st.selectbox("اختر الأستاذ:", _all_profs_tk, key="sel_prof_tk")
                    _profs_to_send = [_sel_prof]
                else:
                    _profs_to_send = [p for p in _all_profs_tk if p in _email_dict]
                    st.caption(f"سيُرسل لـ {len(_profs_to_send)} أستاذ")

                _send_fmt = st.radio("صيغة الإرسال:", ["📧 HTML فقط", "📄 PDF + HTML"], horizontal=True, key="send_fmt_tk")
                st.session_state["_send_fmt_val"] = _send_fmt

                if st.button("📧 إرسال التكاليف", type="primary", use_container_width=True, key="send_takleef"):
                    _sent = 0; _failed = []
                    _progress = st.progress(0)
                    _status = st.empty()

                    for _pi, _prof_name in enumerate(_profs_to_send):
                        _email_addr = _email_dict.get(_prof_name,"")
                        if not _email_addr:
                            _failed.append(f"{_prof_name} (لا بريد)")
                            continue

                        _prof_rows = []
                        for _, _row in sched_tk.iterrows():
                            _profs_r = [str(_row.get(c,"")).strip() for c in ["الأستاذ","الرئيس","المناقش1","المناقش2"]]
                            if _prof_name in _profs_r:
                                _role = next((c for c in ["الأستاذ","الرئيس","المناقش1","المناقش2"] if str(_row.get(c,"")).strip()==_prof_name), "")
                                _role_ar = {"الأستاذ":"مشرفاً","الرئيس":"رئيساً","المناقش1":"ممتحناً","المناقش2":"ممتحناً"}.get(_role,"عضواً")
                                _prof_rows.append({
                                    "رقم": str(_row.get("رقم المذكرة","")),
                                    "عنوان": str(_row.get("عنوان المذكرة","")),
                                    "يوم": str(_row.get(_col_w,"")),
                                    "توقيت": str(_row.get(_col_x,"")),
                                    "قاعة": str(_row.get(_col_y,"")),
                                    "صفة": _role_ar,
                                    "رابط": str(_row.get("رابط الملف","")).strip()
                                })
                        _prof_rows.sort(key=lambda r: (r["يوم"], r["توقيت"]))

                        _html_body = (
                            "<!DOCTYPE html>"
                            "<html dir='rtl' lang='ar'><head><meta charset='UTF-8'>"
                            "<meta name='viewport' content='width=device-width,initial-scale=1.0,maximum-scale=3.0'>"
                            "<style>"
                            "*{box-sizing:border-box;margin:0;padding:0}"
                            "body{font-family:Arial,sans-serif;direction:rtl;background:#f0f4f8;padding:8px}"
                            ".wrap{max-width:680px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 4px 20px rgba(0,0,0,.1)}"
                            ".hdr{background:linear-gradient(135deg,#0F2942,#1a3a6b);padding:20px;text-align:center}"
                            ".hdr h1{color:#FFD700;font-size:1.1rem;margin-bottom:5px}"
                            ".hdr p{color:#cdd8e8;font-size:.8rem;margin-top:3px}"
                            ".bdy{padding:20px 18px}"
                            ".greet{font-size:.95rem;color:#0F2942;font-weight:bold;margin-bottom:10px}"
                            ".txt{font-size:.88rem;color:#333;line-height:1.8;margin-bottom:8px}"
                            ".warn{background:#fff8e1;border-right:4px solid #f59e0b;padding:10px 12px;border-radius:6px;font-size:.83rem;color:#78350f;margin:14px 0}"
                            ".plat{background:#e8f0fe;border-radius:8px;padding:10px 14px;margin-top:16px;text-align:center;font-size:.88rem}"
                            ".plat a{color:#0F2942;font-weight:bold;text-decoration:none}"
                            ".sig{margin-top:18px;font-size:.85rem;color:#1a3a6b;line-height:1.8}"
                            ".ftr{background:#f4f6f9;padding:12px;text-align:center;font-size:.75rem;color:#888;border-top:1px solid #e0e0e0}"
                            "table.dt{width:100%;border-collapse:collapse;margin-top:14px;font-size:.8rem}"
                            "table.dt th{background:#0F2942;color:#FFD700;padding:9px 6px;text-align:center}"
                            "table.dt td{border:1px solid #e0e0e0;padding:7px 6px;text-align:center;color:#333}"
                            "table.dt tr:nth-child(even) td{background:#f9fafc}"
                            "table.dt .ttl{text-align:center;min-width:140px;font-size:.75rem}"
                            ".cards{display:none}"
                            ".card{background:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;padding:14px;margin-bottom:12px;border-right:4px solid #0F2942}"
                            ".card .cn{font-size:.72rem;color:#64748b;margin-bottom:5px}"
                            ".card .ct{font-size:.88rem;font-weight:bold;color:#0F2942;margin-bottom:9px;line-height:1.5}"
                            ".card .cr{display:flex;justify-content:space-between;font-size:.8rem;margin-bottom:4px;color:#333}"
                            ".card .cl{color:#64748b}"
                            ".card .role{display:inline-block;background:#0F2942;color:#FFD700;padding:2px 9px;border-radius:10px;font-size:.72rem;margin-top:7px}"
                            ".card .lnk{display:block;text-align:center;margin-top:10px;background:#1a3a6b;color:#FFD700;padding:7px;border-radius:7px;text-decoration:none;font-size:.83rem}"
                            "@media(max-width:550px){table.dt{display:none}.cards{display:block;margin-top:14px}}"
                            "</style></head><body>"
                            "<div class='wrap'>"
                            "<div class='hdr'><h1>&#128203; برنامج مناقشة مذكرات الماستر</h1>"
                            "<p>الدورة العادية — السنة الجامعية 2025–2026</p>"
                            "<p>كلية الحقوق والعلوم السياسية — جامعة محمد البشير الإبراهيمي، برج بوعريريج</p></div>"
                            "<div class='bdy'>"
                            f"<p class='greet'>الأستاذ(ة) الفاضل(ة): {_prof_name}</p>"
                            "<p class='txt'>تحية طيبة وبعد،<br>مرفق لكم البرنامج الرسمي لمناقشة مذكرات الماستر، للدورة العادية للسنة الجامعية 2025–2026.</p>"
                            "<p class='txt'>يمكنكم معاينة المذكرات وتحميلها مباشرةً من خلال الضغط على &#128065; المرفق بكل مذكرة.</p>"
                            "<div class='warn'>&#9888;&#65039; نظرًا لدقة رزنامة نهاية السنة الجامعية، <strong>يُمنع تأجيل المناقشات أو تعديل توقيتها</strong>. في حال وجود أي ملاحظات، يُرجى التواصل مع مكتب فريق التكوين.</div>"
                            "<table class='dt'><thead><tr><th>#</th><th>&#128065;</th><th>رقم</th><th>العنوان</th><th>التاريخ</th><th>التوقيت</th><th>القاعة</th><th>الصفة</th></tr></thead><tbody>"
                            + "".join(
                                "<tr><td>" + str(idx_r+1) + "</td><td>" +
                                ('<a href="' + r["رابط"] + '" target="_blank">&#128065;</a>' if r['رابط'] and r['رابط'] not in ['','nan'] else '—') +
                                "</td><td>" + r['رقم'] + "</td><td class='ttl'>" + r['عنوان'] + "</td>" +
                                "<td>" + r['يوم'] + "</td><td>" + r['توقيت'] + "</td><td>" + r['قاعة'] + "</td><td>" + r['صفة'] + "</td></tr>"
                                for idx_r,r in enumerate(_prof_rows))
                            + "</tbody></table>"
                            + "<div class='cards'>"
                            + "".join(
                                f"<div class='card'><div class='cn'>مذكرة #{r['رقم']}</div>"
                                f"<div class='ct'>{r['عنوان']}</div>"
                                f"<div class='cr'><span class='cl'>&#128197;</span><span>{r['يوم']}</span></div>"
                                f"<div class='cr'><span class='cl'>&#128336;</span><span>{r['توقيت']}</span></div>"
                                f"<div class='cr'><span class='cl'>&#127963;&#65039;</span><span>{r['قاعة']}</span></div>"
                                f"<span class='role'>{r['صفة']}</span>"
                                + (f"<a class='lnk' href='{r['رابط']}' target='_blank'>&#128065; معاينة المذكرة</a>" if r['رابط'] and r['رابط'] not in ['','nan'] else "")
                                + "</div>"
                                for r in _prof_rows)
                            + "</div>"
                            "<div class='plat'>&#127760; <a href='https://memoires2026.streamlit.app'>memoires2026.streamlit.app</a></div>"
                            "<div class='sig'><p>عيدكم مبارك، وكل عام وأنتم بخير. &#127769;</p>"
                            "<p><strong>مسؤول الميدان — البروفيسور رفاف لخضر</strong></p></div>"
                            "</div><div class='ftr'>هذا البريد أُرسل تلقائيًا من منصة إدارة مذكرات الماستر — جامعة برج بوعريريج 2025/2026</div>"
                            "</div></body></html>"
                        )

                        try:
                            import smtplib
                            from email.mime.multipart import MIMEMultipart
                            from email.mime.text import MIMEText
                            from email.mime.base import MIMEBase
                            from email import encoders as _enc
                            _rows_em = df_p_tk[df_p_tk["الأستاذ"].astype(str).str.strip()==_prof_name.strip()]
                            _email_to = get_email_smart(_rows_em.iloc[0]) if not _rows_em.empty else ""
                            if not _email_to or "@" not in _email_to:
                                _ok, _msg = False, "لا بريد"
                            else:
                                _msg_em = MIMEMultipart("mixed")
                                _msg_em["Subject"] = "تكليف بمناقشة مذكرات الماستر — الدورة العادية الأولى 2025/2026"
                                _msg_em["From"] = EMAIL_SENDER
                                _msg_em["To"] = _email_to
                                _alt = MIMEMultipart("alternative")
                                _alt.attach(MIMEText(_html_body, "html", "utf-8"))
                                _msg_em.attach(_alt)
                                _fmt_val = st.session_state.get("_send_fmt_val","")
                                if _fmt_val == "📄 PDF + HTML":
                                    try:
                                        from weasyprint import HTML as _WH
                                        _pdf_bytes = _WH(string=_html_body).write_pdf()
                                        _pdf_part = MIMEBase("application","pdf")
                                        _pdf_part.set_payload(_pdf_bytes)
                                        _enc.encode_base64(_pdf_part)
                                        _pdf_name = f"تكليف_{_prof_name.replace(' ','_')}.pdf"
                                        _pdf_part.add_header("Content-Disposition","attachment",filename=_pdf_name)
                                        _msg_em.attach(_pdf_part)
                                        _ok = True; _msg = "تم + PDF"
                                    except Exception as _pdf_err:
                                        _failed.append(f"PDF فشل: {_pdf_err}")
                                with smtplib.SMTP_SSL("smtp.gmail.com", 465) as _srv:
                                    _srv.login(EMAIL_SENDER, EMAIL_PASSWORD)
                                    _srv.sendmail(EMAIL_SENDER, _email_to, _msg_em.as_string())
                                _ok, _msg = True, "تم"
                        except Exception as _ex_em:
                            _ok, _msg = False, str(_ex_em)

                        if _ok: _sent += 1
                        else: _failed.append(f"{_prof_name}: {_msg}")

                        _progress.progress(int((_pi+1)/len(_profs_to_send)*100))
                        _status.text(f"جاري الإرسال... {_pi+1}/{len(_profs_to_send)}")

                    _progress.empty(); _status.empty()
                    if _sent > 0: st.success(f"✅ تم إرسال {_sent} تكليف بنجاح")
                    if _failed: st.warning("⚠️ فشل: " + " | ".join(_failed[:5]))

                st.markdown("---")
                st.markdown("### 📨 إرسال تكليف مذكرة واحدة")
                st.info("اختر مذكرة محددة وأرسل تكليفاً لجميع أعضاء اللجنة والطلبة المعنيين")

                try:
                    _sorted_tk1 = sched_tk.copy()
                    _sorted_tk1["_num"] = pd.to_numeric(_sorted_tk1["رقم المذكرة"], errors="coerce")
                    _sorted_tk1 = _sorted_tk1.sort_values("_num").drop(columns=["_num"])
                except:
                    _sorted_tk1 = sched_tk.copy()

                def _fmt_tk1(r):
                    num = str(r.get("رقم المذكرة","")).strip()
                    s1 = str(r.get("الطالب الأول","")).strip()
                    prof = str(r.get("الأستاذ","")).strip()
                    return f"{num.zfill(3)} | {prof} | {s1}"

                _opts_tk1 = [_fmt_tk1(r) for _, r in _sorted_tk1.iterrows()]
                _mids_tk1 = _sorted_tk1["رقم المذكرة"].astype(str).tolist()
                _sel_tk1 = st.selectbox("اختر المذكرة:", _opts_tk1, index=None, placeholder="اكتب للبحث...", key="sel_tk1_memo")
                _mid_tk1 = _mids_tk1[_opts_tk1.index(_sel_tk1)] if _sel_tk1 else None

                if _mid_tk1:
                    _row_tk1 = _sorted_tk1[_sorted_tk1["رقم المذكرة"].astype(str)==str(_mid_tk1)].iloc[0]
                    _title_tk1 = str(_row_tk1.get("عنوان المذكرة","")).strip()
                    _sup_tk1   = str(_row_tk1.get("الأستاذ","")).strip()
                    _pres_tk1  = str(_row_tk1.get("الرئيس","")).strip()
                    _ex1_tk1   = str(_row_tk1.get("المناقش1","")).strip()
                    _ex2_tk1   = str(_row_tk1.get("المناقش2","")).strip()
                    _s1_tk1    = str(_row_tk1.get("الطالب الأول","")).strip()
                    _s2_tk1    = str(_row_tk1.get("الطالب الثاني","")).strip()
                    _date_tk1  = str(_row_tk1.get(_col_w_tk,"")).strip()
                    _slot_tk1  = str(_row_tk1.get(_col_x_tk,"")).strip()
                    _room_tk1  = str(_row_tk1.get(_col_y_tk,"")).strip()
                    _link_tk1  = str(_row_tk1.get("رابط الملف","")).strip()
                    _students_label = "الطالب(ة)" if not _s2_tk1 or _s2_tk1 in ["","nan"] else "الطالبَين"
                    _students_names = _s1_tk1 + (f" — {_s2_tk1}" if _s2_tk1 and _s2_tk1 not in ["","nan"] else "")

                    def _build_email_tk1(prof_name, role_ar):
                        _link_btn = f'<a href="{_link_tk1}" style="color:#1a3a6b;font-weight:bold;">👁️ معاينة المذكرة</a>' if _link_tk1 and _link_tk1 not in ["","nan"] else ""
                        return f"""<!DOCTYPE html><html dir="rtl" lang="ar">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<style>body{{font-family:Arial,sans-serif;direction:rtl;background:#f0f4f8;padding:8px}}
.wrap{{max-width:640px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 4px 20px rgba(0,0,0,.1)}}
.hdr{{background:linear-gradient(135deg,#0F2942,#1a3a6b);padding:20px;text-align:center;color:#FFD700;font-size:1.1rem}}
.bdy{{padding:22px 20px;font-size:0.93rem;color:#333;line-height:1.9}}
.box{{background:#f8fafc;border-right:4px solid #0F2942;padding:12px 16px;border-radius:8px;margin:14px 0}}
.role{{display:inline-block;background:#0F2942;color:#FFD700;padding:3px 12px;border-radius:12px;font-size:0.82rem}}
.ftr{{background:#f4f6f9;padding:12px;text-align:center;font-size:0.78rem;color:#777;border-top:1px solid #e0e0e0}}
</style></head><body>
<div class="wrap">
<div class="hdr">📋 برمجة مناقشة مذكرة الماستر</div>
<div class="bdy">
<p>الأستاذ(ة) الفاضل(ة) <strong>{prof_name}</strong>،</p>
<p>تحية طيبة وبعد،</p>
<p>نُحيطكم علمًا بأنه تمّت برمجة مناقشة مذكرة الماستر رقم <strong>{_mid_tk1}</strong>، الموسومة بـ:</p>
<div class="box"><strong>« {_title_tk1} »</strong></div>
<p>من إعداد {_students_label}: <strong>{_students_names}</strong><br>
وتحت إشراف الأستاذ(ة): <strong>{_sup_tk1}</strong></p>
<p>وذلك يوم <strong>{_date_tk1}</strong> على الساعة <strong>{_slot_tk1}</strong> بقاعة <strong>{_room_tk1}</strong>،
بصفتكم <span class="role">{role_ar}</span></p>
{f'<p style="margin-top:14px;">{_link_btn}</p>' if _link_btn else ""}
<p style="margin-top:12px;">🌐 منصة المذكرات: <a href="https://memoires2026.streamlit.app">memoires2026.streamlit.app</a></p>
<p style="margin-top:18px;color:#1a3a6b;">تقبلوا تحياتي،<br><strong>مسؤول الميدان — البروفيسور رفاف لخضر</strong></p>
</div>
<div class="ftr">كلية الحقوق والعلوم السياسية — جامعة محمد البشير الإبراهيمي، برج بوعريريج</div>
</div></body></html>"""

                    def _build_student_email_tk1(student_name):
                        return f"""<!DOCTYPE html><html dir="rtl" lang="ar">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<style>body{{font-family:Arial,sans-serif;direction:rtl;background:#f0f4f8;padding:8px}}
.wrap{{max-width:640px;margin:0 auto;background:#fff;border-radius:12px;overflow:hidden;box-shadow:0 4px 20px rgba(0,0,0,.1)}}
.hdr{{background:linear-gradient(135deg,#0F2942,#1a3a6b);padding:20px;text-align:center;color:#FFD700;font-size:1.1rem}}
.bdy{{padding:22px 20px;font-size:0.93rem;color:#333;line-height:1.9}}
.box{{background:#f8fafc;border-right:4px solid #10B981;padding:12px 16px;border-radius:8px;margin:14px 0}}
.ftr{{background:#f4f6f9;padding:12px;text-align:center;font-size:0.78rem;color:#777;border-top:1px solid #e0e0e0}}
</style></head><body>
<div class="wrap">
<div class="hdr">🎓 موعد مناقشة مذكرة الماستر</div>
<div class="bdy">
<p>الطالب(ة) العزيز(ة) <strong>{student_name}</strong>،</p>
<p>تحية طيبة وبعد،</p>
<p>يسعدنا إعلامكم بأنه تمّت برمجة مناقشة مذكرتكم الموسومة بـ:</p>
<div class="box"><strong>« {_title_tk1} »</strong></div>
<p>وذلك يوم <strong>{_date_tk1}</strong> على الساعة <strong>{_slot_tk1}</strong> بقاعة <strong>{_room_tk1}</strong><br>
تحت إشراف الأستاذ(ة): <strong>{_sup_tk1}</strong></p>
<p style="margin-top:12px;">🌐 منصة المذكرات: <a href="https://memoires2026.streamlit.app">memoires2026.streamlit.app</a></p>
<p style="margin-top:18px;color:#1a3a6b;">تقبلوا تحياتي،<br><strong>مسؤول الميدان — البروفيسور رفاف لخضر</strong></p>
</div>
<div class="ftr">كلية الحقوق والعلوم السياسية — جامعة محمد البشير الإبراهيمي، برج بوعريريج</div>
</div></body></html>"""

                    _recipients_tk1 = []
                    for _p, _r in [(_sup_tk1,"مشرفًا"), (_pres_tk1,"رئيسًا للجنة"), (_ex1_tk1,"عضوًا ممتحنًا"), (_ex2_tk1,"عضوًا ممتحنًا")]:
                        if _p and _p not in ["","nan"]:
                            _em = _email_dict.get(_p,"")
                            _recipients_tk1.append({"نوع":"أستاذ","اسم":_p,"صفة":_r,"بريد":_em})

                    _df_st_tk1 = load_students()
                    _st_email_col = next((c for c in ["البريد المهني","الإيميل","البريد الإلكتروني","email"] if c in _df_st_tk1.columns), None)
                    _reg1_tk1 = str(_row_tk1.get("رقم تسجيل الطالب 1","")).strip()
                    _reg2_tk1 = str(_row_tk1.get("رقم تسجيل الطالب 2","")).strip()
                    for _sn, _reg in [(_s1_tk1,_reg1_tk1), (_s2_tk1,_reg2_tk1)]:
                        if not _sn or _sn in ["","nan"]: continue
                        _st_email = ""
                        if _st_email_col and _reg and _reg not in ["","nan"]:
                            _st_row = _df_st_tk1[_df_st_tk1.iloc[:,2].astype(str).str.strip()==_reg.strip()]
                            if not _st_row.empty:
                                _st_email = str(_st_row.iloc[0].get(_st_email_col,"")).strip()
                        _recipients_tk1.append({"نوع":"طالب","اسم":_sn,"صفة":"طالب","بريد":_st_email})

                    st.markdown("**المستلمون:**")
                    for _rec in _recipients_tk1:
                        _icon = "👨‍🏫" if _rec["نوع"]=="أستاذ" else "🎓"
                        _em_show = _rec["بريد"] if _rec["بريد"] else "⚠️ لا بريد"
                        st.caption(f"{_icon} {_rec['اسم']} — {_rec['صفة']} — {_em_show}")

                    if st.button("📧 إرسال التكليف للجميع", type="primary", use_container_width=True, key="send_tk1_all"):
                        import smtplib
                        from email.mime.multipart import MIMEMultipart
                        from email.mime.text import MIMEText
                        _sent_tk1=0; _fail_tk1=[]
                        _prog_tk1=st.progress(0); _stat_tk1=st.empty()
                        for _pi, _rec in enumerate(_recipients_tk1):
                            if not _rec["بريد"]:
                                _fail_tk1.append(f"{_rec['اسم']} (لا بريد)"); continue
                            try:
                                _html_tk1 = _build_email_tk1(_rec["اسم"],_rec["صفة"]) if _rec["نوع"]=="أستاذ" else _build_student_email_tk1(_rec["اسم"])
                                _msg_tk1=MIMEMultipart("alternative")
                                _msg_tk1["Subject"]=f"برمجة مناقشة مذكرة الماستر رقم {_mid_tk1}"
                                _msg_tk1["From"]=EMAIL_SENDER; _msg_tk1["To"]=_rec["بريد"]
                                _msg_tk1.attach(MIMEText(_html_tk1,"html","utf-8"))
                                with smtplib.SMTP_SSL("smtp.gmail.com",465) as _srv_tk1:
                                    _srv_tk1.login(EMAIL_SENDER,EMAIL_PASSWORD)
                                    _srv_tk1.sendmail(EMAIL_SENDER,_rec["بريد"],_msg_tk1.as_string())
                                _sent_tk1+=1
                            except Exception as _etk1:
                                _fail_tk1.append(f"{_rec['اسم']}: {_etk1}")
                            _prog_tk1.progress(int((_pi+1)/len(_recipients_tk1)*100))
                            _stat_tk1.text(f"جاري الإرسال... {_pi+1}/{len(_recipients_tk1)}")
                        _prog_tk1.empty(); _stat_tk1.empty()
                        if _sent_tk1>0: st.success(f"✅ تم إرسال {_sent_tk1} رسالة")
                        if _fail_tk1: st.warning("⚠️ فشل: " + " | ".join(_fail_tk1[:5]))


        # ================================================================
        # TAB جدولة ذكية
        # ================================================================
        if not _is_printer_user and not _is_library_user and tab_siyar:
         with tab_siyar:
            st.subheader("📊 سير المناقشات")

            df_siyar = load_memos()
            _col_w_s = "تاريخ المناقشة"
            _col_hal = "الحالة" if "الحالة" in df_siyar.columns else None
            _col_note = "ملاحظات" if "ملاحظات" in df_siyar.columns else None
            _col_pdf = "رابط المحضر" if "رابط المحضر" in df_siyar.columns else None

            sched_siyar = df_siyar[
                df_siyar[_col_w_s].astype(str).str.strip().apply(lambda x: x not in ["","nan"])
            ].copy() if _col_w_s in df_siyar.columns else pd.DataFrame()

            if sched_siyar.empty:
                st.warning("⚠️ لا توجد مذكرات مبرمجة بعد")
            else:
                # ── إحصائيات ──
                _total = len(sched_siyar)
                _tammat = len(sched_siyar[sched_siyar.get(_col_hal, pd.Series()).astype(str)=="تمت"]) if _col_hal else 0
                _moajala = len(sched_siyar[sched_siyar.get(_col_hal, pd.Series()).astype(str)=="مؤجلة"]) if _col_hal else 0
                _molgha = len(sched_siyar[sched_siyar.get(_col_hal, pd.Series()).astype(str)=="ملغاة"]) if _col_hal else 0
                _montadhar = _total - _tammat - _moajala - _molgha

                _c1,_c2,_c3,_c4 = st.columns(4)
                _c1.metric("📋 الكل", _total)
                _c2.metric("✅ تمت", _tammat)
                _c3.metric("⏳ منتظرة", _montadhar)
                _c4.metric("🔄 مؤجلة/ملغاة", _moajala+_molgha)

                st.markdown("---")

                # ── فلتر اليوم ──
                _days_s = ["الكل"] + sorted(sched_siyar[_col_w_s].astype(str).unique().tolist())
                _sel_day_s = st.selectbox("📅 فلتر اليوم:", _days_s, key="siyar_day")
                if _sel_day_s != "الكل":
                    sched_siyar = sched_siyar[sched_siyar[_col_w_s].astype(str)==_sel_day_s]

                st.markdown("---")
                st.markdown("### تحديث حالة مذكرة")

                # ── اختيار المذكرة ──
                def _fmt_siyar(r):
                    mid = str(r.get("رقم المذكرة","")).strip()
                    s1 = str(r.get("الطالب الأول","")).strip()
                    day = str(r.get(_col_w_s,"")).strip()
                    hal = str(r.get(_col_hal,"")).strip() if _col_hal else ""
                    return f"{mid} | {s1} | {day}" + (f" | {hal}" if hal else "")

                try:
                    _sorted_s = sched_siyar.copy()
                    _sorted_s["_num"] = pd.to_numeric(_sorted_s["رقم المذكرة"], errors="coerce")
                    _sorted_s = _sorted_s.sort_values([_col_w_s,"_num"]).drop(columns=["_num"])
                except:
                    _sorted_s = sched_siyar.copy()

                _opts_s = [_fmt_siyar(r) for _, r in _sorted_s.iterrows()]
                _mids_s = _sorted_s["رقم المذكرة"].astype(str).tolist()

                _sel_s = st.selectbox("اختر المذكرة:", _opts_s, index=None, placeholder="اكتب للبحث...", key="siyar_sel")
                _sel_mid_s = _mids_s[_opts_s.index(_sel_s)] if _sel_s else None

                if _sel_mid_s:
                    _row_s = _sorted_s[_sorted_s["رقم المذكرة"].astype(str)==_sel_mid_s].iloc[0]

                    # الحالة الحالية
                    _cur_hal = str(_row_s.get(_col_hal,"")).strip() if _col_hal else ""
                    _cur_note = str(_row_s.get(_col_note,"")).strip() if _col_note else ""
                    _cur_pdf = str(_row_s.get(_col_pdf,"")).strip() if _col_pdf else ""

                    _col_a, _col_b = st.columns(2)
                    with _col_a:
                        _new_hal = st.selectbox("الحالة:", ["", "تمت", "مؤجلة", "ملغاة"],
                            index=["","تمت","مؤجلة","ملغاة"].index(_cur_hal) if _cur_hal in ["","تمت","مؤجلة","ملغاة"] else 0,
                            key="siyar_hal")
                    with _col_b:
                        _new_note = st.text_input("ملاحظات (سبب التأجيل / اليوم الجديد):", value=_cur_note, key="siyar_note")

                    # رفع PDF المحضر
                    st.markdown("**📄 رفع محضر المناقشة الممضى (PDF):**")
                    _upl_pdf = st.file_uploader("اختر ملف PDF:", type=["pdf"], key=f"pdf_upload_{_sel_mid_s}")

                    if _cur_pdf and _cur_pdf not in ["","nan"]:
                        st.success(f"✅ محضر موجود: [فتح المحضر]({_cur_pdf})")

                    if st.button("💾 حفظ التحديثات", type="primary", use_container_width=True, key="siyar_save"):
                        with st.spinner("⏳ جاري الحفظ..."):
                            # ابحث عن رقم الصف
                            _all_m = load_memos()
                            _row_map_s = {}
                            for _i_s, (_, _rw_s) in enumerate(_all_m.iterrows()):
                                _mk = str(_rw_s.get("رقم المذكرة","")).strip()
                                try: _mk = str(int(float(_mk)))
                                except: pass
                                if _mk: _row_map_s[_mk] = _i_s + 2

                            _rn_s = _row_map_s.get(str(_sel_mid_s))
                            if not _rn_s:
                                st.error("❌ لم يوجد الصف")
                            else:
                                _upd_siyar = []
                                # عمود AN = الحالة
                                _upd_siyar.append({"range": f"Feuille 1!AN{_rn_s}", "values": [[_new_hal]]})
                                # عمود AO = ملاحظات
                                _upd_siyar.append({"range": f"Feuille 1!AO{_rn_s}", "values": [[_new_note]]})

                                # رفع PDF إذا موجود
                                if _upl_pdf:
                                    _pdf_bytes = _upl_pdf.read()
                                    _pdf_name = f"محضر_{_sel_mid_s}.pdf"
                                    _ok_pdf, _link_pdf = upload_mahdar_pdf(_pdf_bytes, _pdf_name)
                                    if _ok_pdf:
                                        _upd_siyar.append({"range": f"Feuille 1!AP{_rn_s}", "values": [[_link_pdf]]})
                                        st.success(f"✅ رُفع المحضر: [فتح]({_link_pdf})")
                                    else:
                                        st.error(f"❌ فشل رفع PDF: {_link_pdf}")

                                sheets_service.spreadsheets().values().batchUpdate(
                                    spreadsheetId=MEMOS_SHEET_ID,
                                    body={"valueInputOption": "USER_ENTERED", "data": _upd_siyar}
                                ).execute()
                                clear_cache_and_reload()
                                st.success(f"✅ تم تحديث المذكرة {_sel_mid_s}")
                                st.rerun()

                st.markdown("---")
                # ── جدول الحالات ──
                st.markdown("### 📋 جدول المناقشات")
                _display_cols = ["رقم المذكرة","الطالب الأول",_col_w_s,"توقيت المناقشة","القاعة","الأستاذ","الرئيس"]
                if _col_hal: _display_cols.append(_col_hal)
                if _col_note: _display_cols.append(_col_note)
                _display_cols = list(dict.fromkeys([c for c in _display_cols if c in sched_siyar.columns]))
                _siyar_show = sched_siyar[_display_cols].copy()
                _siyar_show.columns = pd.io.common.dedup_names(_siyar_show.columns.tolist(), is_potential_multiindex=False) if hasattr(pd.io.common, 'dedup_names') else _siyar_show.columns
                st.dataframe(_siyar_show, use_container_width=True, hide_index=True)

        if not _is_printer_user and not _is_library_user and tab_seq:
         with tab_seq:
            st.subheader("📥 استيراد أرقام المحاضر")
            st.info("ارفع ملف Excel فيه: **A = رقم المذكرة | B = رقم المحضر** — سيُكتب في عمود AM في شيت المذكرات")
            _upl_seq = st.file_uploader("📂 ملف أرقام المحاضر:", type=["xlsx","xls"], key="import_seq_file")
            if _upl_seq:
                import openpyxl as _xl_s, io as _io_s
                try:
                    _wb_s = _xl_s.load_workbook(_io_s.BytesIO(_upl_seq.read()))
                    _ws_s = _wb_s.active
                    _seq_data = []
                    for _r_s in range(2, _ws_s.max_row+1):
                        _mid_s = _ws_s.cell(_r_s, 1).value
                        _seq_s = _ws_s.cell(_r_s, 2).value
                        if not _mid_s or not _seq_s: continue
                        try: _mid_s = str(int(float(str(_mid_s))))
                        except: _mid_s = str(_mid_s).strip()
                        _seq_data.append({'رقم المذكرة': _mid_s, 'رقم المحضر': str(_seq_s).strip()})
                    import pandas as _pd_s
                    st.success(f"✅ {len(_seq_data)} محضر جاهز")
                    st.dataframe(_pd_s.DataFrame(_seq_data).head(10), use_container_width=True)
                    if st.button("🚀 استيراد إلى الشيت", type="primary", use_container_width=True, key="do_import_seq"):
                        with st.spinner("⏳ جاري الاستيراد..."):
                            _am = load_memos()
                            _rm = {}
                            for _i_s, (_, _rw_s) in enumerate(_am.iterrows()):
                                _mk_s = str(_rw_s.get("رقم المذكرة","")).strip()
                                try: _mk_s = str(int(float(_mk_s)))
                                except: pass
                                if _mk_s: _rm[_mk_s] = _i_s + 2
                            _upd_s = []; _nf_s = []
                            for _d_s in _seq_data:
                                _rn_s = _rm.get(_d_s['رقم المذكرة'])
                                if not _rn_s: _nf_s.append(_d_s['رقم المذكرة']); continue
                                _upd_s.append({"range": f"Feuille 1!AM{_rn_s}", "values": [[_d_s['رقم المحضر']]]})
                            if _upd_s:
                                for _i2_s in range(0, len(_upd_s), 100):
                                    sheets_service.spreadsheets().values().batchUpdate(
                                        spreadsheetId=MEMOS_SHEET_ID,
                                        body={"valueInputOption": "USER_ENTERED", "data": _upd_s[_i2_s:_i2_s+100]}
                                    ).execute()
                                clear_cache_and_reload()
                                st.success(f"✅ تم استيراد {len(_seq_data)-len(_nf_s)} رقم محضر!")
                                if _nf_s: st.warning(f"⚠️ لم توجد في الشيت: {_nf_s[:10]}")
                            else:
                                st.error("❌ لا توجد تطابقات")
                except Exception as _e_s:
                    st.error(f"❌ {_e_s}")

        if not _is_printer_user and not _is_library_user and tab_stud:
         with tab_stud:
            st.subheader("👥 استيراد أرقام ملفات الطلبة")
            st.info("ارفع ملف Excel فيه: **A = رقم التسجيل | B = رقم الملف** — المقارنة مع عمود C في شيت الطلبة، الكتابة في عمود V")
            _upl_st = st.file_uploader("📂 ملف أرقام الملفات:", type=["xlsx","xls"], key="import_stud_file")
            if _upl_st:
                import openpyxl as _xl_t, io as _io_t
                try:
                    _wb_t = _xl_t.load_workbook(_io_t.BytesIO(_upl_st.read()))
                    _ws_t = _wb_t.active
                    _stud_data = []
                    for _r_t in range(2, _ws_t.max_row+1):
                        _reg_t = _ws_t.cell(_r_t, 1).value  # رقم التسجيل
                        _fil_t = _ws_t.cell(_r_t, 2).value  # رقم الملف
                        if not _reg_t or not _fil_t: continue
                        _stud_data.append({
                            'رقم التسجيل': str(_reg_t).strip(),
                            'رقم الملف': str(int(float(str(_fil_t)))) if str(_fil_t).replace('.','').isdigit() else str(_fil_t).strip()
                        })
                    import pandas as _pd_t
                    st.success(f"✅ {len(_stud_data)} طالب جاهز")
                    st.dataframe(_pd_t.DataFrame(_stud_data).head(10), use_container_width=True)
                    if st.button("🚀 استيراد إلى شيت الطلبة", type="primary", use_container_width=True, key="do_import_stud"):
                        with st.spinner("⏳ جاري الاستيراد..."):
                            # قراءة شيت الطلبة
                            _st_res = sheets_service.spreadsheets().values().get(
                                spreadsheetId=STUDENTS_SHEET_ID,
                                range="Feuille 1!A1:V5000"
                            ).execute()
                            _st_vals = _st_res.get('values', [])
                            # بناء خريطة رقم التسجيل (عمود C = index 2) → رقم الصف
                            _reg_map_t = {}
                            for _i_t, _row_t in enumerate(_st_vals[1:], 2):
                                if len(_row_t) >= 3:
                                    _rk_t = str(_row_t[2]).strip()
                                    if _rk_t: _reg_map_t[_rk_t] = _i_t
                            _upd_t = []; _nf_t = []
                            for _d_t in _stud_data:
                                _rn_t = _reg_map_t.get(_d_t['رقم التسجيل'])
                                if not _rn_t: _nf_t.append(_d_t['رقم التسجيل']); continue
                                _upd_t.append({"range": f"Feuille 1!V{_rn_t}", "values": [[_d_t['رقم الملف']]]})
                            if _upd_t:
                                for _i2_t in range(0, len(_upd_t), 100):
                                    sheets_service.spreadsheets().values().batchUpdate(
                                        spreadsheetId=STUDENTS_SHEET_ID,
                                        body={"valueInputOption": "USER_ENTERED", "data": _upd_t[_i2_t:_i2_t+100]}
                                    ).execute()
                                clear_cache_and_reload()
                                st.success(f"✅ تم تحديث {len(_stud_data)-len(_nf_t)} طالب!")
                                if _nf_t: st.warning(f"⚠️ لم يُوجد رقم تسجيلهم: {_nf_t[:10]}")
                            else:
                                st.error("❌ لا توجد تطابقات — تحقق من أن أرقام التسجيل صحيحة")
                except Exception as _e_t:
                    st.error(f"❌ {_e_t}")

        if not _is_printer_user and not _is_library_user and tab_archive:
         with tab_archive:
            st.info("📌 تم تنفيذ الجدولة — البرنامج محفوظ في الشيت.")
        if False:
         with tab9_dummy:
            st.subheader("📅 جدولة المناقشات الذكية")
            df_memos_j = load_memos()

            # فلتر نوع المذكرات
            def has_jury(row):
                def filled(val): return str(val).strip() not in ["","nan","None","—"]
                return filled(row.get("الأستاذ","")) and filled(row.get("الرئيس","")) and filled(row.get("المناقش1","")) and filled(row.get("المناقش2",""))

            # المذكرات القابلة للجدولة:
            # تُبرمج: كل ما ليس مرفوضاً وليس فارغاً (حتى لو ناقصة لجنة)
            # لا تُبرمج: المرفوضة + التي لم يُودَع فيها (T فارغ)
            def is_schedulable(r):
                t = str(r.get("حالة الإيداع","")).strip()
                # مرفوضة أو فارغة → لا
                if t in ["مرفوضة","","nan"]: return False
                return True

            ready_memos_j = df_memos_j[df_memos_j.apply(is_schedulable, axis=1)].copy()

            # إضافة عمود "سبب عدم الجدولة" للمذكرات الناقصة اللجنة
            def get_missing_members(r):
                missing = []
                for col, role in [("الأستاذ","المشرف"),("الرئيس","الرئيس"),("المناقش1","المناقش1"),("المناقش2","المناقش2")]:
                    v = str(r.get(col,"")).strip()
                    if not v or v in ["","nan","—"]: missing.append(role)
                return missing
            ready_memos_j["أعضاء_ناقصة"] = ready_memos_j.apply(lambda r: ",".join(get_missing_members(r)), axis=1)
            incomplete = ready_memos_j[ready_memos_j["أعضاء_ناقصة"] != ""]
            if not incomplete.empty:
                st.warning(f"⚠️ {len(incomplete)} مذكرة ناقصة اللجنة ستُبرمج بمن هو موجود")
                with st.expander("👁️ المذكرات الناقصة"):
                    st.dataframe(incomplete[["رقم المذكرة","أعضاء_ناقصة"]].rename(columns={"أعضاء_ناقصة":"الأعضاء الناقصون"}), use_container_width=True, hide_index=True)

            total_ready_j = len(ready_memos_j)

            # تحميل الجدول المحفوظ من الشيت تلقائياً عند أول فتح
            if "j_schedule" not in st.session_state and not ready_memos_j.empty:
                _saved_schedule = {}
                _has_saved = False
                _memo_members_saved = {}
                _, _memo_members_saved = build_prof_memo_map(ready_memos_j)
                for _, _r in ready_memos_j.iterrows():
                    _mid = str(_r.get("رقم المذكرة","")).strip()
                    _day = str(_r.get("تاريخ المناقشة","")).strip()
                    _slot = str(_r.get("توقيت المناقشة","")).strip()
                    _room = str(_r.get("القاعة","")).strip()
                    if _day and _day not in ["","nan"] and _slot and _slot not in ["","nan"]:
                        _saved_schedule[_mid] = (_day, _slot, _room if _room not in ["","nan"] else "")
                        _has_saved = True
                    else:
                        _saved_schedule[_mid] = None
                if _has_saved:
                    st.session_state["j_schedule"] = _saved_schedule
                    st.session_state["j_memo_members"] = _memo_members_saved
                    st.session_state["j_rej_log"] = {}
                    st.info(f"📂 تم تحميل الجدول المحفوظ من الشيت ({sum(1 for v in _saved_schedule.values() if v)} مذكرة مبرمجة)")

            already_sched = len(ready_memos_j[ready_memos_j["تاريخ المناقشة"].astype(str).str.strip().apply(lambda x: x not in ["","nan"])]) if "تاريخ المناقشة" in ready_memos_j.columns else 0

            # KPIs
            k1,k2,k3 = st.columns(3)
            with k1: st.markdown(f'''<div class="kpi-card"><div class="kpi-value" style="color:#FFD700;">{total_ready_j}</div><div class="kpi-label">🎓 جاهزة للجدولة</div></div>''', unsafe_allow_html=True)
            with k2: st.markdown(f'''<div class="kpi-card"><div class="kpi-value" style="color:#10B981;">{already_sched}</div><div class="kpi-label">📅 مبرمجة</div></div>''', unsafe_allow_html=True)
            with k3: st.markdown(f'''<div class="kpi-card"><div class="kpi-value" style="color:#F59E0B;">{total_ready_j - already_sched}</div><div class="kpi-label">⏳ غير مبرمجة</div></div>''', unsafe_allow_html=True)

            if ready_memos_j.empty:
                st.warning("⏳ لا توجد مذكرات جاهزة بعد.")
            else:
                st.markdown("---")
                st.markdown("### ⚙️ إعداد معطيات الجدولة")
                # قراءة الإعدادات من الشيت
                _days_from_sheet, _day_limits = load_schedule_days()
                _slots_from_sheet = load_schedule_slots()
                _rooms_from_sheet = load_schedule_rooms()

                col_p1, col_p2 = st.columns(2)
                with col_p1:
                    st.markdown("**📆 الأيام**")
                    if _days_from_sheet:
                        st.success(f"✅ {len(_days_from_sheet)} يوم مُحمَّل من الشيت")
                        gen_days_j = _days_from_sheet
                        with st.expander("👁️ عرض الأيام", expanded=False):
                            st.write(" | ".join(gen_days_j))
                    else:
                        st.warning("⚠️ لا توجد أيام في شيت 'الأيام' — أدخلها يدوياً")
                        import datetime as _dt
                        num_days_j = st.number_input("عدد الأيام", min_value=3, max_value=20, value=10, key="j_ndays")
                        base_date_j = st.date_input("تاريخ البداية", value=date.today(), key="j_basedate")
                        gen_days_j = []
                        d_j = base_date_j
                        while len(gen_days_j) < num_days_j:
                            if d_j.weekday() not in [4,5]: gen_days_j.append(d_j.strftime("%Y-%m-%d"))
                            d_j += _dt.timedelta(days=1)

                with col_p2:
                    st.markdown("**🕐 التوقيتات والقاعات**")
                    if _slots_from_sheet:
                        st.success(f"✅ {len(_slots_from_sheet)} توقيت مُحمَّل")
                        gen_slots_j = _slots_from_sheet
                        with st.expander("👁️ التوقيتات", expanded=False):
                            st.write(" | ".join(gen_slots_j))
                    else:
                        st.warning("⚠️ أدخل التوقيتات يدوياً")
                        slots_txt = st.text_area("التوقيتات","09:30\n11:30\n14:00\n16:00",height=100,key="j_slots_txt")
                        gen_slots_j = [s.strip() for s in slots_txt.strip().split("\n") if s.strip()]

                    if _rooms_from_sheet:
                        st.success(f"✅ {len(_rooms_from_sheet)} قاعة مُحمَّلة")
                        gen_rooms_j = _rooms_from_sheet
                        with st.expander("👁️ القاعات", expanded=False):
                            st.write(" | ".join(gen_rooms_j))
                    else:
                        st.warning("⚠️ أدخل القاعات يدوياً")
                        rooms_txt = st.text_area("القاعات","\n".join([f"قاعة {i+1:02d}" for i in range(10)]),height=120,key="j_rooms_txt")
                        gen_rooms_j = [r.strip() for r in rooms_txt.strip().split("\n") if r.strip()]

                    capacity = len(gen_days_j)*len(gen_slots_j)*len(gen_rooms_j)
                    cap_color = "#10B981" if capacity >= total_ready_j else "#EF4444"
                    st.markdown(f'''<div style="background:rgba(47,111,126,0.1);border:1px solid #2F6F7E;border-radius:10px;padding:10px 14px;margin-top:8px;">
                        الطاقة: <strong style="color:{cap_color};">{capacity} خانة</strong> لـ <strong style="color:#FFD700;">{total_ready_j} مذكرة</strong>
                    </div>''', unsafe_allow_html=True)
                
                # زر تحديث الإعدادات
                if st.button("🔄 تحديث من الشيت", key="refresh_settings", use_container_width=False):
                    st.cache_data.clear()
                    st.rerun()

                st.markdown("---")

                # ================================================================
                # الاستثناءات
                # ================================================================
                with st.expander("⚙️ الاستثناءات والقيود", expanded=False):
                    df_memo_exc = load_memo_exceptions()
                    df_prof_exc = load_prof_exceptions()

                    exc_tab1, exc_tab2 = st.tabs(["📋 استثناءات المذكرات","👨‍🏫 استثناءات الأساتذة"])

                    with exc_tab1:
                        # عرض الاستثناءات الحالية
                        if not df_memo_exc.empty:
                            st.dataframe(df_memo_exc, use_container_width=True, hide_index=True)
                        else:
                            st.info("لا توجد استثناءات مذكرات بعد")

                        st.markdown("**➕ إضافة استثناء مذكرة:**")
                        ex_m1, ex_m2 = st.columns(2)
                        with ex_m1:
                            exc_memo_id = st.text_input("رقم المذكرة", key="exc_memo_id")
                            exc_memo_day = st.selectbox("يوم مثبت (اختياري)", ["—"]+gen_days_j, key="exc_memo_day")
                            exc_memo_slot = st.selectbox("توقيت مثبت (اختياري)", ["—"]+gen_slots_j, key="exc_memo_slot")
                        with ex_m2:
                            exc_memo_room = st.selectbox("قاعة مثبتة (اختياري)", ["—"]+gen_rooms_j, key="exc_memo_room")
                            exc_memo_early = st.text_input("أقرب تاريخ (YYYY-MM-DD)", key="exc_memo_early")
                            exc_memo_late = st.text_input("أبعد تاريخ (YYYY-MM-DD)", key="exc_memo_late")
                            exc_memo_alt = st.multiselect("📅 أيام بديلة (أو/أو)", gen_days_j, key="exc_memo_alt")

                        if st.button("💾 حفظ استثناء المذكرة", key="save_memo_exc", use_container_width=True):
                            if exc_memo_id.strip():
                                ok = save_memo_exception({
                                    "رقم المذكرة": exc_memo_id.strip(),
                                    "يوم مثبت": exc_memo_day if exc_memo_day != "—" else "",
                                    "توقيت مثبت": exc_memo_slot if exc_memo_slot != "—" else "",
                                    "قاعة مثبتة": exc_memo_room if exc_memo_room != "—" else "",
                                    "أقرب تاريخ": exc_memo_early.strip(),
                                    "أبعد تاريخ": exc_memo_late.strip(),
                                    "أيام بديلة": ",".join(exc_memo_alt),
                                })
                                if ok:
                                    st.success("✅ تم الحفظ!")
                                    st.cache_data.clear()
                                    st.rerun()
                            else:
                                st.error("❌ أدخل رقم المذكرة")

                    with exc_tab2:
                        # عرض الاستثناءات الحالية
                        if not df_prof_exc.empty:
                            st.dataframe(df_prof_exc, use_container_width=True, hide_index=True)
                        else:
                            st.info("لا توجد استثناءات أساتذة بعد")

                        # ── استخراج تلقائي بالذكاء الاصطناعي ──
                        st.markdown("---")
                        st.markdown("**🤖 استخراج تلقائي من رسالة الأستاذ (دارجة أو عربية):**")
                        ai_msg = st.text_area("الصق رسالة الأستاذ هنا:", height=120, key="ai_exc_msg",
                            placeholder="مثال: برمجلي المناقشات تاوعي يوم 6 و3 جوان مديرليش مناقشات 31 ماي...")

                        if st.button("🤖 استخراج الاستثناءات", key="ai_extract_btn", use_container_width=True):
                            if ai_msg.strip():
                                with st.spinner("🧠 جاري التحليل..."):
                                    try:
                                        import re as _re

                                        def _parse_prof_message(msg, available_days):
                                            """محلل محلي للرسائل بالدارجة والعربية"""
                                            result = {
                                                "اسم_الأستاذ": "",
                                                "أيام_ممنوعة": [],
                                                "أيام_مسموحة_فقط": [],
                                                "لا_قبل": "",
                                                "لا_بعد": "",
                                                "يوم_واحد": False,
                                                "مجمد": False,
                                                "ملاحظات": ""
                                            }

                                            msg_lower = msg.lower()

                                            # خريطة الأشهر
                                            months = {
                                                "جانفي":1,"فيفري":2,"مارس":3,"أفريل":4,"ماي":5,"mai":5,
                                                "جوان":6,"juin":6,"جويلية":7,"أوت":8,"سبتمبر":9,
                                                "أكتوبر":10,"نوفمبر":11,"ديسمبر":12
                                            }

                                            # استخراج التاريخ من رقم + شهر
                                            def extract_dates(text, month_hint=None):
                                                dates = []
                                                text_l = text  # لا نستخدم lower مع العربية
                                                for month_name, month_num in months.items():
                                                    if month_name not in text_l: continue
                                                    parts = text_l.split(month_name)
                                                    # أرقام قبل الشهر
                                                    seg_before = parts[0][-80:]
                                                    nums_before = _re.findall(r"\b(\d{1,2})\b", seg_before)
                                                    for n in nums_before:
                                                        day = int(n)
                                                        if 1 <= day <= 31:
                                                            dates.append(f"2026-{month_num:02d}-{day:02d}")
                                                    # أرقام بعد الشهر
                                                    if len(parts) > 1:
                                                        seg_after = parts[1][:80]
                                                        nums_after = _re.findall(r"\b(\d{1,2})\b", seg_after)
                                                        for n in nums_after:
                                                            day = int(n)
                                                            if 1 <= day <= 31:
                                                                dates.append(f"2026-{month_num:02d}-{day:02d}")
                                                return list(set(dates))

                                            all_days = set(available_days) if available_days else set()

                                            # التجميد الكلي
                                            frozen_keywords = ["إلى إشعار آخر","إلى إشعار","عدم برمجته","عدم برمجتي","لا تبرمج","عدم برمجتك"]
                                            if any(kw in msg for kw in frozen_keywords):
                                                result["مجمد"] = True
                                                result["ملاحظات"] = "تجميد كلي مطلوب"
                                                return result

                                            # جمل الأيام المسموحة
                                            allowed_triggers = ["برمجلي","برمجني","نحب","اختار"]
                                            banned_triggers = ["مديرليش","ماتبرمجنيش","ماتبرمج","لا تبرمج","مانكونش"]
                                            # قسّم النص على الكلمات المفتاحية
                                            _segments = _re.split(r"(مديرليش|ماتبرمجنيش|ماتبرمج|لا تبرمج|مانكونش|برمجلي|برمجني)", msg)
                                            _current_mode = "مسموح"  # default
                                            for _seg in _segments:
                                                if any(kw in _seg for kw in banned_triggers): _current_mode = "ممنوع"
                                                elif any(kw in _seg for kw in allowed_triggers): _current_mode = "مسموح"
                                                else:
                                                    _dates = extract_dates(_seg)
                                                    if _current_mode == "ممنوع":
                                                        result["أيام_ممنوعة"].extend(_dates)
                                                    else:
                                                        result["أيام_مسموحة_فقط"].extend(_dates)

                                            # إذا لم نجد شيئاً، استخرج كل التواريخ
                                            if not result["أيام_ممنوعة"] and not result["أيام_مسموحة_فقط"]:
                                                all_dates = extract_dates(msg)
                                                if any(kw in msg for kw in ["مديرليش","ماتبرمج","لا يوم"]):
                                                    result["أيام_ممنوعة"] = all_dates
                                                else:
                                                    result["أيام_مسموحة_فقط"] = all_dates

                                            # التوقيت
                                            time_match = _re.search(r'(\d{1,2})[h:h]\s*(\d{0,2})', msg)
                                            if time_match:
                                                h = int(time_match.group(1))
                                                mn = time_match.group(2) or "00"
                                                t_str = f"{h:02d}:{mn.zfill(2)}"
                                                if any(kw in msg for kw in ["ابتداء","من الساعة","بعد","après"]):
                                                    result["لا_قبل"] = t_str
                                                elif any(kw in msg for kw in ["قبل","avant","حتى"]):
                                                    result["لا_بعد"] = t_str

                                            # إزالة التكرار
                                            result["أيام_ممنوعة"] = sorted(set(result["أيام_ممنوعة"]))
                                            result["أيام_مسموحة_فقط"] = sorted(set(result["أيام_مسموحة_فقط"]) - set(result["أيام_ممنوعة"]))

                                            if not result["ملاحظات"]:
                                                result["ملاحظات"] = "راجع النتيجة قبل الحفظ"

                                            return result

                                        _extracted = _parse_prof_message(ai_msg.strip(), gen_days_j)
                                        st.session_state["ai_extracted"] = _extracted
                                    except Exception as _ae:
                                        st.error(f"❌ خطأ في الاستخراج: {str(_ae)}")

                        # عرض النتيجة المستخرجة للمراجعة
                        if "ai_extracted" in st.session_state:
                            _ext = st.session_state["ai_extracted"]
                            st.markdown("**📋 النتيجة المستخرجة — راجع قبل الحفظ:**")

                            ai_col1, ai_col2 = st.columns(2)
                            with ai_col1:
                                ai_prof = st.text_input("اسم الأستاذ", value=_ext.get("اسم_الأستاذ",""), key="ai_prof_name")
                                ai_banned = st.multiselect("أيام ممنوعة", gen_days_j,
                                    default=[d for d in _ext.get("أيام_ممنوعة",[]) if d in gen_days_j], key="ai_banned")
                                ai_allowed = st.multiselect("أيام مسموحة فقط", gen_days_j,
                                    default=[d for d in _ext.get("أيام_مسموحة_فقط",[]) if d in gen_days_j], key="ai_allowed")
                            with ai_col2:
                                ai_before = st.selectbox("لا قبل", ["—"]+(_slots_from_sheet or ["09:30","11:30","14:00","16:00"]),
                                    index=(["—"]+(_slots_from_sheet or [])).index(_ext.get("لا_قبل","—")) if _ext.get("لا_قبل","") in (["—"]+(_slots_from_sheet or [])) else 0,
                                    key="ai_before")
                                ai_after = st.selectbox("لا بعد", ["—"]+(_slots_from_sheet or ["09:30","11:30","14:00","16:00"]),
                                    index=(["—"]+(_slots_from_sheet or [])).index(_ext.get("لا_بعد","—")) if _ext.get("لا_بعد","") in (["—"]+(_slots_from_sheet or [])) else 0,
                                    key="ai_after")
                                ai_oneday = st.checkbox("يوم واحد فقط", value=bool(_ext.get("يوم_واحد", False)), key="ai_oneday")
                                ai_frozen = st.checkbox("تجميد كلي", value=bool(_ext.get("مجمد", False)), key="ai_frozen")

                            if _ext.get("ملاحظات"):
                                st.info(f"💡 ملاحظات الذكاء الاصطناعي: {_ext['ملاحظات']}")

                            if st.button("💾 حفظ الاستثناء المستخرج", type="primary", use_container_width=True, key="ai_save_btn"):
                                if ai_prof.strip():
                                    _ok = save_prof_exception({
                                        "اسم الأستاذ": ai_prof.strip(),
                                        "أيام ممنوعة": ",".join(ai_banned),
                                        "أيام مسموحة فقط": ",".join(ai_allowed),
                                        "لا قبل": ai_before if ai_before != "—" else "",
                                        "لا بعد": ai_after if ai_after != "—" else "",
                                        "يوم واحد": "نعم" if ai_oneday else "",
                                        "أيام متتالية": "",
                                        "مجمّد": "نعم" if ai_frozen else "",
                                        "يقبل 18:00": "",
                                        "يقبل 18:00": "",
                                        "عدد مناقشات الفترة الأولى": "",
                                        "بداية الفترة الثانية": "",
                                    })
                                    if _ok:
                                        st.success("✅ تم الحفظ!")
                                        del st.session_state["ai_extracted"]
                                        st.cache_data.clear()
                                        st.rerun()
                                    else:
                                        st.error("❌ فشل الحفظ")
                                else:
                                    st.error("❌ أدخل اسم الأستاذ")

                        st.markdown("---")
                        st.markdown("**➕ إضافة استثناء أستاذ:**")
                        ex_p1, ex_p2 = st.columns(2)
                        with ex_p1:
                            _profs_all = set()
                            for _pc in ["الأستاذ","الرئيس","المناقش1","المناقش2"]:
                                if _pc in df_memos_j.columns:
                                    _profs_all.update(df_memos_j[_pc].astype(str).str.strip().unique())
                            profs_for_exc = sorted(_profs_all - {"","nan","None"})
                            exc_prof_name = st.selectbox("اسم الأستاذ", ["—"]+profs_for_exc, key="exc_prof_name")
                            exc_prof_banned = st.multiselect("🚫 أيام ممنوعة", gen_days_j, key="exc_prof_banned")
                            exc_prof_allowed = st.multiselect("✅ أيام مسموحة فقط", gen_days_j, key="exc_prof_allowed")
                        with ex_p2:
                            exc_prof_before = st.selectbox("⏰ لا قبل (توقيت)", ["—"]+gen_slots_j, key="exc_prof_before")
                            exc_prof_after = st.selectbox("⏰ لا بعد (توقيت)", ["—"]+gen_slots_j, key="exc_prof_after")
                            exc_prof_oneday = st.checkbox("📅 يوم واحد فقط", key="exc_prof_oneday")
                            exc_prof_consec = st.checkbox("📅 أيام متتالية", key="exc_prof_consec")
                            exc_prof_cluster = st.checkbox("🏠 تجميع الأيام (يسكن بعيداً)", key="exc_prof_cluster")
                            exc_prof_frozen = st.checkbox("🔒 تجميد كلي (لا يُبرمج حتى إشعار آخر)", key="exc_prof_frozen")
                            exc_prof_late = st.checkbox("🌙 يقبل توقيت 18:00", key="exc_prof_late")
                            exc_prof_18 = st.checkbox("🌙 يقبل توقيت 18:00", key="exc_prof_18")
                            st.markdown("**تقسيم على فترتين:**")
                            exc_prof_n_first = st.number_input("عدد مناقشات الفترة الأولى", min_value=0, max_value=20, value=0, key="exc_prof_n_first")
                            exc_prof_start2 = st.selectbox("بداية الفترة الثانية", ["—"]+gen_days_j, key="exc_prof_start2")

                        if st.button("💾 حفظ استثناء الأستاذ", key="save_prof_exc", use_container_width=True):
                            if exc_prof_name != "—":
                                ok = save_prof_exception({
                                    "اسم الأستاذ": exc_prof_name,
                                    "أيام ممنوعة": ",".join(exc_prof_banned),
                                    "أيام مسموحة فقط": ",".join(exc_prof_allowed),
                                    "لا قبل": exc_prof_before if exc_prof_before != "—" else "",
                                    "لا بعد": exc_prof_after if exc_prof_after != "—" else "",
                                    "يوم واحد": "نعم" if exc_prof_oneday else "",
                                    "أيام متتالية": "نعم" if exc_prof_consec else "",
                                    "تجميع الأيام": "نعم" if exc_prof_cluster else "",
                                    "مجمّد": "نعم" if exc_prof_frozen else "",
                                    "يقبل 18:00": "نعم" if exc_prof_late else "",
                                    "يقبل 18:00": "نعم" if exc_prof_18 else "",
                                    "عدد مناقشات الفترة الأولى": str(exc_prof_n_first) if exc_prof_n_first > 0 else "",
                                    "بداية الفترة الثانية": exc_prof_start2 if exc_prof_start2 != "—" else "",
                                })
                                if ok:
                                    st.success("✅ تم الحفظ!")
                                    st.cache_data.clear()
                                    st.rerun()
                            else:
                                st.error("❌ اختر اسم الأستاذ")

                st.markdown("---")
                # ── إعدادات الفترة المكثفة ──
                with st.expander("⚙️ إعدادات الفترة المكثفة", expanded=False):
                    _use_intensive = st.checkbox("تفعيل الفترة المكثفة", value=True, key="j_use_intensive")
                    if _use_intensive:
                        _col_i1, _col_i2 = st.columns(2)
                        with _col_i1:
                            st.date_input(
                                "📅 نهاية البرمجة المكثفة:",
                                value=date(2026, 6, 8),
                                min_value=date(2026, 5, 31),
                                max_value=date(2026, 6, 30),
                                key="j_cutoff_date",
                            )
                        with _col_i2:
                            st.number_input(
                                "🎯 النسبة الهدف (%):",
                                min_value=50, max_value=95, value=70, step=5,
                                key="j_target_ratio",
                            )
                        _cd = st.session_state.get("j_cutoff_date", date(2026,6,8))
                        _tr = st.session_state.get("j_target_ratio", 70)
                        st.caption(f"**{_tr}%** من المذكرات قبل **{_cd.strftime('%d/%m/%Y')}**")
                    else:
                        st.info("كل الأيام سواسية — لا أولوية زمنية")

                # بذرة التوليد — تغييرها يعطي جدولاً مختلفاً
                _seed = st.session_state.get("j_seed", 42)
                c_seed1, c_seed2 = st.columns([3,1])
                with c_seed1:
                    st.caption(f"🎲 بذرة التوليد الحالية: **{_seed}**")
                with c_seed2:
                    if st.button("🎲 بذرة جديدة", key="j_new_seed"):
                        import random as _rnd
                        st.session_state["j_seed"] = _rnd.randint(1, 9999)
                        st.rerun()

                algo_choice = st.radio(
                    "🧠 اختر الخوارزمية:",
                    ["🧱 كتل الأساتذة", "📅 الجدول أولاً", "⚡ الأثقل أولاً (Greedy)", "🧬 GA + Tabu Search"],
                    horizontal=True, key="algo_choice",
                    help="كتل: تجميع مضمون | الجدول أولاً: توزيع متوازن | Greedy: الأسرع"
                )

                c_g1, c_g2, c_g3 = st.columns(3)
                with c_g1:
                    _n_tries = st.number_input("عدد المحاولات:", min_value=1, max_value=50, value=10, step=1, key="j_n_tries")
                    if st.button("🏆 أفضل جدول (Multi-Start)", use_container_width=True, key="j_multi_start"):
                        _pt = st.empty(); _pb = st.progress(0)
                        def _ms_cb(i,n,sc,pl):
                            _pt.text(f"محاولة {i}/{n} | score: {sc:.0f} | مجدول: {pl}")
                            _pb.progress(int(i/n*100))
                        _dme=load_memo_exceptions(); _dpe=load_prof_exceptions()
                        _f2,_dl2,_bd2,_nb2,_na2,_o2,_ad2,_co2,_fr2,_ph2,_alt2,_a182,_cl2=build_constraints(_dme,_dpe,gen_slots_j)
                        _cms=(_f2,_dl2,_bd2,_nb2,_na2,_o2,_ad2,_co2,_fr2,_ph2,_alt2,_a182,_cl2)
                        _bs,_bmm,_br,_bsc=multi_start_best(ready_memos_j,gen_days_j,gen_slots_j,gen_rooms_j,_cms,
                            n_tries=int(_n_tries),algo_name=st.session_state.get("algo_choice","🧱 كتل الأساتذة"),progress_cb=_ms_cb)
                        _pt.empty(); _pb.empty()
                        if _bs:
                            _q,_pl,_upl,_id,_td,_=calc_schedule_quality(_bs,_bmm,gen_days_j,gen_slots_j)
                            st.session_state.update({"j_schedule":_bs,"j_memo_members":_bmm,"j_rej_log":_br,
                                "j_quality":_q,"j_placed":_pl,"j_unplaced":_upl,
                                "j_slots_list":gen_slots_j,"j_days_list":gen_days_j,"j_rooms_list":gen_rooms_j})
                            st.success(f"✅ أفضل جدول: {_pl}/{_pl+_upl} مجدول | جودة: {_q}%")
                            st.rerun()
                    if st.button("🚀 توليد الجدول الذكي", type="primary", use_container_width=True, key="j_generate"):
                        if capacity < total_ready_j:
                            st.error(f"❌ الطاقة ({capacity}) أقل من عدد المذكرات ({total_ready_j})")
                        else:
                            # قراءة الاستثناءات وفحص التعارضات أولاً
                            _df_memo_exc = load_memo_exceptions()
                            _df_prof_exc = load_prof_exceptions()
                            _fixed, _date_lim, _ban_days, _not_bef, _not_aft, _one_day, _allow_days, _consec, _frozen, _phase, _alt_days, _acc18, _cluster = build_constraints(
                                _df_memo_exc, _df_prof_exc, gen_slots_j
                            )

                            # دمج قيود التوقيت اليومي
                            _day_time_limits = _days_from_sheet and _day_limits or {}
                            _conflicts = detect_constraint_conflicts(
                                ready_memos_j, _fixed, _date_lim, _ban_days, _not_bef, _not_aft, gen_slots_j
                            )
                            if _conflicts:
                                st.error("⛔ تم اكتشاف تعارضات في الاستثناءات — يجب حلها قبل التوليد:")
                                for _cf in _conflicts:
                                    st.markdown(f"- {_cf}")
                                st.stop()
                            else:
                                with st.spinner("🧠 DSatur + Tabu Search + Post-Optimization..."):
                                    _fixed, _date_lim, _ban_days, _not_bef, _not_aft, _one_day, _allow_days, _consec, _frozen, _phase, _alt_days, _acc18, _cluster = build_constraints(
                                        _df_memo_exc, _df_prof_exc, gen_slots_j
                                    )
                                    _constraints = (_fixed, _date_lim, _ban_days, _not_bef, _not_aft,
                                        _one_day, _allow_days, _consec, _frozen, _phase,
                                        _alt_days, _acc18, _cluster)

                                    _seed = st.session_state.get("j_seed", 42)
                                    _algo = st.session_state.get("algo_choice", "🧱 كتل الأساتذة")
                                    schedule_j, quality_j, placed_j, unplaced_j, idle_j, days_j, memo_members_j, rej_log_j = run_algorithm(
                                        _algo, ready_memos_j, gen_days_j, gen_slots_j, gen_rooms_j,
                                        _constraints, improve=True, seed=_seed
                                    )
                                    st.session_state["j_schedule"] = schedule_j
                                    st.session_state["j_score"] = quality_j
                                    st.session_state["j_unplaced"] = unplaced_j
                                    st.session_state["j_slots_list"] = gen_slots_j
                                    st.session_state["j_days_list"] = gen_days_j
                                    st.session_state["j_rooms_list"] = gen_rooms_j
                                    st.session_state["j_memo_members"] = memo_members_j
                                    st.session_state["j_rej_log"] = rej_log_j
                                    # ملخص القيود المطبقة
                                    _cs = []
                                    if _fixed: _cs.append(f"📌 {len(_fixed)} موعد مثبت")
                                    if _ban_days: _cs.append(f"🚫 {len(_ban_days)} أستاذ أيام ممنوعة")
                                    if _allow_days: _cs.append(f"✅ {len(_allow_days)} أستاذ أيام مسموحة فقط")
                                    if _not_bef: _cs.append(f"⏰ {len(_not_bef)} لا قبل توقيت")
                                    if _not_aft: _cs.append(f"⏰ {len(_not_aft)} لا بعد توقيت")
                                    if _frozen: _cs.append(f"🔒 {len(_frozen)} أستاذ مجمّد")
                                    if _acc18: _cs.append(f"🌙 {len(_acc18)} يقبل 18:00")
                                    if _alt_days: _cs.append(f"📅 {len(_alt_days)} مذكرة أيام بديلة")
                                    if _cluster: _cs.append(f"🏠 {len(_cluster)} أستاذ تجميع أيام")
                                    st.session_state["j_constraints_summary"] = _cs
                                    st.session_state["_j_ban_days"] = _ban_days
                                    st.session_state["_j_allow_days"] = _allow_days
                                st.success(f"✅ مجدول: {placed_j}/{placed_j+unplaced_j} | غير مجدول: {unplaced_j} | جودة: {quality_j}%")
                                st.rerun()
                with c_g2:
                    if st.button("⚡ تحسين إضافي", use_container_width=True, key="j_improve", disabled="j_schedule" not in st.session_state):
                        if "j_schedule" in st.session_state:
                            with st.spinner("⚡ جاري التحسين..."):
                                cur = st.session_state["j_schedule"]
                                sl_j = st.session_state.get("j_slots_list", gen_slots_j)
                                dy_j = st.session_state.get("j_days_list", gen_days_j)
                                rm_j = st.session_state.get("j_rooms_list", gen_rooms_j)
                                _, conflicts_j2, memo_members_j2 = build_conflict_matrix(ready_memos_j)
                                _pbd2 = st.session_state.get("_j_ban_days", {})
                                _pad2 = st.session_state.get("_j_allow_days", {})
                                _fix2 = _constraints[0] if "_constraints" in dir() else {}

                                # ── مرحلة ملء ذكية: استغلال الفراغات ──
                                _occ2 = {sv: m for m, sv in cur.items() if sv}
                                _pbusy2 = {}
                                _pdc2 = {}
                                for _m2, _sv2 in cur.items():
                                    if not _sv2: continue
                                    _d2, _s2, _r2 = _sv2
                                    for _p2 in memo_members_j2.get(_m2, set()):
                                        _pbusy2[(_d2, _s2, _p2)] = _m2
                                        _pdc2[(_p2, _d2)] = _pdc2.get((_p2, _d2), 0) + 1

                                _can2, _ = make_can_place(_occ2, _pbusy2, _pdc2, memo_members_j2,
                                    _fix2, {}, _pbd2, {}, {}, {}, set(), set(), {}, {}, sl_j, None)
                                _place2 = make_place_fn(cur, _occ2, _pbusy2, _pdc2, memo_members_j2, {}, {})

                                # ابحث عن المذكرات غير المجدولة وضعها في الفراغات
                                _unscheduled = [m for m, sv in cur.items() if not sv]
                                _filled = 0
                                for _umid in _unscheduled:
                                    for _uday in sorted(dy_j, key=lambda d: sum(1 for sv in cur.values() if sv and sv[0]==d)):
                                        for _uslot in sl_j:
                                            for _uroom in rm_j:
                                                if _can2(_umid, _uday, _uslot, _uroom, log=False):
                                                    _place2(_umid, _uday, _uslot, _uroom)
                                                    _filled += 1; break
                                            else: continue; break
                                        else: continue; break

                                # احمل القيود الكاملة
                                _df_memo_exc_i = load_memo_exceptions()
                                _df_prof_exc_i = load_prof_exceptions()
                                _fix_i, _, _pbd_i, _, _, _, _pad_i, _, _, _, _, _acc18_i, _ = build_constraints(_df_memo_exc_i, _df_prof_exc_i, sl_j)
                                better_j = improve_schedule(cur, memo_members_j2, dy_j, sl_j, rm_j, iterations=2000,
                                    prof_banned_days=_pbd_i, prof_allowed_days=_pad_i,
                                    profs_accept_18=_acc18_i, fixed_slots=_fix_i)
                                # أعد تطبيق المثبتات بالقوة
                                for _fmid_i, _fsv_i in _fix_i.items():
                                    _fd_i, _fs_i, _fr_i = _fsv_i
                                    if _fd_i and _fs_i and str(_fmid_i) in better_j:
                                        better_j[str(_fmid_i)] = (_fd_i, _fs_i, _fr_i if _fr_i else rm_j[0])
                                q2, p2, u2, i2, d2, _ = calc_schedule_quality(better_j, memo_members_j2, dy_j, sl_j)
                                st.session_state["j_schedule"] = better_j
                                st.session_state["j_score"] = q2
                                st.session_state["j_unplaced"] = u2
                            st.success(f"✅ تم التحسين! جودة: {q2}%")
                            st.rerun()
                with c_g3:
                    if st.button("🗑️ مسح الجدول", use_container_width=True, key="j_reset"):
                        for k in ["j_schedule","j_score","j_unplaced","j_slots_list","j_days_list","j_rooms_list","j_memo_members","j_confirm_step"]:
                            st.session_state.pop(k, None)
                        st.rerun()

                # عرض الجدول
                if "j_schedule" in st.session_state:
                    j_sched = st.session_state["j_schedule"]
                    j_score = st.session_state.get("j_score", 0)
                    j_unpl = st.session_state.get("j_unplaced", 0)
                    j_sl = st.session_state.get("j_slots_list", gen_slots_j)
                    j_dy = st.session_state.get("j_days_list", gen_days_j)
                    j_rm = st.session_state.get("j_rooms_list", gen_rooms_j)

                    sc_color = "#10B981" if j_score>=90 else "#F59E0B" if j_score>=70 else "#EF4444"
                    placed_count = len([s for s in j_sched.values() if s])
                    st.markdown(f'''<div style="background:linear-gradient(135deg,#0F2942,#1A3A5C);border-radius:16px;padding:16px 24px;margin:16px 0;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:12px;border:1px solid rgba(47,111,126,0.3);">
                        <div><div style="font-size:1.1rem;font-weight:800;color:#FFD700;">📊 جودة الجدول</div>
                        <div style="color:#E2E8F0;font-size:0.85rem;margin-top:4px;">مجدول: {placed_count}/{len(j_sched)} | غير مجدول: {j_unpl}</div></div>
                        <div style="font-size:3rem;font-weight:900;color:{sc_color};">{j_score}%</div>
                    </div>''', unsafe_allow_html=True)

                    sched_rows_j = schedule_to_rows(j_sched, ready_memos_j)
                    # ترتيب زمني
                    sched_rows_j = sorted(sched_rows_j, key=lambda r: (r.get("اليوم","9999"), r.get("التوقيت","99:99"), r.get("القاعة","")))
                    df_sched_j = pd.DataFrame(sched_rows_j)

                    t_grid, t_prof, t_unpl_tab = st.tabs(["📋 الجدول الكامل","👨‍🏫 حسب الأستاذ","⚠️ غير مجدولة"])

                    with t_grid:
                        st.markdown("**🖊️ يمكنك التعديل مباشرة:**")
                        edited_j = st.data_editor(df_sched_j,
                            column_config={
                                "رقم المذكرة": st.column_config.TextColumn("رقم المذكرة", disabled=True, width="small"),
                                "العنوان": st.column_config.TextColumn("العنوان", disabled=True, width="large"),
                                "اليوم": st.column_config.SelectboxColumn("📆 اليوم", options=j_dy+["غير مجدول"], width="medium"),
                                "التوقيت": st.column_config.SelectboxColumn("🕐 التوقيت", options=j_sl+[""], width="small"),
                                "القاعة": st.column_config.SelectboxColumn("🏛️ القاعة", options=j_rm+[""], width="small"),
                                "المشرف": st.column_config.TextColumn("المشرف", disabled=True, width="medium"),
                                "الرئيس": st.column_config.TextColumn("الرئيس", disabled=True, width="medium"),
                                "المناقش1": st.column_config.TextColumn("مناقش1", disabled=True, width="medium"),
                                "المناقش2": st.column_config.TextColumn("مناقش2", disabled=True, width="medium"),
                                "رابط الملف": st.column_config.LinkColumn("📄", width="small"),
                            },
                            hide_index=True, use_container_width=True,
                            height=min(600, 60+len(df_sched_j)*38), key="j_editor")

                        # تحقق تعارضات
                        slot_groups_j = {}; conflicts_ui = []
                        for _, re_j in edited_j.iterrows():
                            if re_j["اليوم"]=="غير مجدول" or not re_j["التوقيت"]: continue
                            key_j = (re_j["اليوم"], re_j["التوقيت"])
                            profs_e = set(filter(None,[str(re_j.get(c,"")).strip() for c in ["المشرف","الرئيس","المناقش1","المناقش2"]]))-{"","nan"}
                            for pm,pp in slot_groups_j.get(key_j,[]):
                                if profs_e & pp: conflicts_ui.append(f"⚠️ **{re_j['رقم المذكرة']}** و **{pm}** — تعارض!")
                            slot_groups_j.setdefault(key_j,[]).append((re_j["رقم المذكرة"], profs_e))

                        if conflicts_ui:
                            for cf in conflicts_ui[:5]: st.error(cf)
                        else: st.success("✅ لا تعارضات")

                        if st.button("💾 حفظ التعديلات اليدوية", key="j_save_manual"):
                            new_j_sched = {}
                            for _, re_j in edited_j.iterrows():
                                m = str(re_j["رقم المذكرة"])
                                new_j_sched[m] = (re_j["اليوم"], re_j["التوقيت"], re_j["القاعة"]) if re_j["اليوم"]!="غير مجدول" and re_j["التوقيت"] else None
                            st.session_state["j_schedule"] = new_j_sched
                            st.success("✅ تم حفظ التعديلات محلياً"); st.rerun()

                    with t_prof:
                        memo_members_display = st.session_state.get("j_memo_members", {})
                        prof_set_j = set()
                        for mid, slot in j_sched.items():
                            if not slot: continue
                            for p in memo_members_display.get(mid, set()): prof_set_j.add(p)
                        for pj in sorted(prof_set_j):
                            pitems = []
                            for row_p in sched_rows_j:
                                if row_p["اليوم"]=="غير مجدول": continue
                                role_j = None
                                if str(row_p.get("المشرف","")).strip()==pj: role_j="مشرف"
                                elif str(row_p.get("الرئيس","")).strip()==pj: role_j="رئيس"
                                elif str(row_p.get("المناقش1","")).strip()==pj: role_j="مناقش1"
                                elif str(row_p.get("المناقش2","")).strip()==pj: role_j="مناقش2"
                                if role_j: pitems.append({**row_p,"الصفة":role_j})
                            if not pitems: continue
                            days_pj = set(i["اليوم"] for i in pitems)
                            with st.expander(f"👨‍🏫 {pj} — {len(pitems)} مناقشة في {len(days_pj)} يوم", expanded=False):
                                for dj, its in sorted({i["اليوم"]:[x for x in pitems if x["اليوم"]==i["اليوم"]] for i in pitems}.items()):
                                    st.markdown(f"**📅 {dj}**")
                                    for it in sorted(its, key=lambda x: x["التوقيت"]):
                                        st.markdown(f"&nbsp;&nbsp;🕐 **{it['التوقيت']}** | 🏛️ {it['القاعة']} | {it['رقم المذكرة']} | **{it['الصفة']}**")

                    with t_unpl_tab:
                        unpl_list = [r for r in sched_rows_j if r["اليوم"]=="غير مجدول"]
                        if not unpl_list:
                            st.success("✅ جميع المذكرات مجدولة!")
                        else:
                            st.warning(f"⚠️ {len(unpl_list)} مذكرة لم تُجدَّل")
                            # تحليل سبب عدم الجدولة لكل مذكرة
                            memo_map_unpl = {str(r["رقم المذكرة"]): r for _, r in ready_memos_j.iterrows()}
                            _rej_log_disp = st.session_state.get("j_rej_log", {})
                            for u in unpl_list:
                                mid = str(u["رقم المذكرة"])
                                reasons = _rej_log_disp.get(mid, set())
                                if reasons:
                                    # اختر أكثر الأسباب فائدة (الأولى مختلفة)
                                    unique_reasons = list(reasons)[:3]
                                    reason_str = " | ".join(unique_reasons)
                                else:
                                    reason_str = "الطاقة ممتلئة أو لا توجد خانة مناسبة"
                                st.markdown(f"- **{mid}** — {u.get('العنوان','')[:40]}")
                                st.caption(f"  🔍 السبب: {reason_str}")

                    st.markdown("---")
                    # تأكيد واعتماد
                    # ── ملخص القيود المطبقة ──
                    if st.session_state.get("j_constraints_summary"):
                        with st.expander("📋 القيود المطبقة في هذا الجدول", expanded=False):
                            for _c in st.session_state["j_constraints_summary"]:
                                st.markdown(f"- {_c}")

                    # ── تحقق صارم من الاستثناءات ──
                    _ban_violations = []
                    _sched_check = st.session_state.get("j_schedule", {})
                    _mbrs_check = st.session_state.get("j_memo_members", {})
                    for _mid, _sv in _sched_check.items():
                        if not _sv: continue
                        _day = _sv[0]
                        for _prof in _mbrs_check.get(_mid, set()):
                            if _day in st.session_state.get("_j_ban_days", {}).get(_prof, set()):
                                _ban_violations.append(f"🔴 {_prof}: مبرمج في يوم ممنوع {_day} (مذكرة {_mid})")
                            if st.session_state.get("_j_allow_days", {}).get(_prof) and _day not in st.session_state.get("_j_allow_days", {})[_prof]:
                                _ban_violations.append(f"🔴 {_prof}: مبرمج في يوم غير مسموح {_day} (مذكرة {_mid})")
                    if _ban_violations:
                        st.error(f"⚠️ {len(_ban_violations)} انتهاك للاستثناءات:")
                        for v in _ban_violations[:10]: st.markdown(f"- {v}")

                    # ── تقرير التحقق ──
                    st.markdown("---")
                    with st.expander("🔍 تقرير التحقق من الجدول", expanded=True):
                        _violations = validate_schedule(
                            st.session_state["j_schedule"],
                            st.session_state.get("j_memo_members", {}),
                            st.session_state.get("j_days_list", gen_days_j),
                            st.session_state.get("j_slots_list", gen_slots_j)
                        )
                        if not _violations:
                            st.success("✅ الجدول سليم — لا انتهاكات")
                        else:
                            red = [v for v in _violations if v.startswith("🔴")]
                            yellow = [v for v in _violations if v.startswith("🟡")]
                            if red:
                                st.error(f"🔴 {len(red)} انتهاك حرج:")
                                for v in red: st.markdown(f"- {v}")
                            if yellow:
                                st.warning(f"🟡 {len(yellow)} تنبيه:")
                                for v in yellow: st.markdown(f"- {v}")

                    # ── زر الحفظ بدون نشر ──
                    st.markdown("---")
                    col_save1, col_save2 = st.columns(2)
                    with col_save1:
                        if st.button("💾 حفظ في الشيت (بدون نشر)", use_container_width=True, key="j_save_no_publish"):
                            with st.spinner("⏳ جاري الحفظ..."):
                                _sched_to_save = st.session_state.get("j_schedule", {})
                                ok_s, msg_s = save_full_schedule_to_sheets(_sched_to_save, ready_memos_j)
                            if ok_s:
                                st.success("✅ تم الحفظ في الشيت — لم يُنشر بعد. غيّر عمود 'نشر البرنامج' يدوياً عند الاعتماد.")
                            else:
                                st.error(msg_s)
                    with col_save2:
                        st.info("💡 لنشر البرنامج: غيّر عمود **'نشر البرنامج'** إلى **نعم** في الشيت")

                    if not conflicts_ui:
                        st.markdown("---")
                        st.markdown('''<div style="background:linear-gradient(135deg,#0D2010,#0F2020);border:2px solid rgba(16,185,129,0.42);border-radius:18px;padding:20px 24px;margin:16px 0;">
                            <div style="font-size:1.1rem;font-weight:800;color:#10B981;margin-bottom:8px;">🎯 تأكيد واعتماد البرنامج النهائي</div>
                            <div style="color:#E2E8F0;font-size:0.85rem;">عند الضغط: ✅ حفظ في قاعدة البيانات + 📧 إيميل لكل أستاذ + 📱 إشعار للطلبة</div>
                        </div>''', unsafe_allow_html=True)

                        test_mode = st.checkbox("🧪 وضع التجربة (لا يُرسل ولا يُحفظ)", value=True, key="j_test_mode")

                        if not st.session_state.get("j_confirm_step", False):
                            btn_lbl = "🧪 معاينة النتائج" if test_mode else "✅ تأكيد واعتماد البرنامج"
                            if st.button(btn_lbl, type="primary", use_container_width=True, key="j_confirm_btn"):
                                st.session_state["j_confirm_step"] = True; st.rerun()
                        else:
                            st.warning("⚠️ هذا الإجراء نهائي ولا يمكن التراجع عنه.")
                            ca_j, cb_j = st.columns(2)
                            with ca_j:
                                if st.button("✅ نعم، اعتمد", type="primary", use_container_width=True, key="j_final_confirm"):
                                    final_j = st.session_state["j_schedule"]
                                    test_mode_final = st.session_state.get("j_test_mode", True)
                                    if test_mode_final:
                                        prof_prog_prev = {}
                                        memo_map_prev = {str(r["رقم المذكرة"]):r for _,r in ready_memos_j.iterrows()}
                                        for mf,sf in final_j.items():
                                            if not sf: continue
                                            rf = memo_map_prev.get(mf,{})
                                            for ck,rk in [("الأستاذ","مشرف"),("الرئيس","رئيس لجنة"),("المناقش1","مناقش 1"),("المناقش2","مناقش 2")]:
                                                pf = str(rf.get(ck,"")).strip() if hasattr(rf,"get") else ""
                                                if pf and pf not in ["","nan","—"]: prof_prog_prev.setdefault(pf,[]).append({"المذكرة":mf,"اليوم":sf[0],"التوقيت":sf[1],"القاعة":sf[2],"الصفة":rk})
                                        st.success(f"🧪 معاينة: سيُرسل لـ {len(prof_prog_prev)} أستاذ")
                                        with st.expander("👨‍🏫 معاينة البرامج"):
                                            for pf,prog in sorted(prof_prog_prev.items()):
                                                st.markdown(f"**{pf}** — {len(prog)} مناقشة")
                                        st.session_state["j_confirm_step"] = False
                                    else:
                                        with st.spinner("⏳ جاري الحفظ والإرسال..."):
                                            ok_j, msg_j = save_full_schedule_to_sheets(final_j, ready_memos_j)
                                            if ok_j:
                                                df_profs_j = load_prof_memos()
                                                df_st_j = load_students()
                                                df_st_j["رقم التسجيل_norm"] = df_st_j["رقم التسجيل"].astype(str).apply(normalize_text)
                                                memo_map_fin = {str(r["رقم المذكرة"]):r for _,r in ready_memos_j.iterrows()}
                                                prof_prog = {}
                                                for mf,sf in final_j.items():
                                                    if not sf: continue
                                                    rf = memo_map_fin.get(mf,{})
                                                    lnk_f = str(rf.get("رابط الملف","")).strip() if hasattr(rf,"get") else ""
                                                    for ck,rk in [("الأستاذ","مشرف"),("الرئيس","رئيس لجنة"),("المناقش1","مناقش 1"),("المناقش2","مناقش 2")]:
                                                        pf = str(rf.get(ck,"")).strip() if hasattr(rf,"get") else ""
                                                        if pf and pf not in ["","nan","—"]: prof_prog.setdefault(pf,[]).append({"رقم المذكرة":mf,"اليوم":sf[0],"التوقيت":sf[1],"القاعة":sf[2],"الصفة":rk,"رابط الملف":lnk_f})
                                                sent_p = 0
                                                for pf,prog_f in prof_prog.items():
                                                    ok_p,_ = send_prof_schedule_email(pf, prog_f, df_profs_j)
                                                    if ok_p: sent_p += 1
                                                    time_module.sleep(0.3)
                                                sent_s = 0
                                                for mf,sf in final_j.items():
                                                    if not sf: continue
                                                    rf = memo_map_fin.get(mf,{})
                                                    if not hasattr(rf,"get"): continue
                                                    rv = rf.tolist() if hasattr(rf,"tolist") else []
                                                    r1f = normalize_text(rf.get("رقم تسجيل الطالب 1", rv[16] if len(rv)>16 else ""))
                                                    r2f = normalize_text(rf.get("رقم تسجيل الطالب 2", rv[17] if len(rv)>17 else ""))
                                                    for reg_f in [r1f, r2f]:
                                                        if not reg_f or reg_f in ["","nan"]: continue
                                                        sr = df_st_j[df_st_j["رقم التسجيل_norm"]==reg_f]
                                                        if sr.empty: continue
                                                        ok_s,_ = send_student_schedule_email(sr.iloc[0].to_dict(), mf, sf[0], sf[1], sf[2])
                                                        if ok_s: sent_s += 1
                                                        time_module.sleep(0.2)
                                                for k in ["j_schedule","j_score","j_unplaced","j_confirm_step"]:
                                                    st.session_state.pop(k,None)
                                                st.success(f"🎉 تم! حُفظ | أُرسل لـ {sent_p} أستاذ | أُشعر {sent_s} طالب")
                                                st.balloons(); clear_cache_and_reload(); time_module.sleep(2); st.rerun()
                                            else: st.error(msg_j)
                            with cb_j:
                                if st.button("إلغاء", use_container_width=True, key="j_cancel"):
                                    st.session_state["j_confirm_step"] = False; st.rerun()
                    else:
                        st.error("⛔ صحّح التعارضات أولاً.")

        # ================================================================
        # TAB إحصائيات الأساتذة
        # ================================================================
        # ================================================================
        # TAB إرسال إيميلات اللجان
        # ================================================================
        # نضعها كقسم منفصل داخل تاب الجدولة

